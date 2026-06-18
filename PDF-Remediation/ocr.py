"""OCR processing functions."""

import io

import fitz
import pytesseract
from PIL import Image
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


OCR_DISCLAIMER = (
    "This document was converted from a scanned image to readable text using "
    "Tesseract OCR and Claude AI. While every effort has been made to ensure "
    "accuracy, some errors may remain. Please review the content before distributing."
)


def run_ocr(pil_images, tesseract_path):
    """
    Extract text from pre-processed page images using Tesseract OCR.

    Accepts PIL Images that have already been rendered and corrected by
    preprocess_pages() — orientation and skew fixes are applied before
    this function is called, so it only handles the OCR step itself.

    Approach:
      1. Configure pytesseract to use the binary at tesseract_path.
      2. Run image_to_string() on each PIL Image in order.
         The images are assumed to be 300 DPI — that resolution is set
         by preprocess_pages() when it renders pages from the PDF.
      3. Return the extracted text as a list — one string per page.

    Parameters:
        pil_images (list[PIL.Image]): Pre-processed page images from
                                     preprocess_pages(), in document order.
        tesseract_path (str): Full path to the tesseract executable
                              (e.g. C:\\Program Files\\Tesseract-OCR\\tesseract.exe).

    Returns:
        list[str]: Extracted text for each page, in document order.

    Raises:
        Exception: Re-raised if Tesseract cannot be found or fails to run,
                   so the caller can show a user-facing error message.
    """
    # Tell pytesseract where the Tesseract binary lives.
    # This must be set before any image_to_string() call.
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    # lang="eng" tells Tesseract to use the English language model.
    # In a future iteration this could be made configurable for multilingual documents.
    return [pytesseract.image_to_string(img, lang="eng") for img in pil_images]


def build_docx(pages_text):
    """
    Assemble OCR-extracted page texts into a DOCX file.

    Each page's text is split into lines. Non-empty lines become paragraphs.
    A page break is inserted between pages to preserve the original pagination.
    The OCR disclaimer is appended on a final page per CLAUDE.md OCR Step 6.

    Parameters:
        pages_text (list[str]): One text string per page from run_ocr().

    Returns:
        bytes: The completed DOCX file as raw bytes, ready for st.download_button.
    """
    doc = DocxDocument()

    for i, page_text in enumerate(pages_text):
        if i > 0:
            # Insert a page break before each page after the first.
            # add_page_break() adds a run with a page break character to a
            # new paragraph — this is the standard python-docx approach.
            doc.add_page_break()

        for line in page_text.splitlines():
            stripped = line.strip()
            if stripped:
                doc.add_paragraph(stripped)

    # Append the conversion disclaimer on its own page (CLAUDE.md OCR Step 6).
    doc.add_page_break()
    disclaimer_para = doc.add_paragraph(OCR_DISCLAIMER)
    # Italicize the disclaimer to visually distinguish it from document content.
    disclaimer_para.runs[0].italic = True

    # Write the document to an in-memory buffer and return the bytes.
    # We never write to disk — the bytes go straight to Streamlit's download.
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_pdf(pages_text):
    """
    Assemble OCR-extracted page texts into a plain-text PDF using ReportLab.

    ReportLab's Platypus layout engine is used to flow text onto letter-size
    pages with standard margins. Each original page's content is separated by
    a PageBreak. The OCR disclaimer is appended at the end.

    Parameters:
        pages_text (list[str]): One text string per page from run_ocr().

    Returns:
        bytes: The completed PDF file as raw bytes, ready for st.download_button.
    """
    output = io.BytesIO()

    # SimpleDocTemplate manages page layout, margins, and page breaks.
    doc_rl = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    # getSampleStyleSheet() returns ReportLab's built-in paragraph styles.
    # "Normal" is plain body text; "Italic" is used for the disclaimer.
    styles = getSampleStyleSheet()
    story = []  # "story" is ReportLab's term for the list of content blocks

    for i, page_text in enumerate(pages_text):
        if i > 0:
            story.append(PageBreak())

        for line in page_text.splitlines():
            stripped = line.strip()
            if stripped:
                # ReportLab's Paragraph parses text as XML, so < > & characters
                # from OCR output must be escaped or they crash the renderer.
                import html as _html
                safe = _html.escape(stripped)
                story.append(Paragraph(safe, styles["Normal"]))
                story.append(Spacer(1, 4))

    # Append the conversion disclaimer on its own page (CLAUDE.md OCR Step 6).
    story.append(PageBreak())
    story.append(Paragraph(OCR_DISCLAIMER, styles["Italic"]))

    # build() renders all the story elements into the output buffer.
    doc_rl.build(story)
    return output.getvalue()


# =============================================================================
# OCR QUALITY SCORING
# =============================================================================
# score_ocr_quality() measures how "word-like" the extracted text is.
# Gibberish OCR output tends to have a high proportion of tokens that are
# not real words — random symbol runs, single stray characters, sequences
# of digits mixed with letters, etc.
#
# Approach: split text into whitespace-separated tokens and count what
# fraction are "clean words" — tokens made entirely of letters, apostrophes,
# or hyphens, with length between 2 and 30 characters.
#
# Returns a float 0–100 (percentage) and a label: "Good", "Fair", or "Poor".
# Displayed in the UI as ✅ / ⚠️ / ❌ — intentionally different from the
# 🔴/🟡/🟢 accessibility severity scheme used elsewhere in the app.
# =============================================================================

def score_ocr_quality(pages_text):
    """
    Estimate OCR output quality as a percentage of word-like tokens.

    Parameters:
        pages_text (list[str]): One string per page from run_ocr().

    Returns:
        (float, str): Score 0–100 and a label ("Good", "Fair", or "Poor").
    """
    import re
    full_text = " ".join(pages_text)
    tokens = full_text.split()
    if not tokens:
        return 0.0, "Poor"

    word_pattern = re.compile(r"^[A-Za-z][A-Za-z'\-]{1,29}$")
    clean = sum(1 for t in tokens if word_pattern.match(t))
    score = (clean / len(tokens)) * 100

    if score >= 72:
        label = "Good"
    elif score >= 45:
        label = "Fair"
    else:
        label = "Poor"

    return round(score, 1), label


# =============================================================================
# EASYOCR FALLBACK
# =============================================================================
# run_easyocr() is the second-pass OCR engine. It uses a neural-network model
# (CRAFT text detector + CRNN recogniser) which handles skewed, low-contrast,
# and unusual-font scans better than Tesseract's pattern matching.
#
# EasyOCR is only invoked when:
#   a) Tesseract quality was flagged as Fair or Poor, AND
#   b) The faculty member explicitly chooses to run it.
#
# On first call EasyOCR downloads ~200 MB of model weights. Subsequent calls
# reuse the cached models. gpu=False ensures it runs on any machine.
# =============================================================================

def run_easyocr(pil_images):
    """
    Extract text from pre-processed page images using EasyOCR.

    Accepts PIL Images that have already been rendered and corrected by
    preprocess_pages() — orientation and skew fixes are applied before
    this function is called, so it only handles the OCR step itself.

    Uses a neural-network model (CRAFT text detector + CRNN recogniser)
    which handles skewed, low-contrast, and unusual-font scans better than
    Tesseract's pattern matching. Results are sorted top-to-bottom by
    bounding-box Y coordinate to approximate visual reading order.

    EasyOCR and numpy are imported inside the function because they are
    optional dependencies — importing them lazily avoids import errors on
    machines where EasyOCR is not installed.

    On first call EasyOCR downloads ~200 MB of model weights. Subsequent
    calls reuse the cached models. gpu=False ensures it runs on any machine.

    Parameters:
        pil_images (list[PIL.Image]): Pre-processed page images from
                                     preprocess_pages(), in document order.

    Returns:
        list[str]: Extracted text for each page, in document order.

    Raises:
        Exception: Re-raised so the caller can show a user-facing error.
    """
    import easyocr
    import numpy as np

    reader = easyocr.Reader(["en"], gpu=False, verbose=False)
    pages_text = []

    for img in pil_images:
        # Convert PIL Image to a numpy array for EasyOCR.
        # Strip the alpha channel if present — EasyOCR expects RGB only.
        img_array = np.array(img)
        if img_array.ndim == 3 and img_array.shape[2] == 4:
            img_array = img_array[:, :, :3]

        results = reader.readtext(img_array)

        # Sort detections top-to-bottom by the top-left Y coordinate of the
        # bounding box so text flows in visual reading order.
        results.sort(key=lambda r: r[0][0][1])
        pages_text.append(" ".join(r[1] for r in results))

    return pages_text


# =============================================================================
# PAGE PREPROCESSING — ORIENTATION & SKEW CORRECTION
# =============================================================================
# preprocess_pages() runs before any OCR engine. It renders each PDF page as
# a PIL Image at 300 DPI, then applies two corrections automatically:
#
#   1. Orientation correction — Tesseract OSD detects if the page was scanned
#      sideways (90°), upside-down (180°), or at 270°, and rotates it upright.
#
#   2. Skew correction — OpenCV's minAreaRect finds the dominant angle of text
#      lines and applies a small rotation to straighten them. Only applied when
#      the detected angle is more than 0.5° (below that it's measurement noise).
#
# Returns the corrected PIL Images and a human-readable processing log that is
# shown in the "What we did" summary after OCR completes.
# =============================================================================

def _deskew_image(pil_image):
    """
    Detect and correct slight skew in a PIL Image using OpenCV.

    Converts to grayscale, thresholds to find text pixels, fits a minimum
    bounding rectangle to those pixels, and rotates by the detected angle.

    Parameters:
        pil_image (PIL.Image): The page image to deskew.

    Returns:
        (PIL.Image, float): Corrected image and the angle applied (degrees).
                            Returns (original_image, 0.0) if skew is negligible
                            or detection fails.
    """
    import cv2
    import numpy as np

    try:
        gray = np.array(pil_image.convert("L"))
        _, thresh = cv2.threshold(
            gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU
        )
        coords = np.column_stack(np.where(thresh > 0))
        if len(coords) < 100:   # too few text pixels to measure reliably
            return pil_image, 0.0

        angle = cv2.minAreaRect(coords)[-1]
        # minAreaRect returns angles in [-90, 0). Normalise to a signed value.
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle

        if abs(angle) < 0.5:    # below threshold — not worth correcting
            return pil_image, 0.0

        (h, w) = gray.shape[:2]
        M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
        img_array = np.array(pil_image)
        rotated = cv2.warpAffine(
            img_array, M, (w, h),
            flags=cv2.INTER_CUBIC,
            borderMode=cv2.BORDER_REPLICATE,
        )
        return Image.fromarray(rotated), angle
    except Exception:
        return pil_image, 0.0


def preprocess_pages(file_bytes, tesseract_path):
    """
    Render every PDF page at 300 DPI and apply orientation and skew correction.

    Called once before running Tesseract. The corrected PIL Images are passed
    directly to the OCR engine so the same preprocessing applies whether we
    run Tesseract or EasyOCR.

    Parameters:
        file_bytes (bytes): Raw bytes of the uploaded PDF.
        tesseract_path (str): Path to the Tesseract binary (needed for OSD).

    Returns:
        (list[PIL.Image], list[str]):
            - Corrected page images in document order.
            - Human-readable log entries describing what was done to each page.
    """
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page_count = len(doc)
    pil_images = []
    log_entries = [f"✔ {page_count} page{'s' if page_count != 1 else ''} detected"]

    pages_rotated = 0
    pages_deskewed = 0

    for page_num in range(page_count):
        page = doc[page_num]
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        # Step 1 — Orientation correction via Tesseract OSD.
        # OSD returns the clockwise rotation needed to make text upright.
        # We skip silently if OSD fails (low-text pages confuse it).
        try:
            osd = pytesseract.image_to_osd(
                img, output_type=pytesseract.Output.DICT
            )
            rotate = osd.get("rotate", 0)
            if rotate and rotate != 0:
                # PIL.rotate is counter-clockwise; OSD gives clockwise degrees.
                img = img.rotate(-rotate, expand=True)
                log_entries.append(
                    f"✔ Page {page_num + 1}: orientation corrected (rotated {rotate}°)"
                )
                pages_rotated += 1
        except Exception:
            pass

        # Step 2 — Skew correction via OpenCV.
        img, skew_angle = _deskew_image(img)
        if abs(skew_angle) > 0.5:
            log_entries.append(
                f"✔ Page {page_num + 1}: skew corrected ({skew_angle:+.1f}°)"
            )
            pages_deskewed += 1

        pil_images.append(img)

    doc.close()

    if pages_rotated == 0 and pages_deskewed == 0:
        log_entries.append("✔ All pages: orientation and alignment look correct")

    return pil_images, log_entries
