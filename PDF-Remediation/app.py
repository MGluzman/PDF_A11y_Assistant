# =============================================================================
# app.py — PDF Assistant (Streamlit)
# =============================================================================
# Purpose:
#   A step-by-step conversational tool that helps faculty identify and fix
#   accessibility barriers in PDF course materials to comply with:
#     - WCAG 2.2 Level AA (Web Content Accessibility Guidelines)
#     - Title II of the Americans with Disabilities Act (ADA)
#       as amended by DOJ regulations effective June 2024
#
# Architecture:
#   This is a Streamlit state machine. Streamlit re-runs this entire script
#   from top to bottom every time the user clicks a button or uploads a file.
#   We use st.session_state to remember where in the workflow the user is,
#   and what data has been collected so far.
#
#   The "step" key in session_state drives which screen is shown.
#   Each step is handled by a dedicated render_*() function below.
#
# How to run:
#   streamlit run app.py
# =============================================================================

import io                    # In-memory byte streams for building files without writing to disk
import os                    # File system operations — used for temp file cleanup in Docling helper
import re                    # Regular expressions — used for detecting visual list patterns in text
import shutil                # shutil.which() — locate executables on the system PATH
import subprocess            # Run Tesseract as a subprocess to verify it is callable
import tempfile              # Creates temporary files that Docling's converter can read by path
import time                  # Used for brief delays during state transitions
import urllib.request        # Lightweight HTTP HEAD requests for API endpoint health checks
from collections import Counter  # Frequency counting for font-size heuristics
import fitz                  # PyMuPDF — opens PDFs, renders pages as images, checks password
import pikepdf               # Low-level PDF editing — writes metadata and structural fixes
import pytesseract           # Python wrapper around the Tesseract OCR binary
import streamlit as st       # The main framework — builds the entire UI
from langdetect import detect as detect_language  # Detects the written language of extracted text
from PIL import Image        # Pillow — convert PyMuPDF pixel data to images pytesseract can read
from docx import Document as DocxDocument   # python-docx — builds DOCX output files
from docx.shared import Pt                  # Point sizes for DOCX styles
from reportlab.lib.pagesizes import letter  # ReportLab — builds PDF output files
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from docling.document_converter import DocumentConverter  # Docling — ML-powered document structure analysis
from docling_core.types.doc import DocItemLabel           # Semantic element type labels (heading, list, table, etc.)


# ---------------------------------------------------------------------------
# Module-level constants for the PDF analysis pipeline.
# ---------------------------------------------------------------------------

# Compiled once at import time — regex for visually formatted list lines.
LIST_PATTERN = re.compile(
    r"^\s*([•·▪▸➢◦○●\-\*]|\d{1,2}[.)]\s|[a-zA-Z][.]\s)\s*\S"
)

# Font name substrings identifying condensed typefaces — excluded from the
# letter-spacing heuristic because their narrow widths are intentional design.
CONDENSED_FONT_HINTS = frozenset({"condense", "narrow", "compress", "condens"})


# -----------------------------------------------------------------------------
# PAGE CONFIGURATION
# Must be the FIRST Streamlit call in the script. Sets the browser tab title,
# icon, and layout. "wide" uses the full browser width — better for documents.
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="PDF Assistant",
    page_icon="♿",
    layout="wide",
    initial_sidebar_state="expanded",
)


# =============================================================================
# DEMO ISSUE DATA
# =============================================================================
# Fallback issue list returned by analyze_pdf() when PDF parsing fails entirely.
# Keeps the app functional with a recognizable set of issues rather than crashing.
#
# Each issue is a dictionary with:
#   id          — unique identifier used to track resolution
#   severity    — "red", "yellow", or "green" (drives display and ordering)
#   title       — short label shown in the issue list
#   description — explanation of why this matters for students
#   wcag        — WCAG 2.2 Success Criterion reference
#   ada         — Title II / 28 CFR Part 35 reference
#   fix_preview — sample text shown during the QA confirmation step
# =============================================================================
DEMO_ISSUES = [
    {
        "id": "missing_alt_text",
        "severity": "red",
        "title": "Images with missing text descriptions",
        "description": (
            "This document contains images that don't have text descriptions (also called "
            "alt text). Students who use screen readers — software that reads content aloud "
            "— won't be able to access the information those images convey."
        ),
        "wcag": "WCAG 2.2 SC 1.1.1 — Non-text Content",
        "ada": "Title II ADA, 28 CFR Part 35",
        "fix_preview": (
            "I've reviewed each image in your document and generated a text description "
            "for each one. I'll walk you through them one at a time so you can review "
            "and approve each description before it's saved."
        ),
    },
    {
        "id": "untagged_pdf",
        "severity": "red",
        "title": "Untagged PDF — no document structure",
        "description": (
            "This document has no accessibility tags. Tags are the behind-the-scenes "
            "structure that tells screen readers what is a heading, what is a paragraph, "
            "and what order to read things in. Without them, a screen reader user hears "
            "the content as an undifferentiated stream of text — or nothing at all."
        ),
        "wcag": "WCAG 2.2 SC 1.3.1 — Info and Relationships",
        "ada": "Title II ADA, 28 CFR Part 35",
        "fix_preview": (
            "I've analyzed the visual layout of your document and generated a tag "
            "structure that reflects the correct reading order and heading hierarchy. "
            "Here's a summary of the structure I'll apply."
        ),
    },
    {
        "id": "heading_hierarchy",
        "severity": "yellow",
        "title": "Incorrect or missing heading hierarchy",
        "description": (
            "The heading levels in this document are inconsistent — some levels are "
            "skipped, and the structure doesn't match what a reader would expect from "
            "the visual layout. Screen reader users navigate long documents by jumping "
            "between headings, so an incorrect hierarchy makes it hard to find content."
        ),
        "wcag": "WCAG 2.2 SC 1.3.1 — Info and Relationships",
        "ada": "Title II ADA, 28 CFR Part 35",
        "fix_preview": (
            "I've mapped your document's visual layout to a correct heading hierarchy "
            "(Heading 1 → Heading 2 → Heading 3) with no skipped levels. Here's the "
            "proposed structure."
        ),
    },
    {
        "id": "missing_lang",
        "severity": "yellow",
        "title": "Missing document language tag",
        "description": (
            "This document doesn't declare what language it's written in. Screen readers "
            "use the language tag to select the correct pronunciation rules and speech "
            "engine. Without it, text may be read aloud in the wrong language or with "
            "incorrect pronunciation — especially for technical terms and proper names."
        ),
        "wcag": "WCAG 2.2 SC 3.1.1 — Language of Page",
        "ada": "Title II ADA, 28 CFR Part 35",
        "fix_preview": (
            "I've detected that this document is written in English and will add the "
            "language tag 'en-US' to the document metadata."
        ),
    },
    {
        "id": "missing_title",
        "severity": "green",
        "title": "Missing document title in metadata",
        "description": (
            "The document's title field is empty — it currently defaults to the filename. "
            "Screen readers announce the document title when a file opens, and search "
            "tools use it for indexing. A descriptive title helps all users understand "
            "what they're looking at before they start reading."
        ),
        "wcag": "WCAG 2.2 SC 2.4.2 — Page Titled",
        "ada": "Title II ADA, 28 CFR Part 35",
        "fix_preview": (
            "I'll add a descriptive title to the document metadata based on the content "
            "I found at the top of the first page. You can review and edit it before I save."
        ),
    },
    {
        "id": "missing_bookmarks",
        "severity": "green",
        "title": "Missing bookmarks / navigation",
        "description": (
            "This document is long enough that it would benefit from bookmarks — a "
            "clickable table of contents that lets readers jump directly to any section. "
            "Without bookmarks, everyone (not just screen reader users) has to scroll "
            "through the entire document to find what they need."
        ),
        "wcag": "WCAG 2.2 SC 2.4.5 — Multiple Ways",
        "ada": "Title II ADA, 28 CFR Part 35",
        "fix_preview": (
            "I've generated a bookmark list based on the heading structure in your "
            "document. Here's what it will look like in the navigation panel."
        ),
    },
]

# Labels and colors for each severity tier.
# We always pair the emoji icon with a text label (never icon alone).
# This is required by WCAG 2.2 SC 1.4.1 — Color Not Used as Only Visual Means.
SEVERITY_LABELS = {
    "red":    "🔴 Red Light",
    "yellow": "🟡 Yellow Light",
    "green":  "🟢 Green Light",
}


# Disclaimer appended to every OCR-converted file per CLAUDE.md OCR Step 6.
# Informs recipients that the text was machine-extracted and may contain errors.
OCR_DISCLAIMER = (
    "This document was converted from a scanned image to readable text using "
    "Tesseract OCR and Claude AI. While every effort has been made to ensure "
    "accuracy, some errors may remain. Please review the content before distributing."
)


# =============================================================================
# PREFLIGHT CHECK
# =============================================================================
# preflight_check() is called at the start of every pipeline that touches
# external resources. It catches three classes of failure before the pipeline
# gets underway so the user sees a clear, actionable message instead of a
# cryptic traceback mid-process.
#
# The three checks:
#   1. Tesseract — is the binary on PATH and does it respond to --version?
#   2. API endpoints — do all required HTTP endpoints return a 2xx status?
#      (This list is empty today; add URLs here when external services are added.)
#   3. Output paths — can the app write to every path it will need to use?
#      Passing tempfile.gettempdir() catches a full or read-only temp directory
#      before Docling tries to create its temporary PDF and fails silently.
#
# Return value:
#   {"ok": bool, "errors": list[str]}
#   "ok" is True only when all requested checks pass.
#   "errors" contains one plain-English message for every failed check.
#   Callers should display every message in the list, then stop the pipeline.
# =============================================================================

def preflight_check(tesseract_path=None, api_endpoints=None, output_paths=None):
    """
    Verify external dependencies before starting a pipeline.

    Parameters:
        tesseract_path (str | None):
            Full path to the Tesseract executable (e.g. the value from the
            sidebar settings field). When provided, the function checks that
            Tesseract can be found — either at this exact path OR somewhere
            on the system PATH — and that it actually responds to --version.
            Pass None to skip the Tesseract check (e.g. when Tesseract is
            not needed for the pipeline being started).

        api_endpoints (list[str] | None):
            HTTP/HTTPS URLs that must return a 2xx status before the pipeline
            can proceed. The check uses a lightweight HEAD request with a
            5-second timeout. Pass None or an empty list to skip.

        output_paths (list[str] | None):
            File-system paths the pipeline will write to. The check opens each
            path in append-binary mode to confirm write access and that no other
            process holds an exclusive lock. Pass None to skip.
            Tip: always include tempfile.gettempdir() for pipelines that use
            Docling, since Docling writes a temporary PDF there.

    Returns:
        dict with keys:
            "ok"     (bool)       — True if every requested check passed
            "errors" (list[str])  — one human-readable message per failure
    """
    errors = []

    # ------------------------------------------------------------------
    # Check 1 — Tesseract is installed and callable
    # ------------------------------------------------------------------
    # Why two sub-checks?
    #   shutil.which("tesseract") finds the binary through the system PATH.
    #   The configured path in the sidebar may differ from the PATH entry
    #   (e.g. a custom install location on Windows). We accept either one
    #   as proof that Tesseract is available, but we always try to run the
    #   configured path so the subprocess call later will actually work.
    # ------------------------------------------------------------------
    if tesseract_path is not None:
        # Sub-check A: is "tesseract" discoverable on the system PATH?
        on_path = shutil.which("tesseract") is not None

        # Sub-check B: does the configured path respond to --version?
        configured_ok = False
        try:
            result = subprocess.run(
                [tesseract_path, "--version"],
                capture_output=True,    # suppress console output
                timeout=5,              # bail out after 5 seconds
            )
            # returncode 0 means Tesseract ran successfully.
            configured_ok = result.returncode == 0
        except FileNotFoundError:
            # The path doesn't point to any file at all.
            pass
        except PermissionError:
            # The file exists but the OS won't let us execute it.
            pass
        except subprocess.TimeoutExpired:
            # Tesseract started but didn't respond in time.
            pass

        if not on_path and not configured_ok:
            # Neither location works — Tesseract is not usable.
            errors.append(
                f"Tesseract OCR could not be found. "
                f"Verify that Tesseract is installed and that the path in the "
                f"sidebar settings is correct.\n"
                f"Configured path: `{tesseract_path}`\n"
                f"Download Tesseract: https://github.com/UB-Mannheim/tesseract/wiki"
            )
        elif not configured_ok:
            # Found on PATH but the configured path is wrong — tell the user
            # exactly which setting to fix.
            errors.append(
                f"Tesseract was found on the system PATH but could not be run "
                f"from the path entered in sidebar settings.\n"
                f"Configured path: `{tesseract_path}`\n"
                f"Please update the Tesseract path in the ⚙️ Settings panel."
            )

    # ------------------------------------------------------------------
    # Check 2 — Required API endpoints return 2xx
    # ------------------------------------------------------------------
    # Currently this list is empty — the app uses no external APIs.
    # Add endpoint URLs here when external services (e.g. Claude API health
    # check, a storage backend) are introduced.
    # ------------------------------------------------------------------
    if api_endpoints:
        for url in api_endpoints:
            try:
                req = urllib.request.Request(url, method="HEAD")
                with urllib.request.urlopen(req, timeout=5) as resp:
                    # Any status in the 400–599 range is a failure.
                    if resp.status >= 400:
                        errors.append(
                            f"Required API endpoint returned HTTP {resp.status}: {url}\n"
                            f"The service may be down or misconfigured."
                        )
            except urllib.error.URLError as exc:
                errors.append(
                    f"Required API endpoint is unreachable: {url}\n"
                    f"Network error: {exc.reason}"
                )
            except Exception as exc:
                errors.append(
                    f"Could not check API endpoint: {url}\n"
                    f"Error: {exc}"
                )

    # ------------------------------------------------------------------
    # Check 3 — Output paths are writable and not locked
    # ------------------------------------------------------------------
    # On Windows, opening a file in 'a+b' (append-binary) mode raises
    # PermissionError if another process holds an exclusive write lock.
    # On all platforms it raises OSError if the directory doesn't exist
    # or the file system is read-only.
    #
    # Special case: callers may pass a directory (e.g. tempfile.gettempdir())
    # to confirm the temp folder is writable. On Windows, open() on a directory
    # path always raises PermissionError regardless of actual permissions, so we
    # probe directories by creating and immediately deleting a small file inside.
    # ------------------------------------------------------------------
    if output_paths:
        for path in output_paths:
            if os.path.isdir(path):
                # Probe a directory by writing a temporary file inside it.
                probe = os.path.join(path, ".preflight_write_probe")
                try:
                    with open(probe, "a+b"):
                        pass
                    os.remove(probe)
                except PermissionError:
                    errors.append(
                        f"Output path is locked by another process or the app "
                        f"does not have write permission: `{path}`\n"
                        f"Close any applications that may have this file open and try again."
                    )
                except OSError as exc:
                    errors.append(
                        f"Output path is not writable: `{path}`\n"
                        f"OS error: {exc}"
                    )
            else:
                try:
                    # Open in append-binary mode so we don't truncate any
                    # existing content — we only need to confirm write access.
                    with open(path, "a+b"):
                        pass
                except PermissionError:
                    errors.append(
                        f"Output path is locked by another process or the app "
                        f"does not have write permission: `{path}`\n"
                        f"Close any applications that may have this file open and try again."
                    )
                except OSError as exc:
                    errors.append(
                        f"Output path is not writable: `{path}`\n"
                        f"OS error: {exc}"
                    )

    return {"ok": len(errors) == 0, "errors": errors}


# =============================================================================
# OCR HELPER FUNCTIONS
# =============================================================================
# These functions handle the OCR pipeline:
#   run_ocr()      — runs Tesseract on pre-processed PIL Images, returns text per page
#   run_easyocr()  — runs EasyOCR on pre-processed PIL Images, returns text per page
#   build_docx()   — assembles extracted text into a downloadable DOCX file
#   build_pdf()    — assembles extracted text into a downloadable PDF file
#
# Page rendering and preprocessing (300 DPI, orientation + skew correction) is
# handled upstream by preprocess_pages() before either OCR function is called.
#
# Both build_* functions append the OCR_DISCLAIMER at the end per CLAUDE.md.
# Both return raw bytes so Streamlit's st.download_button can serve them
# directly without writing anything to disk.
# =============================================================================

def run_ocr(pil_images, tesseract_path):
    """
    Extract text from pre-processed page images using Tesseract OCR.

    Accepts PIL Images that have already been rendered and corrected by
    preprocess_pages() — orientation and skew fixes are applied before
    this function is called, so it only handles the OCR step itself.

    Approach:
      1. Configure pytesseract to use the binary at tesseract_path.
      2. Run image_to_string() on each PIL Image in order.
         The images are assumed to be 300 DPI — that resolution is set
         by preprocess_pages() when it renders pages from the PDF.
      3. Return the extracted text as a list — one string per page.

    Parameters:
        pil_images (list[PIL.Image]): Pre-processed page images from
                                     preprocess_pages(), in document order.
        tesseract_path (str): Full path to the tesseract executable
                              (e.g. C:\\Program Files\\Tesseract-OCR\\tesseract.exe).

    Returns:
        list[str]: Extracted text for each page, in document order.

    Raises:
        Exception: Re-raised if Tesseract cannot be found or fails to run,
                   so the caller can show a user-facing error message.
    """
    # Tell pytesseract where the Tesseract binary lives.
    # This must be set before any image_to_string() call.
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    # lang="eng" tells Tesseract to use the English language model.
    # In a future iteration this could be made configurable for multilingual documents.
    return [pytesseract.image_to_string(img, lang="eng") for img in pil_images]


def build_docx(pages_text):
    """
    Assemble OCR-extracted page texts into a DOCX file.

    Each page's text is split into lines. Non-empty lines become paragraphs.
    A page break is inserted between pages to preserve the original pagination.
    The OCR disclaimer is appended on a final page per CLAUDE.md OCR Step 6.

    Parameters:
        pages_text (list[str]): One text string per page from run_ocr().

    Returns:
        bytes: The completed DOCX file as raw bytes, ready for st.download_button.
    """
    doc = DocxDocument()

    for i, page_text in enumerate(pages_text):
        if i > 0:
            # Insert a page break before each page after the first.
            # add_page_break() adds a run with a page break character to a
            # new paragraph — this is the standard python-docx approach.
            doc.add_page_break()

        for line in page_text.splitlines():
            stripped = line.strip()
            if stripped:
                doc.add_paragraph(stripped)

    # Append the conversion disclaimer on its own page (CLAUDE.md OCR Step 6).
    doc.add_page_break()
    disclaimer_para = doc.add_paragraph(OCR_DISCLAIMER)
    # Italicize the disclaimer to visually distinguish it from document content.
    disclaimer_para.runs[0].italic = True

    # Write the document to an in-memory buffer and return the bytes.
    # We never write to disk — the bytes go straight to Streamlit's download.
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_pdf(pages_text):
    """
    Assemble OCR-extracted page texts into a plain-text PDF using ReportLab.

    ReportLab's Platypus layout engine is used to flow text onto letter-size
    pages with standard margins. Each original page's content is separated by
    a PageBreak. The OCR disclaimer is appended at the end.

    Parameters:
        pages_text (list[str]): One text string per page from run_ocr().

    Returns:
        bytes: The completed PDF file as raw bytes, ready for st.download_button.
    """
    output = io.BytesIO()

    # SimpleDocTemplate manages page layout, margins, and page breaks.
    doc_rl = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    # getSampleStyleSheet() returns ReportLab's built-in paragraph styles.
    # "Normal" is plain body text; "Italic" is used for the disclaimer.
    styles = getSampleStyleSheet()
    story = []  # "story" is ReportLab's term for the list of content blocks

    for i, page_text in enumerate(pages_text):
        if i > 0:
            story.append(PageBreak())

        for line in page_text.splitlines():
            stripped = line.strip()
            if stripped:
                # ReportLab's Paragraph parses text as XML, so < > & characters
                # from OCR output must be escaped or they crash the renderer.
                import html as _html
                safe = _html.escape(stripped)
                story.append(Paragraph(safe, styles["Normal"]))
                story.append(Spacer(1, 4))

    # Append the conversion disclaimer on its own page (CLAUDE.md OCR Step 6).
    story.append(PageBreak())
    story.append(Paragraph(OCR_DISCLAIMER, styles["Italic"]))

    # build() renders all the story elements into the output buffer.
    doc_rl.build(story)
    return output.getvalue()


# =============================================================================
# OCR QUALITY SCORING
# =============================================================================
# score_ocr_quality() measures how "word-like" the extracted text is.
# Gibberish OCR output tends to have a high proportion of tokens that are
# not real words — random symbol runs, single stray characters, sequences
# of digits mixed with letters, etc.
#
# Approach: split text into whitespace-separated tokens and count what
# fraction are "clean words" — tokens made entirely of letters, apostrophes,
# or hyphens, with length between 2 and 30 characters.
#
# Returns a float 0–100 (percentage) and a label: "Good", "Fair", or "Poor".
# Displayed in the UI as ✅ / ⚠️ / ❌ — intentionally different from the
# 🔴/🟡/🟢 accessibility severity scheme used elsewhere in the app.
# =============================================================================

def score_ocr_quality(pages_text):
    """
    Estimate OCR output quality as a percentage of word-like tokens.

    Parameters:
        pages_text (list[str]): One string per page from run_ocr().

    Returns:
        (float, str): Score 0–100 and a label ("Good", "Fair", or "Poor").
    """
    import re
    full_text = " ".join(pages_text)
    tokens = full_text.split()
    if not tokens:
        return 0.0, "Poor"

    word_pattern = re.compile(r"^[A-Za-z][A-Za-z'\-]{1,29}$")
    clean = sum(1 for t in tokens if word_pattern.match(t))
    score = (clean / len(tokens)) * 100

    if score >= 72:
        label = "Good"
    elif score >= 45:
        label = "Fair"
    else:
        label = "Poor"

    return round(score, 1), label


# =============================================================================
# EASYOCR FALLBACK
# =============================================================================
# run_easyocr() is the second-pass OCR engine. It uses a neural-network model
# (CRAFT text detector + CRNN recogniser) which handles skewed, low-contrast,
# and unusual-font scans better than Tesseract's pattern matching.
#
# EasyOCR is only invoked when:
#   a) Tesseract quality was flagged as Fair or Poor, AND
#   b) The faculty member explicitly chooses to run it.
#
# On first call EasyOCR downloads ~200 MB of model weights. Subsequent calls
# reuse the cached models. gpu=False ensures it runs on any machine.
# =============================================================================

def run_easyocr(pil_images):
    """
    Extract text from pre-processed page images using EasyOCR.

    Accepts PIL Images that have already been rendered and corrected by
    preprocess_pages() — orientation and skew fixes are applied before
    this function is called, so it only handles the OCR step itself.

    Uses a neural-network model (CRAFT text detector + CRNN recogniser)
    which handles skewed, low-contrast, and unusual-font scans better than
    Tesseract's pattern matching. Results are sorted top-to-bottom by
    bounding-box Y coordinate to approximate visual reading order.

    EasyOCR and numpy are imported inside the function because they are
    optional dependencies — importing them lazily avoids import errors on
    machines where EasyOCR is not installed.

    On first call EasyOCR downloads ~200 MB of model weights. Subsequent
    calls reuse the cached models. gpu=False ensures it runs on any machine.

    Parameters:
        pil_images (list[PIL.Image]): Pre-processed page images from
                                     preprocess_pages(), in document order.

    Returns:
        list[str]: Extracted text for each page, in document order.

    Raises:
        Exception: Re-raised so the caller can show a user-facing error.
    """
    import easyocr
    import numpy as np

    reader = easyocr.Reader(["en"], gpu=False, verbose=False)
    pages_text = []

    for img in pil_images:
        # Convert PIL Image to a numpy array for EasyOCR.
        # Strip the alpha channel if present — EasyOCR expects RGB only.
        img_array = np.array(img)
        if img_array.ndim == 3 and img_array.shape[2] == 4:
            img_array = img_array[:, :, :3]

        results = reader.readtext(img_array)

        # Sort detections top-to-bottom by the top-left Y coordinate of the
        # bounding box so text flows in visual reading order.
        results.sort(key=lambda r: r[0][0][1])
        pages_text.append(" ".join(r[1] for r in results))

    return pages_text


# =============================================================================
# PAGE PREPROCESSING — ORIENTATION & SKEW CORRECTION
# =============================================================================
# preprocess_pages() runs before any OCR engine. It renders each PDF page as
# a PIL Image at 300 DPI, then applies two corrections automatically:
#
#   1. Orientation correction — Tesseract OSD detects if the page was scanned
#      sideways (90°), upside-down (180°), or at 270°, and rotates it upright.
#
#   2. Skew correction — OpenCV's minAreaRect finds the dominant angle of text
#      lines and applies a small rotation to straighten them. Only applied when
#      the detected angle is more than 0.5° (below that it's measurement noise).
#
# Returns the corrected PIL Images and a human-readable processing log that is
# shown in the "What we did" summary after OCR completes.
# =============================================================================

def _deskew_image(pil_image):
    """
    Detect and correct slight skew in a PIL Image using OpenCV.

    Converts to grayscale, thresholds to find text pixels, fits a minimum
    bounding rectangle to those pixels, and rotates by the detected angle.

    Parameters:
        pil_image (PIL.Image): The page image to deskew.

    Returns:
        (PIL.Image, float): Corrected image and the angle applied (degrees).
                            Returns (original_image, 0.0) if skew is negligible
                            or detection fails.
    """
    import cv2
    import numpy as np

    try:
        gray = np.array(pil_image.convert("L"))
        _, thresh = cv2.threshold(
            gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU
        )
        coords = np.column_stack(np.where(thresh > 0))
        if len(coords) < 100:   # too few text pixels to measure reliably
            return pil_image, 0.0

        angle = cv2.minAreaRect(coords)[-1]
        # minAreaRect returns angles in [-90, 0). Normalise to a signed value.
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle

        if abs(angle) < 0.5:    # below threshold — not worth correcting
            return pil_image, 0.0

        (h, w) = gray.shape[:2]
        M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
        img_array = np.array(pil_image)
        rotated = cv2.warpAffine(
            img_array, M, (w, h),
            flags=cv2.INTER_CUBIC,
            borderMode=cv2.BORDER_REPLICATE,
        )
        return Image.fromarray(rotated), angle
    except Exception:
        return pil_image, 0.0


def preprocess_pages(file_bytes, tesseract_path):
    """
    Render every PDF page at 300 DPI and apply orientation and skew correction.

    Called once before running Tesseract. The corrected PIL Images are passed
    directly to the OCR engine so the same preprocessing applies whether we
    run Tesseract or EasyOCR.

    Parameters:
        file_bytes (bytes): Raw bytes of the uploaded PDF.
        tesseract_path (str): Path to the Tesseract binary (needed for OSD).

    Returns:
        (list[PIL.Image], list[str]):
            - Corrected page images in document order.
            - Human-readable log entries describing what was done to each page.
    """
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page_count = len(doc)
    pil_images = []
    log_entries = [f"✔ {page_count} page{'s' if page_count != 1 else ''} detected"]

    pages_rotated = 0
    pages_deskewed = 0

    for page_num in range(page_count):
        page = doc[page_num]
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        # Step 1 — Orientation correction via Tesseract OSD.
        # OSD returns the clockwise rotation needed to make text upright.
        # We skip silently if OSD fails (low-text pages confuse it).
        try:
            osd = pytesseract.image_to_osd(
                img, output_type=pytesseract.Output.DICT
            )
            rotate = osd.get("rotate", 0)
            if rotate and rotate != 0:
                # PIL.rotate is counter-clockwise; OSD gives clockwise degrees.
                img = img.rotate(-rotate, expand=True)
                log_entries.append(
                    f"✔ Page {page_num + 1}: orientation corrected (rotated {rotate}°)"
                )
                pages_rotated += 1
        except Exception:
            pass

        # Step 2 — Skew correction via OpenCV.
        img, skew_angle = _deskew_image(img)
        if abs(skew_angle) > 0.5:
            log_entries.append(
                f"✔ Page {page_num + 1}: skew corrected ({skew_angle:+.1f}°)"
            )
            pages_deskewed += 1

        pil_images.append(img)

    doc.close()

    if pages_rotated == 0 and pages_deskewed == 0:
        log_entries.append("✔ All pages: orientation and alignment look correct")

    return pil_images, log_entries


# =============================================================================
# IMAGE EXTRACTION HELPER
# =============================================================================
# Used by the per-image alt text workflow to render each image from the PDF
# so the faculty member can see what they're describing.
# =============================================================================

def auto_classify_image(img_bytes, bbox, page_rect):
    """
    Automatically classify a PDF image as background, edge_artifact, artifact,
    or content using pixel statistics and geometry.

    This runs when the faculty member enters the alt text workflow — not during
    the initial analyze_pdf() call — so it never slows down the analysis step.

    Detection strategy:
      - Background: large coverage of the page + low color variance (few distinct
        colors). Common for scanned pages with a uniform paper color wash.
      - Edge artifact: thin strip near a page edge (dark border from book spine,
        page curl, or scanner shadow). Identified by extreme aspect ratio + proximity
        to page boundary, or by high darkness near an edge.
      - Artifact: small, very dark region anywhere on the page (smudge, speck).
      - Content: everything else — needs human review.

    Parameters:
        img_bytes (bytes): Raw extracted image bytes.
        bbox (fitz.Rect): Image bounding box in PDF point space.
        page_rect (fitz.Rect): Full page dimensions in PDF point space.

    Returns:
        dict with keys:
            classification — "background" | "edge_artifact" | "artifact" | "content"
            confidence     — float 0.0–1.0 (>= 0.75 = high-confidence, shown in bulk review)
            reason         — human-readable string shown in the UI
    """
    import numpy as np

    try:
        img  = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        arr  = np.array(img, dtype=float)

        # --- Pixel statistics ---
        # color_std: mean standard deviation across R, G, B channels.
        # Low value = image is mostly one colour (solid wash or near-solid background).
        color_std       = float(arr.std(axis=(0, 1)).mean())
        mean_brightness = float(arr.mean())

        # --- Geometric analysis ---
        page_w   = page_rect.width  or 1
        page_h   = page_rect.height or 1
        img_w    = bbox.width       or 0
        img_h    = bbox.height      or 0
        coverage = (img_w * img_h) / (page_w * page_h)
        aspect   = (img_w / img_h) if img_h > 0 else 0

        # near_edge: True if the bbox is within 5% of any page border.
        margin      = 0.05
        near_left   = bbox.x0 < page_w * margin
        near_right  = bbox.x1 > page_w * (1 - margin)
        near_top    = bbox.y0 < page_h * margin
        near_bottom = bbox.y1 > page_h * (1 - margin)
        near_edge   = near_left or near_right or near_top or near_bottom

        # --- Classification rules (first match wins) ---

        # Background: large coverage + low colour variance.
        # coverage bonus lifts confidence for images that dominate the page.
        if coverage > 0.40 and color_std < 30:
            confidence = min(0.95, 0.55 + (coverage - 0.40) + (30 - color_std) / 60)
            return {
                "classification": "background",
                "confidence": round(confidence, 2),
                "reason": (
                    f"Covers {coverage:.0%} of the page with very uniform colour "
                    f"(variation score {color_std:.1f}/255 — likely a paper background wash)."
                ),
            }

        # Edge artifact: thin strip right at a page border.
        thin = aspect > 8 or (aspect < 0.125 and aspect > 0)
        if near_edge and thin:
            return {
                "classification": "edge_artifact",
                "confidence": 0.85,
                "reason": (
                    f"Thin strip near the page edge "
                    f"(aspect ratio {max(aspect, 1/aspect if aspect else 1):.1f}:1) — "
                    "likely a scanner shadow or book-spine border."
                ),
            }

        # Dark edge artifact: near a page edge and very dark (spine shadow, etc.).
        if near_edge and mean_brightness < 60:
            return {
                "classification": "edge_artifact",
                "confidence": 0.80,
                "reason": (
                    f"Dark region (brightness {mean_brightness:.0f}/255) near the page edge — "
                    "likely a scan shadow or dark border."
                ),
            }

        # Small dark artifact anywhere on the page (smudge, speck).
        if mean_brightness < 40 and coverage < 0.10:
            return {
                "classification": "artifact",
                "confidence": 0.70,
                "reason": (
                    f"Small ({coverage:.1%} of page), very dark image "
                    f"(brightness {mean_brightness:.0f}/255) — likely a scan smudge or mark."
                ),
            }

        # Default: needs human review.
        return {
            "classification": "content",
            "confidence": 0.50,
            "reason": "No clear artifact pattern detected — please review manually.",
        }

    except Exception as exc:
        return {
            "classification": "content",
            "confidence": 0.50,
            "reason": f"Classification could not run ({exc}).",
        }


def run_auto_classification(images, pdf_bytes):
    """
    Run auto_classify_image() on every image in the list and return a dict
    mapping img_key → classification result.

    Handles both image sources:
      "pymupdf" — extracts the XObject directly and gets bbox via get_image_rects()
      "docling"  — crops the rendered page at the stored bbox for pixel stats;
                   uses the stored bbox for geometry checks

    Parameters:
        images    (list[dict]): Image metadata dicts from analyze_pdf().
                                Each dict must have "img_key", "page", and "source".
        pdf_bytes (bytes): Current working copy of the PDF.

    Returns:
        dict: {img_key: classification_result (dict)}
              img_key is an int (xref) for pymupdf images and a string for docling images.
    """
    results = {}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for img_meta in images:
            img_key  = img_meta.get("img_key", img_meta.get("xref"))
            source   = img_meta.get("source", "pymupdf")
            page_idx = img_meta["page"] - 1   # convert to 0-based
            try:
                page      = doc[page_idx]
                page_rect = page.rect

                if source == "docling":
                    # Docling image: bbox is already stored in TOPLEFT PDF point space.
                    # Wrap in a fitz.Rect so the geometry calculations in
                    # auto_classify_image() work the same way as for pymupdf images.
                    raw_bbox      = img_meta.get("bbox")
                    fitz_bbox     = fitz.Rect(*raw_bbox) if raw_bbox else fitz.Rect(0, 0, 0, 0)
                    img_bytes_raw, _ = extract_image_by_source(pdf_bytes, img_meta)
                else:
                    # PyMuPDF image: look up bbox via get_image_rects and extract XObject.
                    xref          = img_meta["xref"]
                    rects         = page.get_image_rects(xref)
                    fitz_bbox     = rects[0] if rects else fitz.Rect(0, 0, 0, 0)
                    img_info      = doc.extract_image(xref)
                    img_bytes_raw = img_info["image"] if img_info else None

                if img_bytes_raw:
                    results[img_key] = auto_classify_image(img_bytes_raw, fitz_bbox, page_rect)
                else:
                    results[img_key] = {
                        "classification": "content",
                        "confidence": 0.50,
                        "reason": "Image data could not be extracted for classification.",
                    }
            except Exception:
                results[img_key] = {
                    "classification": "content",
                    "confidence": 0.50,
                    "reason": "Classification failed for this image.",
                }
        doc.close()
    except Exception:
        pass
    return results


def render_page_with_image_highlight(pdf_bytes, page_num, xref=None, bbox=None):
    """
    Render a full PDF page and draw a thick red border around a specific image.

    This gives faculty the page-context view: they can see exactly which image
    on the page they are being asked to describe, even when a page contains
    multiple images or the extracted image alone is ambiguous.

    Approach:
      1. Render the page at 150 DPI (good enough for context, fast enough to
         not slow down the per-image workflow noticeably).
      2. Locate the image bounding box — either from get_image_rects(xref)
         for PyMuPDF XObject images, or directly from the bbox parameter for
         Docling PICTURE regions (already in PDF point space, TOPLEFT origin).
      3. Scale the bounding box from PDF points to pixel coordinates.
      4. Draw a thick red rectangle on the rendered pixmap using Pillow's
         ImageDraw — this keeps the dependency on libraries we already have
         and avoids needing to modify the PDF document itself.

    Parameters:
        pdf_bytes (bytes): The current working copy of the PDF.
        page_num  (int):   0-based page index where the image appears.
        xref      (int):   XObject cross-reference number — used for PyMuPDF images.
                           Pass None when using bbox instead.
        bbox      (tuple): (x0, y0, x1, y1) in PDF points, TOPLEFT origin —
                           used for Docling PICTURE regions. Pass None when
                           using xref instead.

    Returns:
        bytes: PNG image bytes of the highlighted page, or None on failure.
               Callers should fall back to the plain extracted image if None.
    """
    try:
        from PIL import ImageDraw

        doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[page_num]

        # Determine the highlight rectangle(s).
        # For PyMuPDF images: get_image_rects() returns one fitz.Rect per
        # location where this xref appears on the page.
        # For Docling images: the bbox is already in PDF point space — wrap
        # it in a list to match the same iteration pattern below.
        if bbox is not None:
            x0, y0, x1, y1 = bbox
            rects = [fitz.Rect(x0, y0, x1, y1)]
        elif xref is not None:
            rects = page.get_image_rects(xref)
        else:
            rects = []

        # Render the page at 150 DPI.
        # 150/72 ≈ 2.08 — a good balance between clarity and speed.
        dpi   = 150
        scale = dpi / 72
        mat   = fitz.Matrix(scale, scale)
        pix   = page.get_pixmap(matrix=mat)
        doc.close()

        # Convert pixmap to a Pillow Image for drawing.
        img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")

        if rects:
            draw = ImageDraw.Draw(img)
            for rect in rects:
                # Scale bounding box from PDF point space to pixel space.
                x0 = int(rect.x0 * scale)
                y0 = int(rect.y0 * scale)
                x1 = int(rect.x1 * scale)
                y1 = int(rect.y1 * scale)

                # Draw a thick red border (4px outline).
                # We draw four nested rectangles to approximate stroke width,
                # since Pillow's rectangle outline doesn't support line width.
                for offset in range(4):
                    draw.rectangle(
                        [x0 - offset, y0 - offset, x1 + offset, y1 + offset],
                        outline=(220, 38, 38),  # Tailwind red-600 — clearly visible
                    )

        out = io.BytesIO()
        img.save(out, format="PNG")
        return out.getvalue()

    except Exception:
        return None


def extract_image_from_pdf(pdf_bytes, xref):
    """
    Extract a single image from a PDF by its cross-reference number (xref).

    PyMuPDF assigns each embedded image a unique xref that persists for the
    lifetime of the document. We use extract_image() rather than rendering
    the full page, so we get the original image data at full resolution.

    Parameters:
        pdf_bytes (bytes): The current working copy of the PDF.
        xref (int): The xref number of the target image (from page.get_images()).

    Returns:
        (bytes, str): Raw image bytes and the file extension (e.g. "png", "jpeg"),
                      or (None, None) if extraction fails.
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        img_info = doc.extract_image(xref)
        doc.close()
        if img_info:
            return img_info["image"], img_info["ext"]
    except Exception:
        pass
    return None, None


def extract_image_by_source(pdf_bytes, img_meta):
    """
    Extract image bytes for an image that may be either a PyMuPDF XObject
    or a Docling PICTURE region.

    Dispatches on img_meta["source"]:
      "pymupdf" — delegates to extract_image_from_pdf() using img_meta["xref"]
      "docling"  — renders the page at 300 DPI and crops to img_meta["bbox"]

    Parameters:
        pdf_bytes (bytes): Current working copy of the PDF.
        img_meta  (dict):  Image metadata dict from analyze_pdf().

    Returns:
        (bytes, str): Image bytes and extension ("png"), or (None, None) on failure.
    """
    source = img_meta.get("source", "pymupdf")

    if source == "pymupdf":
        # Standard path: extract the embedded XObject directly.
        return extract_image_from_pdf(pdf_bytes, img_meta["xref"])

    # Docling path: render the page and crop at the stored bounding box.
    # bbox is (x0, y0, x1, y1) in PDF points, TOPLEFT origin — ready to scale.
    try:
        bbox    = img_meta.get("bbox")
        page_no = img_meta.get("page")   # 1-based
        if bbox is None or page_no is None:
            return None, None

        x0, y0, x1, y1 = bbox

        doc   = fitz.open(stream=pdf_bytes, filetype="pdf")
        page  = doc[page_no - 1]
        dpi   = 300
        scale = dpi / 72
        mat   = fitz.Matrix(scale, scale)
        pix   = page.get_pixmap(matrix=mat)
        doc.close()

        img_full = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")

        # Scale bbox from PDF points to pixel space.
        left   = max(0,             int(x0 * scale))
        top    = max(0,             int(y0 * scale))
        right  = min(img_full.width,  int(x1 * scale))
        bottom = min(img_full.height, int(y1 * scale))

        if right <= left or bottom <= top:
            return None, None

        cropped = img_full.crop((left, top, right, bottom))
        out = io.BytesIO()
        cropped.save(out, format="PNG")
        return out.getvalue(), "png"

    except Exception:
        return None, None


# =============================================================================
# PDF ANALYSIS — Phase 2 (real checks)
# =============================================================================
# analyze_pdf() performs the following checks on every uploaded PDF:
#
#   GATE CHECKS (stop analysis immediately if true):
#     1. Password protection — file cannot be read at all
#     2. Scanned/image-based — no selectable text; OCR required first
#
#   CONTENT CHECKS (run on readable, non-scanned PDFs):
#     Red   — Untagged PDF (no accessibility structure tree)
#     Red   — Images without alt text (images detected on any page)
#     Yellow — Missing document language tag (/Lang in PDF root)
#     Green  — Missing document title (empty /Title in metadata)
#     Green  — Missing bookmarks (no TOC and document > 9 pages)
#
# Each detected issue is returned as a dict with:
#   id, severity, title, description, wcag, ada, fix_preview, fix_data
#
# fix_data carries any extra information the fix function will need later
# (e.g., the detected language code, a title candidate, a list of images).
# =============================================================================


def _detect_title_candidate(doc):
    """
    Heuristic: scan the first page for the largest-font non-empty text span.
    That text is almost always the document title.

    Parameters:
        doc: an open fitz.Document object

    Returns:
        str — the candidate title text, or "" if nothing useful was found
    """
    try:
        page = doc[0]
        blocks = page.get_text("dict")["blocks"]
        best_text = ""
        best_size = 0
        for block in blocks:
            if block.get("type") != 0:   # skip image blocks
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size", 0)
                    text = span.get("text", "").strip()
                    # Accept spans that are clearly heading-sized and
                    # short enough to plausibly be a title (< 120 chars).
                    if text and size > best_size and len(text) < 120:
                        best_size = size
                        best_text = text
        return best_text
    except Exception:
        return ""


def _lang_display_name(lang_code):
    """
    Map a BCP 47 language code to a human-readable name for display in the UI.

    Parameters:
        lang_code (str): e.g. "en-US", "fr-FR"

    Returns:
        str — readable name, or the code itself if not in the table
    """
    names = {
        "en": "English", "en-US": "English", "en-GB": "English",
        "fr": "French",  "fr-FR": "French",
        "es": "Spanish", "es-ES": "Spanish",
        "de": "German",  "de-DE": "German",
        "zh": "Chinese", "ja": "Japanese", "ar": "Arabic",
        "pt": "Portuguese", "it": "Italian",
    }
    return names.get(lang_code, lang_code)


def _generate_toc_from_font_sizes(doc):
    """
    Build a table-of-contents list suitable for fitz.Document.set_toc() by
    treating text spans whose font size is significantly larger than the
    document's body text as headings.

    Approach:
      1. Collect every text span across all pages with its font size and page.
      2. Find the most common rounded font size — that is the body text size.
      3. Any span with size > 120% of body size is a heading candidate.
      4. Rank unique heading sizes descending to assign heading levels 1–3.
      5. Return a list of [level, title, page_number] entries.

    Parameters:
        doc: an open fitz.Document object

    Returns:
        list of [int level, str title, int page] — empty if nothing found
    """
    all_spans = []
    for page_num, page in enumerate(doc):
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    size = span.get("size", 0)
                    if text and size > 0 and len(text) < 150:
                        all_spans.append({
                            "text": text,
                            "size": size,
                            "page": page_num + 1,
                        })

    if not all_spans:
        return []

    # Most common rounded font size = body text
    sizes = [round(s["size"]) for s in all_spans]
    body_size = Counter(sizes).most_common(1)[0][0]

    # Heading candidates are spans noticeably larger than body text
    heading_spans = [s for s in all_spans if s["size"] > body_size * 1.2]
    if not heading_spans:
        return []

    # Map each unique heading size to a level (largest font = level 1)
    unique_sizes = sorted(set(round(s["size"]) for s in heading_spans), reverse=True)
    size_to_level = {sz: min(i + 1, 3) for i, sz in enumerate(unique_sizes)}

    # Build the TOC, de-duplicating identical consecutive headings
    toc = []
    seen = set()
    for span in heading_spans:
        text = span["text"]
        if text in seen:
            continue
        seen.add(text)
        level = size_to_level.get(round(span["size"]), 1)
        toc.append([level, text, span["page"]])

    return toc


# =============================================================================
# HEADING HIERARCHY HELPERS
# =============================================================================
# These two functions detect heading structure problems in tagged PDFs.
# They are called from analyze_pdf() only when is_tagged is True — there is
# no point checking headings in an untagged PDF since that issue is already
# flagged as Red and subsumes any heading problems.
# =============================================================================

def _extract_heading_sequence(file_bytes):
    """
    Walk the PDF StructTreeRoot recursively and return a list of heading
    level integers found in document order.

    PDF heading element types are /H (generic, treated as level 1), and
    /H1 through /H6. We extract only the numeric level — the text content
    of each heading requires parsing page content streams and is out of scope.

    Parameters:
        file_bytes (bytes): Raw PDF bytes to inspect.

    Returns:
        list[int]: Heading levels in order, e.g. [1, 2, 3, 2, 1, 2].
                   Empty list if the document has no StructTreeRoot or no
                   heading elements.
    """
    levels = []

    def _walk(node):
        """Recursively visit every node in the structure tree."""
        try:
            if isinstance(node, pikepdf.Dictionary):
                s_val = node.get("/S")
                if s_val is not None:
                    s = str(s_val)
                    if s == "/H":
                        levels.append(1)             # generic heading = level 1
                    elif len(s) == 3 and s[1] == "H" and s[2].isdigit():
                        levels.append(int(s[2]))     # /H1 … /H6
                # Descend into children (/K). Children can be:
                #   - a single Dictionary (another element)
                #   - an Array of elements and/or MCID integers
                #   - an integer MCID (skip — no children)
                kids = node.get("/K")
                if kids is not None:
                    _walk(kids)
            elif isinstance(node, pikepdf.Array):
                for item in node:
                    _walk(item)
            # Integers (MCID references) and other scalars are silently skipped.
        except Exception:
            pass

    try:
        with pikepdf.open(io.BytesIO(file_bytes)) as pdf:
            if "/StructTreeRoot" not in pdf.Root:
                return []
            _walk(pdf.Root["/StructTreeRoot"])
    except Exception:
        pass

    return levels


def _check_heading_hierarchy(levels):
    """
    Inspect a heading level sequence for structural problems.

    Rules checked:
      - No headings at all — returns "none"
      - A level jump skips one or more levels going deeper
        (e.g. H1 → H3 skips H2) — returns "skip"
      - Going back up (H3 → H1) is always valid — new top-level section.

    Parameters:
        levels (list[int]): Output of _extract_heading_sequence().

    Returns:
        str: "none" | "skip" | "ok"
    """
    if not levels:
        return "none"
    prev = levels[0]
    for lv in levels[1:]:
        if lv > prev + 1:    # jump skips at least one level going deeper
            return "skip"
        prev = lv
    return "ok"


def _describe_heading_sequence(levels):
    """
    Build a short human-readable summary of the heading sequence for display
    in the fix preview. Marks the first detected skip with an arrow so the
    faculty member can immediately see where the problem is.

    Parameters:
        levels (list[int]): Output of _extract_heading_sequence().

    Returns:
        str: Multi-line markdown string, or a fallback message if levels is empty.
    """
    if not levels:
        return "_No heading elements were found in this document's structure._"

    lines = []
    prev = levels[0]
    for lv in levels:
        indent = "  " * (lv - 1)
        if lv > prev + 1:
            lines.append(f"{indent}**H{lv} ← skips H{prev + 1}**")
        else:
            lines.append(f"{indent}H{lv}")
        prev = lv
    # Truncate to 20 lines to avoid overwhelming the screen
    if len(lines) > 20:
        lines = lines[:20] + [f"_… and {len(levels) - 20} more headings_"]
    return "\n".join(lines)


# -----------------------------------------------------------------------------
# FIX FUNCTIONS
# Each function takes (working_bytes, fix_data) and returns (new_bytes, message).
# working_bytes — the current state of the working copy of the PDF.
# fix_data      — the dict stored in the issue when analysis ran (may be empty).
# new_bytes     — updated PDF bytes after the fix (same as input if fix failed).
# message       — a brief confirmation string shown to the user after saving.
# -----------------------------------------------------------------------------

def apply_fix_missing_title(working_bytes, fix_data):
    """
    Write a document title to the PDF's Info dictionary and XMP metadata.

    Uses both pdf.docinfo (old-style Info dict) and pdf.open_metadata() (XMP)
    for maximum reader compatibility.

    Parameters:
        working_bytes (bytes): Current working copy of the PDF.
        fix_data (dict): Must contain "confirmed_title" (set by the UI after
                         faculty edits/approves) or "title_candidate" as fallback.

    Returns:
        (bytes, str): Updated PDF bytes and a short confirmation message.
    """
    title = (fix_data.get("confirmed_title") or fix_data.get("title_candidate", "")).strip()
    if not title:
        return working_bytes, "No title provided — metadata unchanged."
    try:
        with pikepdf.open(io.BytesIO(working_bytes)) as pdf:
            # XMP metadata (preferred by modern PDF readers and assistive tech)
            with pdf.open_metadata() as meta:
                meta["dc:title"] = title
            # Info dict (legacy, but still read by many screen readers)
            pdf.docinfo["/Title"] = title
            out = io.BytesIO()
            pdf.save(out)
            return out.getvalue(), f"Title set to \"{title}\"."
    except Exception as e:
        return working_bytes, f"Could not apply title fix: {e}"


def apply_fix_missing_lang(working_bytes, fix_data):
    """
    Write a /Lang entry to the PDF's document root.

    /Lang is the standard PDF mechanism for declaring the document language.
    Screen readers use it to select the correct speech engine and pronunciation
    rules. (WCAG 2.2 SC 3.1.1)

    Parameters:
        working_bytes (bytes): Current working copy of the PDF.
        fix_data (dict): Must contain "lang_code" (BCP 47, e.g. "en-US").

    Returns:
        (bytes, str): Updated PDF bytes and a short confirmation message.
    """
    lang_code = fix_data.get("lang_code", "en-US")
    try:
        with pikepdf.open(io.BytesIO(working_bytes)) as pdf:
            pdf.Root["/Lang"] = pikepdf.String(lang_code)
            out = io.BytesIO()
            pdf.save(out)
            return out.getvalue(), f"Language tag set to '{lang_code}'."
    except Exception as e:
        return working_bytes, f"Could not apply language fix: {e}"


def apply_fix_bookmarks(working_bytes, fix_data):
    """
    Generate and write a bookmark outline to the PDF using font-size heuristics
    to detect heading structure.

    Uses PyMuPDF's set_toc() which writes a proper PDF /Outline tree that PDF
    readers and assistive technology can navigate. (WCAG 2.2 SC 2.4.5)

    Parameters:
        working_bytes (bytes): Current working copy of the PDF.
        fix_data (dict): Unused; included for consistent function signature.

    Returns:
        (bytes, str): Updated PDF bytes and a short confirmation message.
    """
    try:
        doc = fitz.open(stream=working_bytes, filetype="pdf")
        toc = _generate_toc_from_font_sizes(doc)
        if toc:
            doc.set_toc(toc)
        out = io.BytesIO()
        doc.save(out)
        doc.close()
        count = len(toc)
        return out.getvalue(), f"Added {count} bookmark{'s' if count != 1 else ''}."
    except Exception as e:
        return working_bytes, f"Could not apply bookmarks fix: {e}"


def apply_fix_untagged(working_bytes, fix_data):
    """
    Handle the untagged PDF issue.

    True in-place PDF tagging (adding a full StructTreeRoot) requires
    reconstructing the document's entire logical structure — this is beyond
    what a runtime script can reliably do without a dedicated tagging engine
    (e.g., Adobe Acrobat Pro, axesPDF, CommonLook).

    Instead, we note the issue. When the faculty member downloads their file as
    DOCX (via render_choose_format), the build_docx_from_pdf() function applies
    font-size heuristics to produce a structured DOCX with proper heading styles.
    They can then re-export that DOCX as a properly tagged PDF from Word or
    Google Docs.

    Parameters:
        working_bytes (bytes): Current working copy of the PDF.
        fix_data (dict): Unused.

    Returns:
        (bytes, str): Unchanged PDF bytes and an explanatory message.
    """
    # No change to the PDF bytes — the fix manifests in the DOCX output path.
    return working_bytes, (
        "Noted. When you download your file as a Word document, it will be structured "
        "with proper headings and reading order — ready to re-save as a properly tagged "
        "PDF from Word or Google Docs."
    )


def apply_fix_alt_text(working_bytes, fix_data):
    """
    Handle the missing alt text issue.

    Writing alt text directly into a PDF's structure tree requires the document
    to already have a StructTreeRoot with Figure elements — which untagged PDFs
    lack entirely. For tagged PDFs, it requires locating each Figure element and
    writing the /Alt attribute.

    For now we record the acknowledgement. The alt text descriptions collected
    during the session (stored in st.session_state) are available for inclusion
    in the DOCX output.

    Parameters:
        working_bytes (bytes): Current working copy of the PDF.
        fix_data (dict): Contains "images" list from analysis.

    Returns:
        (bytes, str): Unchanged PDF bytes and a confirmation message.
    """
    return working_bytes, (
        "The text descriptions you've provided will be included in the Word document output. "
        "When you re-export from Word, you can add them as alt text to each image."
    )


def apply_fix_heading_hierarchy(working_bytes, fix_data):
    """
    Handle the heading hierarchy issue.

    Rewriting the StructTreeRoot to correct heading levels requires modifying
    both the structure tree and the page content streams — this cannot be done
    reliably in-place without a dedicated tagging engine.

    The real fix manifests in the DOCX output: build_docx_from_pdf() uses
    font-size heuristics to assign correct Heading 1 / Heading 2 / Heading 3
    styles, producing a document with a valid hierarchy that the faculty member
    can re-export as a properly tagged PDF from Word or Google Docs.

    Parameters:
        working_bytes (bytes): Current working copy of the PDF.
        fix_data (dict): Contains "issue_type" and "heading_levels".

    Returns:
        (bytes, str): Unchanged PDF bytes and an explanatory message.
    """
    return working_bytes, (
        "Noted. The Word document output will use the correct heading structure "
        "based on your document's visual layout. Re-exporting from Word will produce "
        "a properly tagged PDF with a valid heading hierarchy."
    )


def apply_fix_readability_barrier(working_bytes, fix_data):
    """
    Handle severe and moderate readability barrier issues (font size too small).

    Like the untagged PDF and heading hierarchy fixes, this cannot be corrected
    in the original PDF bytes at runtime — changing font sizes requires
    rewriting page content streams, which is beyond safe in-place editing.

    The fix manifests in the DOCX output path: build_docx_from_pdf() enforces
    a minimum font size so all body text is at least 9pt in the Word document.
    Faculty can then re-export from Word as a properly readable PDF.

    Parameters:
        working_bytes (bytes): Current working copy of the PDF.
        fix_data (dict): Contains pct_affected, small_chars, total_chars.

    Returns:
        (bytes, str): Unchanged PDF bytes and an explanatory message.
    """
    pct = fix_data.get("pct_affected", 0)
    return working_bytes, (
        f"Noted — {pct}% of your text is below 9pt. When you download your file "
        "as a Word document, all text will be brought up to a comfortable minimum "
        "size. Re-exporting that Word file as PDF will carry the correction through."
    )


# Dispatch table: maps issue ID → fix function.
# Used by render_resolving_issue() to apply the right fix after faculty confirms.
FIX_DISPATCH = {
    "missing_title":                apply_fix_missing_title,
    "missing_lang":                 apply_fix_missing_lang,
    "missing_bookmarks":            apply_fix_bookmarks,
    "untagged_pdf":                 apply_fix_untagged,
    "missing_alt_text":             apply_fix_alt_text,
    "heading_hierarchy":            apply_fix_heading_hierarchy,
    "readability_barrier_severe":   apply_fix_readability_barrier,
    "readability_barrier_moderate": apply_fix_readability_barrier,
}


def _run_docling(file_bytes):
    """
    Run Docling document conversion on raw PDF bytes.

    Docling uses ML models (downloaded once and cached in ~/.cache/huggingface)
    to understand document structure at a semantic level — it identifies headings,
    paragraphs, list items, tables, images, and reading order across complex page
    layouts. This gives us richer data than PyMuPDF's font-size heuristics alone.

    Why a temp file? Docling's DocumentConverter expects a file path, not bytes.
    We write to a named temp file, convert, then delete it in the finally block
    so the file is always cleaned up even if conversion raises an exception.

    Parameters:
        file_bytes (bytes): Raw PDF bytes.

    Returns:
        DoclingDocument object if conversion succeeds, or None on any failure.
        Callers should treat None as "Docling unavailable — fall back to PyMuPDF."
    """
    tmp_path = None
    try:
        # Write the PDF bytes to a temporary file so Docling can open it by path.
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        # DocumentConverter is Docling's main entry point.
        # By default it uses its own layout model for digital PDFs.
        # For scanned PDFs it would need an OCR backend — but we only call
        # this function on digital PDFs (after the scanned gate check passes).
        converter = DocumentConverter()
        result = converter.convert(tmp_path)
        return result.document

    except Exception:
        # Docling can fail on malformed PDFs, unsupported features, or model
        # errors. We return None silently and let the caller degrade gracefully.
        return None

    finally:
        # Always remove the temp file — even if an exception was raised above.
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


def analyze_pdf(file_bytes):
    """
    Analyze a PDF file and return a result dict describing what was found.

    Gate checks run first (password, scanned). If both pass, content checks
    run to detect specific accessibility issues. Each detected issue is returned
    as a dict that includes fix_data — the extra information the matching fix
    function will need when the user confirms the fix later.

    Parameters:
        file_bytes (bytes): The raw bytes of the uploaded PDF file.

    Returns:
        dict with keys:
            status  — "password_protected" | "scanned" | "no_issues" | "issues_found"
            issues  — list of issue dicts (empty if status is not "issues_found")
            cover_page_only — bool (only present when status == "scanned")
    """
    issues = []

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        # --- GATE CHECK 1: Password protection ---
        # doc.needs_pass is True when the file is encrypted and unreadable.
        if doc.needs_pass:
            doc.close()
            return {"status": "password_protected", "issues": []}

        # --- GATE CHECK 2: Scanned / image-based document ---
        # Scanned PDFs have no selectable text — get_text() returns "".
        # We apply the 2-page sampling rule defined in CLAUDE.md:
        #   page 1 empty → check page 2 → stop analysis regardless, require OCR.
        if len(doc) > 0:
            first_page_text = doc[0].get_text().strip()
            if not first_page_text:
                cover_page_only = len(doc) > 1 and bool(doc[1].get_text().strip())
                doc.close()
                return {
                    "status": "scanned",
                    "issues": [],
                    "cover_page_only": cover_page_only,
                }

        page_count = len(doc)
        metadata = doc.metadata or {}

        # --- DOCLING ANALYSIS ---
        # Run Docling on the digital PDF to extract semantic structure — element
        # types (heading, list, table, image), reading order, and layout data.
        # This supplements PyMuPDF's lower-level block/span extraction.
        #
        # We run it here (after gate checks) so:
        #   1. We know the PDF is digital and readable (not scanned, not encrypted).
        #   2. The result is available for all subsequent content checks below.
        #   3. build_docx_from_pdf() can use it via session_state.docling_doc
        #      without re-running the (slow) conversion a second time.
        #
        # If Docling fails, docling_doc is None and all downstream code falls
        # back to PyMuPDF — no disruption to the user.
        docling_doc = _run_docling(file_bytes)

        # --- CONTENT CHECK: Missing document title (Green) ---
        # PyMuPDF exposes the Info dict as doc.metadata with a "title" key.
        title = (metadata.get("title") or "").strip()
        if not title:
            title_candidate = _detect_title_candidate(doc)
            if title_candidate:
                fix_preview = (
                    f"I'll add the title \"{title_candidate}\" to your document's metadata. "
                    "Screen readers announce this when the file opens, and it helps with "
                    "search indexing. You can edit it below before I save it."
                )
            else:
                fix_preview = (
                    "I'll add a descriptive title to the document metadata. "
                    "Enter the title you'd like to use for this document."
                )
            issues.append({
                "id": "missing_title",
                "severity": "green",
                "title": "Missing document title in metadata",
                "description": (
                    "The document's title field is empty — it currently defaults to the filename. "
                    "Screen readers announce the document title when a file opens, and search "
                    "tools use it for indexing. A descriptive title helps all users understand "
                    "what they're looking at before they start reading."
                ),
                "wcag": "WCAG 2.2 SC 2.4.2 — Page Titled",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": fix_preview,
                "fix_data": {"title_candidate": title_candidate},
            })

        # --- CONTENT CHECK: Missing language tag (Yellow) ---
        # PyMuPDF surfaces /Lang (if present) as doc.metadata["language"].
        lang = (metadata.get("language") or "").strip()
        if not lang:
            # Auto-detect the language from the first ~1000 characters of text.
            sample_text = ""
            for page in doc:
                sample_text += page.get_text()
                if len(sample_text) >= 1000:
                    break
            try:
                detected = detect_language(sample_text[:1000]) if sample_text.strip() else "en"
            except Exception:
                detected = "en"
            # Convert langdetect 2-letter codes to BCP 47 codes for PDF /Lang.
            lang_map = {
                "en": "en-US", "fr": "fr-FR", "es": "es-ES",
                "de": "de-DE", "pt": "pt-PT", "it": "it-IT",
            }
            lang_code = lang_map.get(detected, detected)
            issues.append({
                "id": "missing_lang",
                "severity": "yellow",
                "title": "Missing document language tag",
                "description": (
                    "This document doesn't declare what language it's written in. Screen readers "
                    "use the language tag to select the correct pronunciation rules and speech "
                    "engine. Without it, text may be read aloud in the wrong language or with "
                    "incorrect pronunciation — especially for technical terms and proper names."
                ),
                "wcag": "WCAG 2.2 SC 3.1.1 — Language of Page",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    f"I've detected that this document is written in "
                    f"{_lang_display_name(lang_code)} and will add the language tag "
                    f"'{lang_code}' to the document metadata."
                ),
                "fix_data": {"lang_code": lang_code},
            })

        # --- CONTENT CHECK: Missing bookmarks (Green) ---
        # doc.get_toc() returns a list of [level, title, page] entries.
        # An empty list means no bookmark outline exists.
        toc = doc.get_toc()
        if page_count > 9 and not toc:
            issues.append({
                "id": "missing_bookmarks",
                "severity": "green",
                "title": "Missing bookmarks / navigation",
                "description": (
                    f"This document is {page_count} pages long but has no bookmarks. "
                    "Without bookmarks, everyone — not just screen reader users — has to "
                    "scroll through the entire document to find what they need."
                ),
                "wcag": "WCAG 2.2 SC 2.4.5 — Multiple Ways",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "I'll generate bookmarks based on the heading structure I detect in your "
                    "document. You can review the bookmark list before I save it."
                ),
                "fix_data": {},
            })

        # --- CONTENT CHECK: Untagged PDF (Red) ---
        # A tagged PDF has a /StructTreeRoot in its root dictionary.
        # We use pikepdf for this check because PyMuPDF does not expose the
        # root dict directly in a reliable way across versions.
        is_tagged = False
        try:
            with pikepdf.open(io.BytesIO(file_bytes)) as pdf:
                is_tagged = "/StructTreeRoot" in pdf.Root
        except Exception:
            pass  # If pikepdf fails, assume untagged (safer for accessibility)

        if not is_tagged:
            issues.append({
                "id": "untagged_pdf",
                "severity": "red",
                "title": "Untagged PDF — no document structure",
                "description": (
                    "This document has no accessibility tags. Tags are the behind-the-scenes "
                    "structure that tells screen readers what is a heading, what is a paragraph, "
                    "and what order to read things in. Without them, a screen reader user hears "
                    "the content as an undifferentiated stream of text — or nothing at all."
                ),
                "wcag": "WCAG 2.2 SC 1.3.1 — Info and Relationships",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "I'll convert this document to a structured Word document (DOCX) that "
                    "preserves your headings, paragraphs, and reading order. You can then "
                    "re-save it as a properly tagged PDF from Word or Google Docs. This is "
                    "the most reliable path to a fully accessible PDF."
                ),
                "fix_data": {},
            })

        # --- CONTENT CHECK: Heading hierarchy (Yellow) ---
        # Only meaningful for tagged PDFs — untagged documents are already
        # flagged Red above, and that issue subsumes any heading problems.
        # We check two failure modes:
        #   "none"  — the document is tagged but has no heading elements at all
        #   "skip"  — heading levels jump (e.g. H1 → H3), breaking navigation
        if is_tagged:
            heading_levels = _extract_heading_sequence(file_bytes)
            hierarchy_status = _check_heading_hierarchy(heading_levels)

            if hierarchy_status == "none":
                issues.append({
                    "id": "heading_hierarchy",
                    "severity": "yellow",
                    "title": "No headings found in document structure",
                    "description": (
                        "This document is tagged for accessibility but contains no heading "
                        "elements. Screen reader users navigate long documents by jumping "
                        "between headings — without them, they must listen to the entire "
                        "document from start to finish to find what they need."
                    ),
                    "wcag": "WCAG 2.2 SC 1.3.1 — Info and Relationships",
                    "ada": "Title II ADA, 28 CFR Part 35",
                    "fix_preview": (
                        "The Word document output will use your document's visual layout "
                        "to assign proper Heading 1, Heading 2, and Heading 3 styles. "
                        "Re-exporting from Word will produce a fully navigable PDF."
                    ),
                    "fix_data": {"issue_type": "none", "heading_levels": []},
                })

            elif hierarchy_status == "skip":
                issues.append({
                    "id": "heading_hierarchy",
                    "severity": "yellow",
                    "title": "Incorrect or missing heading hierarchy",
                    "description": (
                        "The heading levels in this document skip one or more levels — "
                        "for example, jumping from Heading 1 directly to Heading 3 with "
                        "no Heading 2 in between. Screen reader users rely on heading "
                        "structure to understand how a document is organized and to jump "
                        "between sections. A broken hierarchy makes that navigation "
                        "confusing or misleading."
                    ),
                    "wcag": "WCAG 2.2 SC 1.3.1 — Info and Relationships",
                    "ada": "Title II ADA, 28 CFR Part 35",
                    "fix_preview": (
                        "Here is the heading structure detected in your document. "
                        "The Word document output will correct the hierarchy based on "
                        "the visual layout of your document."
                    ),
                    "fix_data": {
                        "issue_type": "skip",
                        "heading_levels": heading_levels,
                    },
                })

        # --- CONTENT CHECK: Images without alt text (Red) ---
        # We collect images from two sources and merge them into one list:
        #
        #   PyMuPDF XObjects — page.get_images() returns every embedded image
        #   that lives as its own PDF object. Works well for digital PDFs.
        #
        #   Docling PICTURE regions — Docling's ML vision model analyses rendered
        #   pixels and detects image regions even when they are not separate XObjects.
        #   This covers scanned PDFs where a content image (figure, photo, diagram)
        #   is a pixel region inside a full-page background scan rather than a
        #   standalone XObject. PyMuPDF cannot see these; Docling can.
        #
        # Each image dict carries:
        #   "page"    — 1-based page number
        #   "xref"    — int xref for pymupdf images; None for docling images
        #   "img_key" — hashable dict key used throughout the alt text workflow
        #               (xref int for pymupdf; "docling_{page}_{x0}_{y0}" string for docling)
        #   "source"  — "pymupdf" or "docling"
        #   "bbox"    — (x0, y0, x1, y1) in PDF points, TOPLEFT origin
        #               (only present for docling images; pymupdf uses xref for bbox lookup)
        images = []
        for page_num, page in enumerate(doc):
            for img in page.get_images(full=True):
                xref = img[0]
                images.append({
                    "page": page_num + 1,
                    "xref": xref,
                    "img_key": xref,        # int — used as dict key in alt text workflow
                    "source": "pymupdf",
                })

        # Merge Docling PICTURE detections.
        # Only runs when Docling succeeded (docling_doc is not None).
        # The bbox conversion from BOTTOMLEFT (Docling default) to TOPLEFT
        # (PyMuPDF rendering convention) uses to_top_left_origin(page_height).
        if docling_doc is not None:
            try:
                for item, _ in docling_doc.iterate_items():
                    if not hasattr(item, "label") or item.label != DocItemLabel.PICTURE:
                        continue
                    if not item.prov:
                        continue
                    prov    = item.prov[0]
                    page_no = prov.page_no   # 1-based
                    try:
                        page_height = docling_doc.pages[page_no].size.height
                        tl          = prov.bbox.to_top_left_origin(page_height)
                        x0 = round(tl.l, 2)
                        y0 = round(tl.t, 2)
                        x1 = round(tl.r, 2)
                        y1 = round(tl.b, 2)
                    except Exception:
                        continue   # Can't compute bbox — skip this item
                    img_key = f"docling_{page_no}_{int(x0)}_{int(y0)}"
                    images.append({
                        "page":    page_no,
                        "xref":    None,     # No PDF xref — Docling region, not XObject
                        "img_key": img_key,  # string key — avoids int collision with xrefs
                        "source":  "docling",
                        "bbox":    (x0, y0, x1, y1),
                    })
            except Exception:
                pass  # Docling iteration failure must never break image detection

        if images:
            count = len(images)
            issues.append({
                "id": "missing_alt_text",
                "severity": "red",
                "title": "Images with missing text descriptions",
                "description": (
                    f"This document contains {count} image{'s' if count != 1 else ''} "
                    "that don't have text descriptions (alt text). Students who use screen "
                    "readers — software that reads content aloud — won't be able to access "
                    "the information those images convey."
                ),
                "wcag": "WCAG 2.2 SC 1.1.1 — Non-text Content",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    f"I found {count} image{'s' if count != 1 else ''} in your document. "
                    "I'll walk you through each one so you can tell me whether it's decorative "
                    "or informational, and we'll add the right text description for each."
                ),
                "fix_data": {"images": images},
            })

        # --- CONTENT CHECK: Readability barrier — font size too small ---
        # Iterate all text spans and count characters whose font size falls
        # below 9pt, which is the practical minimum for sustained body reading.
        # Spans shorter than 3 characters are excluded to avoid flagging page
        # numbers, footnote markers, and other incidental small text.
        # Percentage of affected content drives severity:
        #   > 25% of total text → Red (severe barrier)
        #   1–25%              → Yellow (moderate barrier)
        SMALL_FONT_PT = 9.0
        total_chars   = 0
        small_chars   = 0

        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        size = span.get("size", 0)
                        if len(text) < 3:
                            continue
                        total_chars += len(text)
                        if size < SMALL_FONT_PT:
                            small_chars += len(text)

        if total_chars > 0 and small_chars > 0:
            pct = (small_chars / total_chars) * 100
            if pct > 25:
                rb_id       = "readability_barrier_severe"
                rb_severity = "red"
                rb_title    = "Severe readability barrier — text too small"
                rb_desc     = (
                    f"More than a quarter of your document's text ({pct:.0f}%) uses a "
                    "font size below 9pt, which is too small for many readers. Students "
                    "with low vision, reading difficulties, or aging eyes may struggle "
                    "to read this content without significant zooming."
                )
            else:
                rb_id       = "readability_barrier_moderate"
                rb_severity = "yellow"
                rb_title    = "Moderate readability barrier — some text too small"
                rb_desc     = (
                    f"About {pct:.0f}% of your document's text uses a font size below "
                    "9pt. While this affects a smaller portion of the content, it still "
                    "creates barriers for students with low vision or reading difficulties."
                )

            issues.append({
                "id":          rb_id,
                "severity":    rb_severity,
                "title":       rb_title,
                "description": rb_desc,
                "wcag":        "WCAG 2.2 SC 1.4.3, 1.4.6 — Contrast (Minimum/Enhanced)",
                "ada":         "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "When you download your file as a Word document, I'll apply a "
                    "minimum font size so all body text is comfortably readable."
                ),
                "fix_data": {
                    "pct_affected": round(pct, 1),
                    "small_chars":  small_chars,
                    "total_chars":  total_chars,
                },
            })

        # --- COMBINED SCAN: reading order, color, list tagging, line spacing, letter spacing ---
        # page.get_text("dict") re-parses the page binary on every call — no caching.
        # One pass per page eliminates four redundant parse calls on every document page.
        reading_order_problem_pages = 0
        reading_order_pages_checked = 0
        color_usage     = Counter()
        visual_list_lines   = 0
        tight_line_pairs    = 0
        total_line_pairs    = 0
        tight_tracking_spans  = 0
        total_tracking_spans  = 0

        for page in doc:
            page_h = page.rect.height
            page_w = page.rect.width
            blocks = page.get_text("dict")["blocks"]
            text_blocks = [b for b in blocks if b.get("type") == 0 and b.get("lines")]

            # Reading order: flag pages where blocks jump significantly upward within
            # the same column zone (> 25% of page height, < 40% of width apart).
            if len(text_blocks) >= 4:
                reading_order_pages_checked += 1
                violations = 0
                for i in range(1, len(text_blocks)):
                    prev_y = (text_blocks[i-1]["bbox"][1] + text_blocks[i-1]["bbox"][3]) / 2
                    curr_y = (text_blocks[i]["bbox"][1]   + text_blocks[i]["bbox"][3])   / 2
                    prev_x = (text_blocks[i-1]["bbox"][0] + text_blocks[i-1]["bbox"][2]) / 2
                    curr_x = (text_blocks[i]["bbox"][0]   + text_blocks[i]["bbox"][2])   / 2
                    if (curr_y < prev_y - page_h * 0.25) and (abs(curr_x - prev_x) < page_w * 0.40):
                        violations += 1
                if violations >= 2:
                    reading_order_problem_pages += 1

            for block in text_blocks:
                lines = block.get("lines", [])

                # Line spacing: measure gap between adjacent line tops relative to font size.
                # Skips single-line blocks (range is empty) and footnote-scale text (< 7pt).
                for i in range(len(lines) - 1):
                    first_spans = lines[i].get("spans", [])
                    if not first_spans:
                        continue
                    line_fs = first_spans[0].get("size", 0)
                    if line_fs < 7:
                        continue
                    line_step = lines[i + 1]["bbox"][1] - lines[i]["bbox"][1]
                    if line_step <= 0:
                        continue
                    total_line_pairs += 1
                    if (line_step / line_fs) < 1.10:  # below standard single spacing
                        tight_line_pairs += 1

                for line in lines:
                    # List pattern: reconstruct the full line string from its spans so
                    # LIST_PATTERN can check the leading characters without a second
                    # page.get_text() call.
                    line_text = "".join(s.get("text", "") for s in line.get("spans", []))
                    if LIST_PATTERN.match(line_text):
                        visual_list_lines += 1

                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue

                        # Color: accumulate character count per non-black/non-white color.
                        # Excludes near-black (< 40) and near-white (> 215) brightness.
                        if len(text) >= 3:
                            color = span.get("color", 0)
                            r, g, b = (color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF
                            if 40 <= (r + g + b) / 3 <= 215:
                                color_usage[color] += len(text)

                        # Letter spacing: geometric char-width estimate.
                        # (span_width / char_count) / font_size < 0.28 signals extreme compression.
                        # Condensed fonts (narrow by design) are excluded via CONDENSED_FONT_HINTS.
                        if len(text) >= 5:
                            font_name = span.get("font", "").lower()
                            if not any(h in font_name for h in CONDENSED_FONT_HINTS):
                                fs = span.get("size", 0)
                                if fs >= 8:
                                    sw = span["bbox"][2] - span["bbox"][0]
                                    if sw > 0:
                                        total_tracking_spans += 1
                                        if sw / (len(text) * fs) < 0.28:
                                            tight_tracking_spans += 1

        # --- Reading order flag ---
        if (
            reading_order_pages_checked >= 2
            and reading_order_problem_pages >= max(1, reading_order_pages_checked * 0.40)
        ):
            issues.append({
                "id": "reading_order",
                "severity": "yellow",
                "title": "Reading order may not match visual order",
                "description": (
                    "The order in which content was embedded in this PDF doesn't match "
                    "the top-to-bottom order a sighted reader would follow. Screen readers "
                    "and assistive technology encounter content in the order it is stored "
                    "in the file — if that order is scrambled, users hear content out of "
                    "sequence, making the document confusing or unintelligible."
                ),
                "wcag": "WCAG 2.2 SC 1.3.2 — Meaningful Sequence",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "When you download your file as a Word document, the content will be "
                    "re-ordered into a logical top-to-bottom reading sequence. Re-exporting "
                    "from Word will produce a PDF with a correct, consistent reading order."
                ),
                "fix_data": {
                    "problem_pages": reading_order_problem_pages,
                    "pages_checked": reading_order_pages_checked,
                },
            })

        # --- Color only cue / Excessive text colors flags ---
        # 2–5 distinct meaningful colors → intentional color coding, flag for manual review.
        # 6+ → excessive color variety. The two conditions are mutually exclusive.
        meaningful_colors = [c for c, cnt in color_usage.items() if cnt >= 40]

        if 2 <= len(meaningful_colors) <= 5:
            issues.append({
                "id": "color_only_cue",
                "severity": "yellow",
                "title": "Colored text detected — manual accessibility check needed",
                "description": (
                    f"This document uses {len(meaningful_colors)} distinct text colors. "
                    "Color can create accessibility barriers in several ways — most of "
                    "which can't be verified or fixed automatically. A quick manual "
                    "review is needed to confirm your students won't miss anything."
                ),
                "wcag": "WCAG 2.2 SC 1.4.1 — Use of Color",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": "",
                "fix_data": {"color_count": len(meaningful_colors)},
            })

        # --- Inconsistent list tagging flag ---
        docling_list_items = 0
        if docling_doc is not None:
            try:
                for item, _ in docling_doc.iterate_items():
                    if hasattr(item, "label") and item.label == DocItemLabel.LIST_ITEM:
                        docling_list_items += 1
            except Exception:
                pass  # Docling failure must not interrupt the broader analysis

        if visual_list_lines >= 8 and docling_list_items < visual_list_lines * 0.50:
            issues.append({
                "id": "inconsistent_list_tagging",
                "severity": "green",
                "title": "Lists may not be properly tagged in the document structure",
                "description": (
                    f"This document appears to contain {visual_list_lines} list items "
                    "formatted with bullets, dashes, or numbers — but the underlying "
                    "file structure may not label them as actual lists. Screen readers "
                    "announce list context ('list of 5 items') and let users jump "
                    "between items. Without that tagging, list content reads as a "
                    "stream of plain text with no navigation or context cues."
                ),
                "wcag": "WCAG 2.2 SC 1.3.1 — Info and Relationships",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "When you download your file as a Word document, visually formatted "
                    "lists will be converted to properly structured list elements with "
                    "correct indentation and item tags. Re-exporting from Word will "
                    "carry the list structure through into the PDF."
                ),
                "fix_data": {
                    "visual_list_lines":  visual_list_lines,
                    "docling_list_items": docling_list_items,
                },
            })

        # --- Line spacing flag ---
        if total_line_pairs >= 20 and (tight_line_pairs / total_line_pairs) > 0.40:
            pct_tight = round(tight_line_pairs / total_line_pairs * 100)
            issues.append({
                "id": "line_spacing_tight",
                "severity": "green",
                "title": "Line spacing may be too tight",
                "description": (
                    f"About {pct_tight}% of this document's text lines are spaced "
                    "tighter than the recommended minimum. Tight line spacing makes "
                    "sustained reading harder — especially for students with dyslexia, "
                    "low vision, or attention difficulties. WCAG recommends body text "
                    "line spacing of at least 1.5× the font size."
                ),
                "wcag": "WCAG 2.2 SC 1.4.12 — Text Spacing",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "When you download your file as a Word document, I'll apply 1.5× "
                    "line spacing to all body text. This improves readability without "
                    "changing the document's structure or meaning."
                ),
                "fix_data": {"tight_pct": pct_tight},
            })

        # --- Letter spacing flag ---
        if total_tracking_spans >= 40 and (tight_tracking_spans / total_tracking_spans) > 0.45:
            pct_tracking = round(tight_tracking_spans / total_tracking_spans * 100)
            issues.append({
                "id": "letter_spacing_tight",
                "severity": "green",
                "title": "Letter spacing may be too tight",
                "description": (
                    "A significant portion of this document's text appears to use "
                    "compressed letter spacing. Tight character tracking reduces "
                    "legibility for all readers and creates particular barriers for "
                    "students with dyslexia, who rely on distinct spacing between "
                    "letters to distinguish similar characters (b/d, p/q, n/m)."
                ),
                "wcag": "WCAG 2.2 SC 1.4.12 — Text Spacing",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "When you download your file as a Word document, letter spacing "
                    "will be normalized to the font's standard tracking. This removes "
                    "any artificial compression present in the original file."
                ),
                "fix_data": {"tight_pct": pct_tracking},
            })

        # --- Excessive text colors flag ---
        if len(meaningful_colors) > 5:
            issues.append({
                "id": "excessive_text_colors",
                "severity": "green",
                "title": "Excessive use of different text colors",
                "description": (
                    f"This document uses {len(meaningful_colors)} distinct text colors "
                    "across substantial portions of the content. Using many different "
                    "text colors creates visual noise and cognitive distraction, making "
                    "it harder to focus on the material. It can also confuse readers "
                    "about which color differences carry meaning and which are decorative."
                ),
                "wcag": "WCAG 2.2 SC 1.4.1 — Use of Color",
                "ada": "Title II ADA, 28 CFR Part 35",
                "fix_preview": (
                    "When you download your file as a Word document, text colors will "
                    "be reduced to a small, intentional set — body text and headings. "
                    "This keeps the visual hierarchy clear without the distraction of "
                    "many competing colors."
                ),
                "fix_data": {"color_count": len(meaningful_colors)},
            })

        doc.close()

    except Exception:
        # If analysis fails entirely, fall back to demo issues rather than
        # crashing — this keeps the app usable even if a PDF is malformed.
        return {"status": "issues_found", "issues": DEMO_ISSUES}

    if not issues:
        return {"status": "no_issues", "issues": [], "docling_doc": docling_doc}

    return {"status": "issues_found", "issues": issues, "docling_doc": docling_doc}


# =============================================================================
# DOCX OUTPUT BUILDER — Phase 2
# =============================================================================
# build_docx_from_pdf() extracts text and structure from the working PDF and
# produces a properly styled DOCX file using python-docx.
#
# It uses the same font-size heuristic as _generate_toc_from_font_sizes():
#   - Most common rounded font size  = body text → Normal paragraph style
#   - Size > 140% of body            = Heading 1
#   - Size > 120% of body            = Heading 2
#
# This function is used by render_choose_format() when the faculty member
# chooses to download their remediated file as a Word document.
# =============================================================================

def build_docx_from_pdf(pdf_bytes):
    """
    Extract text and structure from a PDF and produce a structured DOCX.

    Uses Docling as the primary path when a cached DoclingDocument is available
    in session state. Docling's ML-based element labels (TITLE, SECTION_HEADER,
    LIST_ITEM, TABLE, PICTURE) map cleanly to DOCX styles and give much better
    heading and list detection than the font-size heuristic alone.

    Falls back to the PyMuPDF font-size heuristic if Docling is unavailable
    (e.g., the PDF was an OCR output generated mid-session, or Docling failed).

    Parameters:
        pdf_bytes (bytes): The working copy of the PDF (with metadata fixes applied).

    Returns:
        bytes: A DOCX file as raw bytes, ready for st.download_button.
    """
    doc_out = DocxDocument()

    # ------------------------------------------------------------------
    # PRIMARY PATH: Use the cached Docling document from session state.
    # The DoclingDocument was produced by _run_docling() during analyze_pdf()
    # and stored in st.session_state.docling_doc by render_analyzing().
    # ------------------------------------------------------------------
    docling_doc = st.session_state.get("docling_doc")

    # ------------------------------------------------------------------
    # Shared setup for both paths: read alt text results once so both
    # the Docling path and the PyMuPDF fallback can use the same data.
    # ------------------------------------------------------------------
    alt_results   = st.session_state.get("alt_text_results", {})
    alt_images    = st.session_state.get("alt_text_images", [])

    # Types whose images should be excluded from the output entirely.
    REMOVED_TYPES     = {"artifact", "background", "edge_artifact"}
    # Types whose images carry an approved description worth writing out.
    INFORMATIONAL_TYPES = {"critical", "supplementary"}

    if docling_doc is not None:
        try:
            # Build exclusion data from alt_text_results so we can skip artifact/
            # background images in the output.
            #
            # Docling PICTURE items don't carry xrefs, so we can't match 1-to-1.
            # Instead we exclude pages where EVERY image has been marked for removal.
            # Pages with at least one non-removed image still get placeholders for
            # all their pictures (we can't tell which Docling item maps to which xref).

            # Group img_keys by page (1-based page numbers from fix_data).
            # Uses img_key (not xref) as the identifier so Docling images — which
            # have xref=None — are correctly tracked alongside PyMuPDF images.
            keys_by_page: dict[int, set] = {}
            for img_meta in alt_images:
                k = img_meta.get("img_key", img_meta.get("xref"))
                keys_by_page.setdefault(img_meta["page"], set()).add(k)

            # A page is fully excluded only if every image on it is marked for removal.
            fully_excluded_pages = {
                pg for pg, keys in keys_by_page.items()
                if all(alt_results.get(k, {}).get("type") in REMOVED_TYPES for k in keys)
            }

            # Build a page → [approved descriptions] mapping for informational images.
            # Docling PICTURE items have a page number but no xref, so we group
            # approved descriptions by page and consume them in order as PICTURE
            # items are encountered during iterate_items().
            #
            # Why a list per page? A page can have multiple informational images.
            # Popping from the front (index 0) preserves the xref order from the
            # alt text workflow, which follows reading order left-to-right, top-to-bottom.
            descriptions_by_page: dict[int, list[str]] = {}
            for xref, entry in alt_results.items():
                if (
                    entry.get("type") in INFORMATIONAL_TYPES
                    and entry.get("alt_text", "").strip()
                ):
                    pg = entry.get("page")
                    if pg is not None:
                        descriptions_by_page.setdefault(pg, []).append(
                            entry["alt_text"].strip()
                        )

            # iterate_items() walks the document in reading order and yields
            # (item, level) tuples. The level is nesting depth, not heading level.
            for item, _ in docling_doc.iterate_items():
                if not hasattr(item, "label"):
                    continue

                label = item.label
                text = getattr(item, "text", "").strip()

                if label == DocItemLabel.TITLE:
                    # Document title — Heading 1 in Word (most prominent)
                    doc_out.add_heading(text, level=1)

                elif label == DocItemLabel.SECTION_HEADER:
                    # Section heading — use Heading 2 as a safe default.
                    # Docling doesn't always expose sub-levels, so we use 2
                    # consistently rather than guessing a deeper hierarchy.
                    doc_out.add_heading(text, level=2)

                elif label == DocItemLabel.LIST_ITEM:
                    # List items — map to Word's built-in List Bullet style
                    # so they render as a proper bulleted list, not plain text.
                    # WCAG 2.2 SC 1.3.1: structure must be programmatically determinable.
                    if text:
                        doc_out.add_paragraph(text, style="List Bullet")

                elif label == DocItemLabel.TABLE:
                    # Tables: export to markdown to extract cell text, then
                    # add a plain-text note pointing faculty to the original.
                    # Full table reconstruction in python-docx is complex and
                    # will be addressed in a future iteration.
                    doc_out.add_paragraph("[Table — see original PDF for layout]")

                elif label == DocItemLabel.PICTURE:
                    # Skip picture placeholders for pages where every image was
                    # marked as an artifact or background during the alt text
                    # workflow. Pages with mixed content (some real, some removed)
                    # still get a placeholder — we can't tell which Docling picture
                    # item corresponds to which xref.
                    item_page = item.prov[0].page_no if item.prov else None
                    if item_page in fully_excluded_pages:
                        continue

                    page_num = item_page or "?"

                    # Look up an approved description for this page.
                    # Pop from the front so each PICTURE item on the same page
                    # gets its own unique description rather than the same one repeated.
                    page_descs = descriptions_by_page.get(item_page, [])
                    if page_descs:
                        approved_text = page_descs.pop(0)
                        # Write the description as the alt text for this image.
                        # Format: "[Image, page N: <approved description>]"
                        # This makes the description visible and searchable in Word,
                        # satisfying WCAG 2.2 SC 1.1.1 for documents that will be
                        # re-exported to PDF.
                        doc_out.add_paragraph(
                            f"[Image, page {page_num}: {approved_text}]"
                        )
                    else:
                        doc_out.add_paragraph(f"[Image — page {page_num} of original PDF]")

                elif label in (DocItemLabel.TEXT, DocItemLabel.PARAGRAPH):
                    # Standard body text paragraph.
                    if text:
                        doc_out.add_paragraph(text)

                else:
                    # Catch-all: any other Docling element type we haven't explicitly
                    # handled (e.g., FORMULA, CODE, CAPTION) gets added as plain text
                    # so content is never silently dropped.
                    if text:
                        doc_out.add_paragraph(text)

            out = io.BytesIO()
            doc_out.save(out)
            return out.getvalue()

        except Exception:
            # If the Docling path fails mid-way, fall through to the PyMuPDF
            # fallback below rather than returning an incomplete document.
            doc_out = DocxDocument()  # Reset to a fresh document

    # ------------------------------------------------------------------
    # FALLBACK PATH: Font-size heuristic using PyMuPDF.
    # Used when no Docling document is cached (e.g., for OCR output PDFs
    # generated mid-session) or when the Docling path raises an exception.
    #
    # Heuristic:
    #   Most common rounded font size  = body text → Normal paragraph style
    #   Size > 140% of body            = Heading 1
    #   Size > 120% of body            = Heading 2
    # ------------------------------------------------------------------
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        # Build a flat img_key → alt_text_results entry lookup.
        # alt_results may contain both int keys (pymupdf xrefs) and string keys
        # (docling img_keys). Only int keys correspond to XObjects accessible via
        # page.get_images(); string keys are filtered out in the per-image loop.
        xref_to_entry = {k: entry for k, entry in alt_results.items()}

        # Determine body text size by finding the most frequently used font size.
        all_sizes = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span.get("text", "").strip():
                            all_sizes.append(span.get("size", 12))
        body_size = Counter(round(s) for s in all_sizes).most_common(1)[0][0] if all_sizes else 12

        for page_num, page in enumerate(doc):
            if page_num > 0:
                doc_out.add_page_break()

            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue

                for line in block.get("lines", []):
                    # Join spans within a line into a single string.
                    line_text = " ".join(
                        span.get("text", "") for span in line.get("spans", [])
                    ).strip()
                    if not line_text:
                        continue

                    # Determine average font size for this line.
                    line_sizes = [span.get("size", body_size) for span in line.get("spans", [])]
                    avg_size = sum(line_sizes) / len(line_sizes) if line_sizes else body_size

                    # Map font size ratio to a DOCX heading level.
                    if avg_size > body_size * 1.4:
                        doc_out.add_heading(line_text, level=1)
                    elif avg_size > body_size * 1.2:
                        doc_out.add_heading(line_text, level=2)
                    else:
                        doc_out.add_paragraph(line_text)

            # ----------------------------------------------------------
            # After text blocks: emit alt text for images on this page.
            # page.get_images(full=True) returns one tuple per image XObject,
            # with the xref at index 0 — the same key used in alt_text_results.
            #
            # This adds image descriptions at the end of each page's content.
            # Inline positioning is not possible with the font-size heuristic
            # path, but all approved descriptions are preserved in the output.
            # ----------------------------------------------------------
            for img in page.get_images(full=True):
                xref  = img[0]
                entry = xref_to_entry.get(xref, {})
                img_type = entry.get("type", "")

                if img_type in REMOVED_TYPES:
                    # Artifact, background, or edge artifact — skip entirely.
                    continue

                if img_type == "decorative":
                    # Decorative images get empty alt text in the final PDF.
                    # No visible placeholder is needed in the Word document.
                    continue

                alt_text = entry.get("alt_text", "").strip()

                if alt_text:
                    # Approved description from the alt text workflow.
                    # Format matches the Docling path for consistency.
                    doc_out.add_paragraph(
                        f"[Image, page {page_num + 1}: {alt_text}]"
                    )
                elif img_type in INFORMATIONAL_TYPES:
                    # Image was classified as informational but description
                    # was not saved (e.g., the workflow was abandoned mid-way).
                    doc_out.add_paragraph(
                        f"[Image, page {page_num + 1}: description not provided]"
                    )
                else:
                    # Image was never reviewed — generic placeholder.
                    doc_out.add_paragraph(
                        f"[Image — page {page_num + 1} of original PDF]"
                    )

        doc.close()

    except Exception:
        doc_out.add_paragraph(
            "The content of this document could not be extracted automatically. "
            "Please open the original PDF and copy the content manually."
        )

    out = io.BytesIO()
    doc_out.save(out)
    return out.getvalue()


# =============================================================================
# SESSION STATE INITIALIZATION
# =============================================================================
# st.session_state persists between Streamlit re-runs for the same browser
# session. We initialize any keys that don't exist yet on the very first run.
# =============================================================================

def init_state():
    """Set default values for all session_state keys on first load."""
    defaults = {
        "step": "upload",           # Current workflow step
        "file_name": None,          # Original uploaded file name
        "file_bytes": None,         # Raw bytes of the uploaded file
        "issues": [],               # Full list of issue dicts from analysis
        "resolved_ids": [],         # List of issue IDs the user has resolved
        "skipped_ids": [],          # Issue IDs the user explicitly skipped via the alternative path
        "showing_alternative": False,  # True when the post-rejection alternative UI is active
        "ocr_download_mode": False,    # True when user chose "download and review first" on OCR screen
        "manual_review_results": {},   # {item_id: "ok" | "attention"} — responses on manual review screen
        "resolved_count": 0,        # Running count of resolved issues (drives re-analysis)
        "current_issue_id": None,   # ID of the issue currently being worked on
        "proposed_fix": None,       # Text of the fix shown during QA confirmation
        "reanalysis_done": False,   # Whether re-analysis was offered at the current checkpoint
        "cover_page_only": False,   # True if only page 1 is unreadable (page 2 has text)
        "ocr_pages_text": [],       # List of strings — one per page — from run_ocr()
        "ocr_docx_bytes": None,     # Prebuilt DOCX bytes for the _readable download
        "ocr_pdf_bytes": None,      # Prebuilt PDF bytes for the _readable download
        "ocr_issues": [],           # Accessibility issues found in the OCR-converted PDF
        "ocr_quality_score": None,  # Float 0–100 from score_ocr_quality()
        "ocr_quality_label": None,  # "Good", "Fair", or "Poor"
        "ocr_engine_used": None,    # "tesseract" or "easyocr" — shown in preview
        "ocr_processing_log": [],   # Step-by-step log shown in "What we did"
        "working_pdf_bytes": None,  # Running copy of the PDF — updated as fixes are applied
        "alt_text_images": [],      # Image list from the missing_alt_text issue's fix_data
        "alt_text_index": 0,        # Index of the image currently being worked on
        "alt_text_phase": "question_1",  # Sub-step within the per-image workflow
        "alt_text_results": {},     # {xref: {type, severity, alt_text, page}} — one per image
        "alt_text_auto_classifications": {},  # {xref: {classification, confidence, reason}}
        "docling_doc": None,        # DoclingDocument from _run_docling() — cached here so
                                    # build_docx_from_pdf() can use it without re-converting
        "edited_docx_bytes": None,  # Cached DOCX output from build_docx_from_pdf() — built
                                    # once on first render of choose_format, reused on reruns
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_state()


# =============================================================================
# SIDEBAR
# =============================================================================
# The sidebar is always visible. It holds settings and a Start Over button.
# We use st.sidebar as a context manager (the 'with' block) so everything
# inside appears in the sidebar panel on the left.
# =============================================================================

with st.sidebar:
    st.header("⚙️ Settings")

    # Tesseract OCR path — pytesseract needs to know where the binary lives.
    # This path will be read by the OCR step in Phase 2.
    st.text_input(
        label="Tesseract OCR Path",
        value=r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        help="Path to tesseract.exe. Download from: https://github.com/UB-Mannheim/tesseract/wiki",
        key="tesseract_path",
    )

    st.divider()

    # "Start Over" button — only show it when the user is past the upload step.
    # Clicking it resets all session state and returns to the upload screen.
    if st.session_state.step != "upload":
        if st.button("↩ Start Over", use_container_width=True):
            # Reset every key back to its default value by clearing state
            # and re-running init_state (which sets the defaults again).
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            init_state()
            st.rerun()

    st.divider()

    # Severity system explainer — shown once in the sidebar as reference.
    # Per WCAG 2.2 SC 1.4.1, we always pair the color indicator with text.
    st.markdown(
        """
        **About issue severity**

        🔴 **Red Light** — Serious issues that create significant obstacles
        for some or all of your students.

        🟡 **Yellow Light** — Important issues required for accessibility
        compliance under Title II of the ADA.

        🟢 **Green Light** — Finishing touches that bring your document
        to full compliance.
        """
    )

    st.divider()

    st.markdown(
        """
        **About this tool**

        Built to support faculty in making digital course content accessible
        under **Title II of the ADA** (28 CFR Part 35), effective
        April 24, 2026 for large public institutions.

        Version 0.3 — Phase 2
        """
    )


# =============================================================================
# RENDER FUNCTIONS — one per workflow step
# =============================================================================
# Each function below handles one step in the workflow. It renders the UI
# for that step and sets up the buttons that transition to the next step.
#
# Pattern for transitioning:
#   st.session_state.step = "next_step_name"
#   st.rerun()   ← tells Streamlit to re-run the script from the top
# =============================================================================


# -----------------------------------------------------------------------------
# STEP: upload
# The starting screen. Shows a file uploader. When a PDF is uploaded,
# store the file data in session state and move to "analyzing".
# -----------------------------------------------------------------------------
def render_upload():
    """
    Purpose: Let the user upload a PDF.
    WCAG reference: File upload control has an accessible label (SC 1.3.1).
    """
    st.title("♿ PDF Assistant")
    st.markdown(
        "Upload a course PDF and I'll check it for accessibility issues that may "
        "create barriers for your students. I'll walk you through any fixes step by step."
    )
    st.divider()

    # Copyright acknowledgment — must be checked before a file can be uploaded.
    copyright_acknowledged = st.checkbox(
        "I acknowledge that the file I am uploading is my own work or that I have "
        "the right to use it, and that I take full responsibility for ensuring it "
        "is free from copyright restrictions. This tool assumes all uploaded materials "
        "are used legitimately and bears no responsibility for copyright infringement.",
        key="copyright_ack",
    )

    uploaded_file = st.file_uploader(
        label="Choose a PDF file to check",
        type=["pdf"],
        help="Upload the course PDF you want to review for accessibility issues.",
        key="pdf_uploader",
        disabled=not copyright_acknowledged,
    )

    if not copyright_acknowledged:
        st.caption("Please check the box above to enable file upload.")
    else:
        # Processing time notice — shown once the uploader is enabled so faculty
        # know upfront that analysis can take a few minutes for larger files.
        # st.caption() renders in smaller, muted text — appropriate for a
        # supplementary note that shouldn't compete with the uploader itself.
        st.caption(
            "Once you upload your file, it may take a few minutes to analyze "
            "depending on its size. Large files can take 5 minutes or more. "
            "Please be patient — the tool is working."
        )

    if uploaded_file is not None and copyright_acknowledged:
        # Store the file name and raw bytes in session state so we can
        # access them in later steps even after Streamlit re-renders.
        # working_pdf_bytes starts as an exact copy of the original and is
        # updated in-place as each fix is confirmed — the original is never touched.
        st.session_state.file_name = uploaded_file.name.strip()
        st.session_state.file_bytes = uploaded_file.read()
        st.session_state.working_pdf_bytes = st.session_state.file_bytes
        st.session_state.step = "analyzing"
        st.rerun()


# -----------------------------------------------------------------------------
# STEP: analyzing
# Shows a spinner while we call analyze_pdf(), then routes to the appropriate
# next step based on what the analysis found.
# -----------------------------------------------------------------------------
def render_analyzing():
    """
    Purpose: Analyze the uploaded PDF and route to the correct next step.
    This step is transient — the user never stays here long.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # ------------------------------------------------------------------
    # Preflight check — verify the temp directory is writable before
    # Docling tries to create its temporary PDF there.
    # Tesseract is not needed for analysis, so tesseract_path is omitted.
    # ------------------------------------------------------------------
    preflight = preflight_check(
        output_paths=[tempfile.gettempdir()],
    )
    if not preflight["ok"]:
        st.error(
            "🔴 **Setup issue — cannot start analysis.**\n\n"
            + "\n\n".join(preflight["errors"])
        )
        if st.button("← Go back", type="primary"):
            st.session_state.step = "upload"
            st.rerun()
        return

    with st.spinner(f"Analyzing **{st.session_state.file_name}** — just a moment..."):
        result = analyze_pdf(st.session_state.file_bytes)

    # Cache the Docling document in session state so build_docx_from_pdf()
    # can use it later without re-running the (slow) ML conversion.
    # If Docling failed, this will be None — all downstream code handles that.
    st.session_state.docling_doc = result.get("docling_doc")

    # Route to the correct next step based on what was found.
    if result["status"] == "password_protected":
        st.session_state.step = "password_protected"
    elif result["status"] == "scanned":
        # Store whether only the cover page was unreadable vs. the whole doc.
        # Used by render_scanned_doc() for tailored messaging.
        st.session_state.cover_page_only = result.get("cover_page_only", False)
        st.session_state.step = "scanned_doc"
    elif result["status"] == "no_issues":
        st.session_state.step = "no_issues"
    else:
        # Store the issues list and go to the issue list screen.
        st.session_state.issues = result["issues"]
        st.session_state.step = "issue_list"

    st.rerun()


# -----------------------------------------------------------------------------
# STEP: password_protected
# Verbatim message from CLAUDE.md. Offers two choices.
# -----------------------------------------------------------------------------
def render_password_protected():
    """
    Purpose: Inform the user the file is password protected and offer options.
    Per CLAUDE.md: never ask for the password. Offer guidance instead.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Security warning — verbatim from CLAUDE.md
    st.warning(
        "⚠️ **DO NOT share your password or any other credentials in this tool.**"
    )

    st.info(
        "It looks like this file is password protected, which means I'm not able to "
        "read or analyze its contents. If you are the original author of this document, "
        "I can walk you through saving a new version of it without the password protection "
        "— just let me know. If someone else created this file, you may need to reach out "
        "to them or your IT support team to get an unlocked version."
    )

    st.markdown("**How would you like to proceed?**")

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Yes, walk me through it.",
            type="primary",
            use_container_width=True,
        ):
            st.session_state.step = "password_walkthrough"
            st.rerun()
    with col2:
        if st.button(
            "I'll take care of it another way.",
            use_container_width=True,
        ):
            st.session_state.step = "done"
            st.session_state.done_reason = "password_exit"
            st.rerun()

    _render_go_back("upload")


# -----------------------------------------------------------------------------
# STEP: password_walkthrough
# Step-by-step instructions for removing password protection. Then prompt
# the user to re-upload the unlocked file.
# -----------------------------------------------------------------------------
def render_password_walkthrough():
    """
    Purpose: Guide the user through removing password protection.
    After following the steps, the user returns to upload a new file.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    st.markdown(
        "Here's how to remove password protection in the most common applications:"
    )

    with st.expander("Adobe Acrobat", expanded=True):
        st.markdown(
            """
            1. Open your PDF in Adobe Acrobat.
            2. Go to **File → Properties → Security**.
            3. Change the Security Method to **No Security**.
            4. Click **OK**, then save the file (**File → Save As**) with a new name.
            5. Upload the new file here.
            """
        )

    with st.expander("Microsoft Word"):
        st.markdown(
            """
            1. Open the file in Word (it may prompt for the password to open it).
            2. Go to **File → Info → Protect Document → Encrypt with Password**.
            3. Delete the password field and click **OK**.
            4. Save the file as PDF: **File → Save As → PDF**.
            5. Upload the new file here.
            """
        )

    with st.expander("macOS Preview"):
        st.markdown(
            """
            1. Open the PDF in Preview. Enter the password when prompted.
            2. Go to **File → Export as PDF**.
            3. In the save dialog, click **Security Options** and remove the password.
            4. Save and upload the new file here.
            """
        )

    with st.expander("Google Drive"):
        st.markdown(
            """
            1. Upload the password-protected PDF to Google Drive.
            2. Right-click it and choose **Open with → Google Docs**.
            3. Google Docs will convert it. Go to **File → Download → PDF Document**.
            4. Upload the downloaded file here.
            """
        )

    st.divider()
    st.markdown("Once you have an unlocked version, click below to start over and upload it.")

    if st.button("↩ Upload a new file", type="primary"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        init_state()
        st.rerun()

    _render_go_back("password_protected")


# -----------------------------------------------------------------------------
# STEP: scanned_doc
# Verbatim Red Light message from CLAUDE.md. Offers two choices.
# -----------------------------------------------------------------------------
def render_scanned_doc():
    """
    Purpose: Inform the user the document contains scanned (image-based) text
    and stop analysis until OCR is complete.

    This is a Red Light issue — the most serious category.

    The 2-page detection logic in analyze_pdf() may have determined that only
    the cover page is unreadable (cover_page_only=True) vs. the full document.
    Either way, analysis is stopped and OCR is required before proceeding.

    Accessibility note: the "ANALYSIS STOPPED" label uses text, not color alone,
    to communicate the severity — per WCAG 2.2 SC 1.4.1.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Tailor the opening line slightly based on whether only the cover is unreadable.
    # Both cases stop analysis — but the description is more accurate this way.
    if st.session_state.cover_page_only:
        scope_note = (
            "At least one page of your PDF contains text that was scanned as an image."
        )
    else:
        scope_note = (
            "Your PDF contains text that was scanned as an image."
        )

    # "ANALYSIS STOPPED" callout — pairs text label with color (WCAG SC 1.4.1)
    st.error(
        f"🔴 **ANALYSIS STOPPED — {SEVERITY_LABELS['red']} Issue**\n\n"
        f"{scope_note} This creates serious accessibility barriers for most users: "
        "screen readers cannot read image-based text, students cannot search or copy "
        "it, and it cannot be resized or translated. "
        "This document cannot be analyzed or improved until it is converted to a "
        "readable PDF.\n\n"
        "**Would you like me to convert it now?**"
    )

    # Standards reference — supporting context, not the lead reason to act
    st.caption(
        "WCAG 2.2 SC 1.4.5 — Images of Text | Title II ADA, 28 CFR Part 35"
    )

    # Scan quality disclaimer — sets expectations without blocking the workflow.
    # Orientation and skew are handled automatically; contrast is the one factor
    # the software cannot correct for.
    st.caption(
        "💡 **Tip:** The conversion tool will automatically detect and correct "
        "page orientation and slight skew. For best results, make sure your scan "
        "has clear contrast between the text and background — very faint or "
        "washed-out scans may still produce inaccurate text regardless of the "
        "scanner used."
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Yes, make this document readable.",
            type="primary",
            use_container_width=True,
        ):
            st.session_state.step = "running_ocr"
            st.rerun()
    with col2:
        if st.button(
            "No, leave it as it is.",
            use_container_width=True,
        ):
            st.session_state.step = "scanned_goodbye"
            st.rerun()

    _render_go_back("upload")


# -----------------------------------------------------------------------------
# STEP: scanned_goodbye
# Shown when the user declines OCR conversion on a scanned document.
# Offers two exit paths: upload a different file or end the session entirely.
# -----------------------------------------------------------------------------
def render_scanned_goodbye():
    """
    Purpose: Close the session gracefully when the user chooses not to convert
    their scanned document.

    Two choices:
      - Upload another file — resets state and returns to the upload screen
      - No, I'm done for now — ends the session
    """
    st.title("♿ PDF Assistant")
    st.divider()

    st.info(
        "No problem — you can return any time to convert this document and work "
        "through its accessibility issues. If you have another PDF you'd like to "
        "check, you can upload it now."
    )

    st.markdown("**What would you like to do?**")

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Upload another file",
            type="primary",
            use_container_width=True,
        ):
            # Reset all session state and return to the upload screen.
            # This is the same reset pattern used by "Start Over" in the sidebar.
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            init_state()
            st.rerun()
    with col2:
        if st.button(
            "No, I'm done for now.",
            use_container_width=True,
        ):
            st.session_state.step = "done"
            st.session_state.done_reason = "scanned_exit"
            st.rerun()

    _render_go_back("scanned_doc")


# -----------------------------------------------------------------------------
# STEP: running_ocr
# Runs Tesseract OCR on every page of the scanned document, then builds both
# DOCX and PDF versions of the readable output and stores them in session state.
# Transitions to ocr_format_select on success, or shows an error on failure.
# -----------------------------------------------------------------------------
def render_running_ocr():
    """
    Purpose: Run real Tesseract OCR on the uploaded PDF.

    Steps performed here (per CLAUDE.md OCR Process):
      1. Configure pytesseract with the Tesseract path from sidebar settings.
      2. Render each page as a 300 DPI image using PyMuPDF.
      3. Extract text from each image using Tesseract.
      4. Build DOCX and PDF output files (with disclaimer) and store as bytes
         in session state so the format selection screen can serve them directly.
      5. Transition to ocr_format_select, or show an error if Tesseract fails.

    Error handling:
      If Tesseract raises any exception (binary not found, wrong path, etc.),
      a clear error message is shown with a "Go back" button. We never crash
      silently — the user always sees actionable guidance.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Read the Tesseract path the user entered in the sidebar.
    # The sidebar text_input stores this value in session_state["tesseract_path"].
    tesseract_path = st.session_state.get(
        "tesseract_path",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    )

    # ------------------------------------------------------------------
    # Preflight check — confirm Tesseract is callable and the temp
    # directory is writable before starting the OCR pipeline.
    # Catching these problems here gives the user a clear, actionable
    # error rather than a cryptic traceback mid-conversion.
    # ------------------------------------------------------------------
    preflight = preflight_check(
        tesseract_path=tesseract_path,
        output_paths=[tempfile.gettempdir()],
    )
    if not preflight["ok"]:
        st.error(
            "🔴 **Setup issue — cannot start OCR.**\n\n"
            + "\n\n".join(preflight["errors"])
        )
        if st.button("← Go back", type="primary"):
            st.session_state.step = "scanned_doc"
            st.rerun()
        return

    preprocess_error = None
    ocr_error        = None

    with st.spinner(
        "Checking page orientation, correcting any skew, then converting to "
        "readable text — this may take a moment for longer documents..."
    ):
        # ---- Step 1: Preprocessing (orientation + skew) ----
        # Separated from OCR so failures give accurate error messages.
        try:
            pil_images, preprocessing_log = preprocess_pages(
                st.session_state.file_bytes, tesseract_path
            )
        except Exception as e:
            preprocess_error = str(e)
            # Fall back to plain page renders with no preprocessing applied.
            try:
                doc = fitz.open(stream=st.session_state.file_bytes, filetype="pdf")
                pil_images = []
                for page_num in range(len(doc)):
                    pix = doc[page_num].get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                    pil_images.append(Image.open(io.BytesIO(pix.tobytes("png"))))
                doc.close()
                preprocessing_log = [
                    f"✔ {len(pil_images)} page{'s' if len(pil_images) != 1 else ''} detected",
                    f"⚠️ Orientation/skew correction skipped — {preprocess_error}",
                ]
            except Exception:
                pil_images = []
                preprocessing_log = []

        # ---- Step 2: OCR ----
        if pil_images:
            try:
                pages_text = run_ocr(pil_images, tesseract_path)

                quality_score, quality_label = score_ocr_quality(pages_text)
                st.session_state.ocr_quality_score  = quality_score
                st.session_state.ocr_quality_label  = quality_label
                st.session_state.ocr_engine_used    = "tesseract"
                st.session_state.ocr_processing_log = preprocessing_log + [
                    "✔ OCR performed using Tesseract",
                    f"✔ Quality estimate: "
                    f"{'✅' if quality_label == 'Good' else '⚠️' if quality_label == 'Fair' else '❌'} "
                    f"{quality_label} ({quality_score}%)",
                ]
                st.session_state.ocr_pages_text = pages_text

                # Don't build DOCX/PDF here — render_ocr_format_select() builds
                # them lazily on first render and caches the result. Building here
                # would be wasted work if the user only downloads one format.
                # Clear any stale cached output from a previous OCR run.
                st.session_state.ocr_docx_bytes = None
                st.session_state.ocr_pdf_bytes  = None

                # analyze_pdf() needs a PDF to inspect, so we still build the PDF
                # here — but only to drive the accessibility analysis, not for download.
                ocr_pdf_bytes = build_pdf(pages_text)
                analysis_result = analyze_pdf(ocr_pdf_bytes)
                st.session_state.ocr_issues = (
                    analysis_result["issues"]
                    if analysis_result["status"] == "issues_found"
                    else []
                )
                st.session_state.working_pdf_bytes = ocr_pdf_bytes

            except Exception as e:
                ocr_error = str(e)
        else:
            ocr_error = "Could not render pages from this PDF."

    if ocr_error:
        st.error(
            "🔴 **Conversion failed — Tesseract OCR could not run.**\n\n"
            "This usually means Tesseract is not installed, or the path in the "
            "sidebar settings is incorrect.\n\n"
            f"**Error details:** `{ocr_error}`\n\n"
            "Please check the Tesseract path in the ⚙️ Settings panel on the left, "
            "then try again."
        )
        if preprocess_error:
            st.warning(
                f"Note: page preprocessing also reported an issue: `{preprocess_error}`"
            )
        if st.button("← Go back", type="primary"):
            st.session_state.step = "scanned_doc"
            st.rerun()
    else:
        # Step 5: OCR succeeded — transition to the format selection screen.
        st.session_state.step = "ocr_format_select"
        st.rerun()


# -----------------------------------------------------------------------------
# STEP: running_easyocr
# Second-pass OCR using EasyOCR, triggered when the faculty member chooses
# to try the enhanced scanner after reviewing the Tesseract preview.
# Follows the same structure as render_running_ocr(): runs inside a spinner,
# stores output in session_state, then transitions back to ocr_format_select
# so the user can compare the new result in the same preview UI.
# -----------------------------------------------------------------------------
def render_running_easyocr():
    """
    Purpose: Run EasyOCR as a second-pass scanner on the original uploaded PDF.

    EasyOCR uses a neural-network text detector and recogniser that handles
    low-quality scans, unusual fonts, and skewed pages better than Tesseract.
    It is slower and requires model weights (~200 MB, downloaded on first run).

    On success, overwrites ocr_pages_text, ocr_docx_bytes, ocr_pdf_bytes,
    ocr_quality_score, ocr_quality_label, and ocr_engine_used in session_state,
    then returns to ocr_format_select for the faculty member to review.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Read the Tesseract path before the spinner so preflight can check it.
    tesseract_path = st.session_state.get(
        "tesseract_path",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    )

    # ------------------------------------------------------------------
    # Preflight check — Tesseract is needed for preprocessing even though
    # EasyOCR handles the final text extraction. preprocess_pages() calls
    # pytesseract.image_to_osd() for orientation detection, which requires
    # the Tesseract binary to be present and callable.
    # ------------------------------------------------------------------
    preflight = preflight_check(
        tesseract_path=tesseract_path,
        output_paths=[tempfile.gettempdir()],
    )
    if not preflight["ok"]:
        st.error(
            "🔴 **Setup issue — cannot start enhanced OCR.**\n\n"
            + "\n\n".join(preflight["errors"])
        )
        if st.button("← Go back", type="primary"):
            st.session_state.step = "ocr_format_select"
            st.rerun()
        return

    easyocr_error = None

    with st.spinner(
        "Running the enhanced scanner — this may take a minute or two for "
        "longer documents. Hang tight..."
    ):
        try:
            # Re-run preprocessing to get corrected images.
            # The preprocessing log from the Tesseract pass is already stored
            # — we keep those entries and only update the OCR-specific lines.
            pil_images, preprocessing_log = preprocess_pages(
                st.session_state.file_bytes, tesseract_path
            )

            # Run EasyOCR on the corrected images.
            pages_text = run_easyocr(pil_images)

            quality_score, quality_label = score_ocr_quality(pages_text)

            # Build updated log: same preprocessing entries, new OCR line.
            processing_log = preprocessing_log + [
                "✔ OCR performed using EasyOCR (enhanced scanner)",
                f"✔ Quality estimate: "
                f"{'✅' if quality_label == 'Good' else '⚠️' if quality_label == 'Fair' else '❌'} "
                f"{quality_label} ({quality_score}%)",
            ]

            st.session_state.ocr_pages_text     = pages_text
            st.session_state.ocr_quality_score  = quality_score
            st.session_state.ocr_quality_label  = quality_label
            st.session_state.ocr_engine_used    = "easyocr"
            st.session_state.ocr_processing_log = processing_log

            # Clear the cached DOCX/PDF from the Tesseract pass so
            # render_ocr_format_select() rebuilds them from the updated text.
            st.session_state.ocr_docx_bytes = None
            st.session_state.ocr_pdf_bytes  = None

            # Build PDF for accessibility analysis only (not for download cache).
            ocr_pdf_bytes = build_pdf(pages_text)
            analysis_result = analyze_pdf(ocr_pdf_bytes)
            st.session_state.ocr_issues = (
                analysis_result["issues"]
                if analysis_result["status"] == "issues_found"
                else []
            )
            st.session_state.working_pdf_bytes = ocr_pdf_bytes

        except Exception as e:
            easyocr_error = str(e)

    if easyocr_error:
        st.error(
            "🔴 **Enhanced scan failed.**\n\n"
            f"**Error details:** {easyocr_error}\n\n"
            "You can still download the Tesseract version of your document."
        )
        if st.button("← Go back to previous result", type="primary"):
            st.session_state.ocr_engine_used = "tesseract"
            st.session_state.step = "ocr_format_select"
            st.rerun()
    else:
        st.session_state.step = "ocr_format_select"
        st.rerun()


# -----------------------------------------------------------------------------
# STEP: ocr_format_select
# Shown after OCR completes successfully. The user picks their output format
# (DOCX or PDF), downloads the readable file, then decides whether to continue
# with the accessibility issue workflow or end the session.
#
# Per CLAUDE.md OCR Steps 4-7: ask format → save with _readable suffix →
# show disclaimer note → ask how to continue.
# -----------------------------------------------------------------------------
def render_ocr_format_select():
    """
    Purpose: Show OCR quality results, ask whether the faculty member is happy,
    then either advance to accessibility analysis or offer a download to review.

    Flow per CLAUDE.md OCR Step 7:
      1. Show quality score, processing log, and text preview.
      2. Show what accessibility issues were found.
      3. Ask: "Are you happy with this result?" — three choices:
           a. Yes, let's keep going      → proceed to issue list (no download required)
           b. Download and review first  → show download buttons, then re-ask
           c. No, I want to stop         → done screen
      Downloads are always available at the end via render_choose_format().

    File naming per CLAUDE.md: original name + '_readable' before the extension.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    pages_text = st.session_state.ocr_pages_text
    page_count = len(pages_text)

    # -------------------------------------------------------------------------
    # BUILD OUTPUT FILES — lazy, cached
    # Built on first render so we never process a format the user won't use.
    # Cached in session state and reused on every rerender of this screen.
    # -------------------------------------------------------------------------
    if not st.session_state.ocr_docx_bytes:
        with st.spinner("Preparing your readable files..."):
            st.session_state.ocr_docx_bytes = build_docx(pages_text)
            st.session_state.ocr_pdf_bytes  = build_pdf(pages_text)

    st.success(
        f"Conversion complete — {page_count} page{'s' if page_count != 1 else ''} "
        "of readable text extracted successfully."
    )

    # -------------------------------------------------------------------------
    # QUALITY BADGE
    # Uses ✅ / ⚠️ / ❌ to stay visually distinct from the red/yellow/green
    # accessibility severity scheme used elsewhere in the tool.
    # -------------------------------------------------------------------------
    quality_score = st.session_state.get("ocr_quality_score")
    quality_label = st.session_state.get("ocr_quality_label", "Unknown")
    engine_used   = st.session_state.get("ocr_engine_used", "Tesseract")

    badge          = {"Good": "✅", "Fair": "⚠️", "Poor": "❌"}.get(quality_label, "⚪")
    engine_display = "Tesseract OCR" if engine_used == "tesseract" else "EasyOCR (enhanced)"

    st.markdown(
        f"**Text quality estimate:** {badge} {quality_label} "
        f"({quality_score}% word-like tokens) — scanned using {engine_display}"
    )

    if quality_label == "Good":
        st.caption(
            "The extracted text looks reliable. Review the preview below to confirm."
        )
    elif quality_label == "Fair":
        st.caption(
            "The extraction looks partially readable but some words may be inaccurate. "
            "Review the preview — if you see significant errors, try the enhanced scanner."
        )
    else:
        st.caption(
            "The extraction quality looks low — the text may contain significant errors. "
            "Review the preview and consider trying the enhanced scanner for a better result."
        )

    # -------------------------------------------------------------------------
    # PROCESSING LOG — "What we did"
    # Transparent step-by-step log: page count, orientation corrections, engine.
    # Collapsed by default so it's available without dominating the screen.
    # -------------------------------------------------------------------------
    processing_log = st.session_state.get("ocr_processing_log", [])
    if processing_log:
        with st.expander("📋 What we did", expanded=False):
            for entry in processing_log:
                st.markdown(entry)

    # -------------------------------------------------------------------------
    # TEXT PREVIEW — auto-opens when quality is Fair or Poor
    # -------------------------------------------------------------------------
    with st.expander(
        f"👁 Preview extracted text ({page_count} page{'s' if page_count != 1 else ''})",
        expanded=(quality_label in ("Fair", "Poor")),
    ):
        st.caption(
            "This is the text extracted from your scanned document. "
            "Review it for accuracy before deciding how to proceed."
        )
        for i, page_text in enumerate(pages_text):
            st.markdown(f"**Page {i + 1}**")
            st.text_area(
                label=f"page_{i + 1}",
                value=page_text.strip() if page_text.strip() else "(No text detected on this page)",
                height=200,
                disabled=True,
                label_visibility="collapsed",
                key=f"ocr_preview_page_{i}",
            )
            if i < len(pages_text) - 1:
                st.divider()

    # -------------------------------------------------------------------------
    # EASYOCR UPGRADE — offered when Tesseract was used and quality is not Good
    # -------------------------------------------------------------------------
    if engine_used == "tesseract" and quality_label in ("Fair", "Poor"):
        st.info(
            "**Not happy with the result?** I can try again using a more powerful "
            "scanner (EasyOCR) that handles difficult scans better. It takes a "
            "little longer, but often produces more accurate text."
        )
        if st.button("Try the enhanced scanner", use_container_width=True, key="try_easyocr"):
            st.session_state.step = "running_easyocr"
            st.rerun()
    elif engine_used == "tesseract" and quality_label == "Good":
        with st.expander("Want to try the enhanced scanner anyway?", expanded=False):
            st.caption(
                "The quality estimate looks good, but automated scoring isn't perfect. "
                "If you spot errors in the preview, you can still try the enhanced scanner."
            )
            if st.button(
                "Try the enhanced scanner",
                use_container_width=True,
                key="try_easyocr_optional",
            ):
                st.session_state.step = "running_easyocr"
                st.rerun()

    st.divider()

    # -------------------------------------------------------------------------
    # ISSUE SUMMARY — shown before asking how to continue so the faculty member
    # knows what they're deciding about before committing to a path.
    # -------------------------------------------------------------------------
    ocr_issues = st.session_state.get("ocr_issues", [])

    if ocr_issues:
        red_count    = sum(1 for i in ocr_issues if i["severity"] == "red")
        yellow_count = sum(1 for i in ocr_issues if i["severity"] == "yellow")
        green_count  = sum(1 for i in ocr_issues if i["severity"] == "green")

        summary_parts = []
        if red_count:
            summary_parts.append(f"**{red_count} 🔴 Red Light**")
        if yellow_count:
            summary_parts.append(f"**{yellow_count} 🟡 Yellow Light**")
        if green_count:
            summary_parts.append(f"**{green_count} 🟢 Green Light**")

        st.info(
            "Here are the issues we found in your converted document, grouped by "
            "severity: " + ", ".join(summary_parts) + ". "
            "If you'd like, I can help you work through most of them."
        )
    else:
        st.success(
            "🎉 Great news — your converted document looks good! No additional "
            "accessibility issues were found. Your students should be able to "
            "access this content without barriers."
        )

    # -------------------------------------------------------------------------
    # SATISFACTION QUESTION — the decision comes before downloads.
    # Downloading is never a gate to proceeding; it's an optional side-path.
    # -------------------------------------------------------------------------
    if not st.session_state.get("ocr_download_mode", False):

        st.markdown("**Are you happy with this result?**")
        col1, col2, col3 = st.columns(3)

        with col1:
            if st.button(
                "Yes, let's keep going.",
                type="primary",
                use_container_width=True,
                key="ocr_continue_direct",
            ):
                if ocr_issues:
                    st.session_state.issues = ocr_issues
                    st.session_state.step   = "issue_list"
                else:
                    st.session_state.step = "no_issues"
                st.rerun()

        with col2:
            if st.button(
                "I want to download and review it first.",
                use_container_width=True,
                key="ocr_download_first",
            ):
                st.session_state.ocr_download_mode = True
                st.rerun()

        with col3:
            if st.button(
                "No, I want to stop.",
                use_container_width=True,
                key="ocr_stop_direct",
            ):
                st.session_state.step        = "done"
                st.session_state.done_reason = "ocr_exit"
                st.rerun()

    else:
        # ---------------------------------------------------------------------
        # DOWNLOAD MODE — shown after "I want to download and review it first"
        # Downloads are presented, then the same continue/stop choice is offered.
        # ---------------------------------------------------------------------
        st.markdown(
            "Your converted file is ready to download. Choose a format below — "
            "it will be saved with **_readable** added to your original file name."
        )

        base_name = st.session_state.file_name.replace(".pdf", "").replace(".PDF", "")
        docx_name = f"{base_name}_readable.docx"
        pdf_name  = f"{base_name}_readable.pdf"

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Word processor document (DOCX)**")
            st.caption(
                "Opens in Microsoft Word, Google Docs, or any compatible application. "
                "Best if you need to edit the content further before re-publishing."
            )
            st.download_button(
                label=f"⬇ Download {docx_name}",
                data=st.session_state.ocr_docx_bytes,
                file_name=docx_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )
        with col2:
            st.markdown("**New PDF document**")
            st.caption(
                "A text-based PDF ready to share or upload directly. "
                "Best if you want a finished file."
            )
            st.download_button(
                label=f"⬇ Download {pdf_name}",
                data=st.session_state.ocr_pdf_bytes,
                file_name=pdf_name,
                mime="application/pdf",
                use_container_width=True,
            )

        st.caption(f'📄 Note added to your file: "{OCR_DISCLAIMER}"')
        st.divider()

        st.markdown("**Ready to keep going?**")
        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "Yes, let's keep going.",
                type="primary",
                use_container_width=True,
                key="ocr_continue_after_download",
            ):
                st.session_state.ocr_download_mode = False
                if ocr_issues:
                    st.session_state.issues = ocr_issues
                    st.session_state.step   = "issue_list"
                else:
                    st.session_state.step = "no_issues"
                st.rerun()
        with col2:
            if st.button(
                "No, I'm done for now.",
                use_container_width=True,
                key="ocr_done_after_download",
            ):
                st.session_state.ocr_download_mode = False
                st.session_state.step              = "done"
                st.session_state.done_reason       = "ocr_exit"
                st.rerun()

    _render_go_back("scanned_doc")


# -----------------------------------------------------------------------------
# STEP: no_issues
# Verbatim clean document message from CLAUDE.md.
# -----------------------------------------------------------------------------
def render_no_issues():
    """
    Purpose: Tell the user the document looks good. Offer download or end session.
    Verbatim message from CLAUDE.md.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Verbatim message from CLAUDE.md
    st.success(
        "🎉 Great news — your document looks good! We didn't find any accessibility "
        "issues that need attention. Your students should be able to access this content "
        "without barriers. Nice work!"
    )

    st.markdown("**What would you like to do?**")

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Download my file",
            type="primary",
            use_container_width=True,
        ):
            # Offer the original file as a download since no changes were made.
            st.download_button(
                label="Download PDF",
                data=st.session_state.file_bytes,
                file_name=st.session_state.file_name,
                mime="application/pdf",
            )
    with col2:
        if st.button("I'm all set, thanks.", use_container_width=True):
            st.session_state.step = "done"
            st.session_state.done_reason = "clean_exit"
            st.rerun()

    _render_go_back("upload")


# -----------------------------------------------------------------------------
# STEP: issue_list
# Shows the prioritized issue list grouped by severity.
# The user clicks any issue to begin working on it.
# -----------------------------------------------------------------------------
def render_issue_list():
    """
    Purpose: Present all remaining issues, grouped by severity (Red → Yellow → Green).
    User clicks any issue to start resolving it in any order they choose.
    Per CLAUDE.md: display the Priority Recommendation message verbatim first.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Count remaining unresolved, unskipped issues
    remaining = [
        issue for issue in st.session_state.issues
        if issue["id"] not in st.session_state.resolved_ids
        and issue["id"] not in st.session_state.get("skipped_ids", [])
    ]

    if not remaining:
        # All issues resolved — go to manual review checklist before download
        st.session_state.step = "manual_review"
        st.rerun()
        return

    # File context
    st.markdown(f"**File:** {st.session_state.file_name}")

    # Priority Recommendation — verbatim from CLAUDE.md
    st.info(
        "We strongly recommend starting with the Red Light and Yellow Light issues. "
        "These create the most significant barriers for your students and are the ones "
        "that bring your document into compliance with Title II accessibility requirements. "
        "Where would you like to start?"
    )

    # Group remaining issues by severity for display
    # Order: red first, then yellow, then green — most serious to least serious
    for severity in ["red", "yellow", "green"]:
        severity_issues = [i for i in remaining if i["severity"] == severity]
        if not severity_issues:
            continue

        st.markdown(f"### {SEVERITY_LABELS[severity]}")

        for issue in severity_issues:
            # Each issue renders as a clickable button.
            # The button label is the issue title.
            if st.button(
                f"{issue['title']}",
                key=f"issue_btn_{issue['id']}",
                use_container_width=True,
            ):
                st.session_state.current_issue_id = issue["id"]
                st.session_state.proposed_fix = issue["fix_preview"]
                if issue["id"] == "untagged_pdf":
                    # Ask whether they have the original source file first.
                    st.session_state.step = "source_file_question"
                elif issue["id"] == "missing_alt_text":
                    # Launch the per-image alt text workflow.
                    # Initialise all workflow state from the issue's fix_data
                    # so the sub-steps always start fresh.
                    st.session_state.alt_text_images = issue.get("fix_data", {}).get("images", [])
                    st.session_state.alt_text_index = 0
                    st.session_state.alt_text_phase = "classifying"  # auto-classify first
                    st.session_state.alt_text_results = {}
                    st.session_state.alt_text_auto_classifications = {}
                    st.session_state.step = "image_alt_text"
                else:
                    st.session_state.step = "resolving_issue"
                st.rerun()

        st.markdown("")  # Spacing between severity groups

    # Progress indicator
    total    = len(st.session_state.issues)
    resolved = len(st.session_state.resolved_ids)
    skipped  = len(st.session_state.get("skipped_ids", []))
    st.divider()
    caption = f"{resolved} of {total} issues resolved"
    if skipped:
        caption += f" · {skipped} skipped"
    st.caption(caption)

    _render_go_back("upload")


# =============================================================================
# STEP: image_alt_text
# =============================================================================
# Per-image alt text workflow. Triggered when the faculty member selects the
# "Images with missing text descriptions" issue from the issue list.
#
# Sub-phases (stored in st.session_state.alt_text_phase):
#   classifying  — runs auto_classify_image() on all images (transient spinner)
#   bulk_review  — shows high-confidence detections for one-click bulk removal
#   question_1   — three-option choice: decorative / artifact / informational
#   question_2   — critical or supplementary? (only if informational)
#   enter_text   — enter or edit the alt text description
#   summary      — review all decisions before confirming
#
# Per CLAUDE.md Image Remediation Workflow:
#   Background / Edge artifact  → Green, excluded from DOCX output
#   Decorative                  → Green, empty alt text, no description needed
#   Not decorative, critical    → stays Red, full description required
#   Not decorative, not critical → Yellow, brief description required
# =============================================================================

def _render_go_back(destination):
    """
    Render the linear "← Go back" button for non-remediation workflow screens.

    This is simple step navigation — it moves the user to the preceding screen
    without touching any session state beyond `step`. No work is lost.

    Parameters:
        destination (str): The step name to return to (e.g. "upload", "scanned_doc").
    """
    st.divider()
    if st.button("← Go back", key=f"go_back_{destination}"):
        st.session_state.step = destination
        st.rerun()


def _render_abandon_button():
    """
    Render the global "✕ Abandon current process" escape button.

    Shown at the bottom of every remediation and human-in-the-loop screen.
    Clicking it discards whatever is currently in progress (any uncommitted
    decision for the current issue/image) and returns to the issue list,
    which is always the last completed summary state.

    Completed work is never lost — resolved_ids and applied fixes remain in
    session state. Only the in-progress interaction is abandoned.
    """
    st.divider()
    if st.button(
        "✕ Abandon current process",
        use_container_width=False,
        key=f"abandon_{st.session_state.get('step', 'unknown')}_{st.session_state.get('alt_text_index', 0)}",
        help="Stop what you're doing and return to the issue list. Completed fixes are kept.",
    ):
        # Clear any in-progress issue state so the issue list starts clean.
        st.session_state.current_issue_id    = None
        st.session_state.proposed_fix        = None
        st.session_state.showing_alternative = False
        st.session_state.step                = "issue_list"
        st.rerun()


def _render_image_nav_buttons(index, total, img_key, current, results):
    """
    Render the three navigation buttons shown at the bottom of every per-image
    phase: Previous, Skip, and Stop.

    - Previous: goes back to the preceding image and clears any partial result
      recorded for the current image, so the faculty member can re-answer.
      Hidden on the first image.
    - Skip: records the image as skipped (no decision made) and advances.
      Skipped images appear in the summary so faculty can revisit them.
    - Stop: jumps straight to the summary screen regardless of how many images
      remain. Any unclassified images will be visible there.

    Parameters:
        index   (int):        0-based index of the current image.
        total   (int):        Total number of images in the workflow.
        img_key (int|str):    Key identifying the current image in alt_text_results.
                              An int xref for pymupdf images; a string for docling images.
        current (dict):       The current image's metadata dict.
        results (dict):       The live alt_text_results dict from session state.
    """
    st.divider()
    nav_cols = st.columns(3)

    with nav_cols[0]:
        # Previous — hidden for the first image (nowhere to go back to).
        if index > 0:
            if st.button(
                "← Previous image",
                use_container_width=True,
                key=f"nav_prev_{index}",
            ):
                # Clear any partial answer for the current image so it starts
                # fresh when the faculty member returns to it.
                results.pop(img_key, None)
                st.session_state.alt_text_results = results
                st.session_state.alt_text_index = index - 1
                st.session_state.alt_text_phase = "question_1"
                st.rerun()

    with nav_cols[1]:
        if st.button(
            "Skip this image →",
            use_container_width=True,
            key=f"nav_skip_{index}",
        ):
            # Record as skipped so the summary can flag it.
            results[img_key] = {
                "type": "skipped",
                "severity": "green",
                "alt_text": "",
                "page": current["page"],
            }
            st.session_state.alt_text_results = results
            _advance_alt_text_image()

    with nav_cols[2]:
        if st.button(
            "Skip fixing images",
            use_container_width=True,
            key=f"nav_stop_{index}",
        ):
            st.session_state.alt_text_phase = "summary"
            st.rerun()


def _advance_alt_text_image():
    """
    Move to the next image in the sequence, or to the summary phase when all
    images have been processed. Always resets the sub-phase to question_1 for
    the next image.
    """
    index = st.session_state.alt_text_index
    total = len(st.session_state.alt_text_images)
    next_index = index + 1
    if next_index >= total:
        st.session_state.alt_text_phase = "summary"
    else:
        st.session_state.alt_text_index = next_index
        st.session_state.alt_text_phase = "question_1"
    st.rerun()


def render_image_alt_text():
    """
    Purpose: Walk the faculty member through each image one at a time,
    collecting a decision (decorative vs. informational) and a text description
    for each non-decorative image.

    Per CLAUDE.md: severity is reassigned based on the faculty member's answers —
    never apply a description without their explicit confirmation.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    images  = st.session_state.alt_text_images
    index   = st.session_state.alt_text_index
    phase   = st.session_state.alt_text_phase
    total   = len(images)
    results = st.session_state.alt_text_results
    auto_classifications = st.session_state.get("alt_text_auto_classifications", {})

    HIGH_CONFIDENCE = 0.75  # threshold for bulk review inclusion

    # -------------------------------------------------------------------------
    # CLASSIFYING PHASE — run auto-classification on all images (transient)
    # -------------------------------------------------------------------------
    if phase == "classifying":
        pdf_source = st.session_state.working_pdf_bytes or st.session_state.file_bytes
        with st.spinner(
            f"Reviewing {total} image{'s' if total != 1 else ''} — "
            "checking for backgrounds and scan artifacts..."
        ):
            classifications = run_auto_classification(images, pdf_source)
        st.session_state.alt_text_auto_classifications = classifications

        # Decide next phase: bulk_review if any high-confidence non-content detections
        # exist, otherwise go straight to per-image question_1.
        # classifications is keyed by img_key (int for pymupdf, str for docling).
        flagged = [
            img_key for img_key, r in classifications.items()
            if r["classification"] != "content" and r["confidence"] >= HIGH_CONFIDENCE
        ]
        if flagged:
            st.session_state.alt_text_phase = "bulk_review"
        else:
            st.session_state.alt_text_phase = "question_1"
        st.rerun()
        return

    # -------------------------------------------------------------------------
    # BULK REVIEW PHASE — one-click removal of high-confidence detections
    # -------------------------------------------------------------------------
    if phase == "bulk_review":
        flagged_images = [
            img for img in images
            if auto_classifications.get(img["img_key"], {}).get("classification") != "content"
            and auto_classifications.get(img["img_key"], {}).get("confidence", 0) >= HIGH_CONFIDENCE
        ]

        n = len(flagged_images)
        bg_count  = sum(1 for img in flagged_images
                        if auto_classifications[img["img_key"]]["classification"] == "background")
        ea_count  = sum(1 for img in flagged_images
                        if auto_classifications[img["img_key"]]["classification"] == "edge_artifact")
        art_count = sum(1 for img in flagged_images
                        if auto_classifications[img["img_key"]]["classification"] == "artifact")

        # Build a plain-language summary of what was found.
        found_parts = []
        if bg_count:
            found_parts.append(f"{bg_count} page background{'s' if bg_count != 1 else ''}")
        if ea_count:
            found_parts.append(f"{ea_count} edge artifact{'s' if ea_count != 1 else ''}")
        if art_count:
            found_parts.append(f"{art_count} scan artifact{'s' if art_count != 1 else ''}")

        st.markdown("### I reviewed your images automatically")
        st.info(
            f"I found **{', '.join(found_parts)}** that look like they should be removed "
            f"rather than described. Here's what I spotted:"
        )

        pdf_source = st.session_state.working_pdf_bytes or st.session_state.file_bytes

        # Show thumbnails of flagged images in a 3-column grid.
        cols = st.columns(3)
        for i, img_meta in enumerate(flagged_images):
            img_key     = img_meta["img_key"]
            clf         = auto_classifications[img_key]
            label_map   = {
                "background":   "Page background",
                "edge_artifact": "Edge artifact",
                "artifact":     "Scan artifact",
            }
            label = label_map.get(clf["classification"], clf["classification"])
            with cols[i % 3]:
                img_bytes, _ = extract_image_by_source(pdf_source, img_meta)
                if img_bytes:
                    st.image(img_bytes, use_container_width=True)
                st.caption(f"**{label}** — page {img_meta['page']}")
                st.caption(clf["reason"])

        st.divider()
        st.markdown("**What would you like to do?**")

        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                f"Remove all {n} flagged items",
                type="primary",
                use_container_width=True,
                key="bulk_remove_all",
            ):
                # Apply removal to all flagged images in alt_text_results.
                for img_meta in flagged_images:
                    img_key = img_meta["img_key"]
                    clf     = auto_classifications[img_key]
                    # background and edge_artifact both stored as "artifact" type
                    # so build_docx_from_pdf skips them consistently.
                    results[img_key] = {
                        "type": clf["classification"],  # "background", "edge_artifact", or "artifact"
                        "severity": "green",
                        "alt_text": "",
                        "page": img_meta["page"],
                    }
                st.session_state.alt_text_results = results
                # Advance to first unclassified image, or summary if all done.
                first_idx = next(
                    (i for i, img in enumerate(images) if img["img_key"] not in results),
                    None,
                )
                if first_idx is not None:
                    st.session_state.alt_text_index = first_idx
                    st.session_state.alt_text_phase = "question_1"
                else:
                    st.session_state.alt_text_phase = "summary"
                st.rerun()

        with col2:
            if st.button(
                "Let me review each one individually",
                use_container_width=True,
                key="bulk_review_each",
            ):
                # Don't apply anything — go through normal flow.
                # Auto-classification badges will still appear as suggestions.
                st.session_state.alt_text_phase = "question_1"
                st.rerun()

        _render_abandon_button()
        return

    # -------------------------------------------------------------------------
    # SUMMARY PHASE — shown after all images have been processed
    # -------------------------------------------------------------------------
    if phase == "summary":
        st.markdown("### Review your image decisions")
        st.markdown(
            "Here's what we have for each image. Confirm to apply all decisions, "
            "or go back to make changes."
        )
        st.divider()

        # Partition images into four groups for the summary.
        REMOVED_TYPES = ("artifact", "background", "edge_artifact")
        removed      = [img for img in images if results.get(img["img_key"], {}).get("type") in REMOVED_TYPES]
        skipped      = [img for img in images if results.get(img["img_key"], {}).get("type") == "skipped"
                        or img["img_key"] not in results]
        others       = [img for img in images if results.get(img["img_key"], {}).get("type")
                        not in REMOVED_TYPES and results.get(img["img_key"], {}).get("type") not in (None, "skipped")]

        if removed:
            type_label_map = {
                "artifact":      "Scan artifact",
                "background":    "Page background",
                "edge_artifact": "Edge artifact",
            }
            st.warning(
                f"**{len(removed)} image{'s' if len(removed) != 1 else ''} will be "
                "excluded from the accessible output.**"
            )
            for img in removed:
                img_key = img["img_key"]
                pg      = results.get(img_key, {}).get("page", img.get("page", "?"))
                kind    = results.get(img_key, {}).get("type", "artifact")
                label   = type_label_map.get(kind, "Removed image")
                with st.expander(f"Page {pg} — {label} (will be excluded)", expanded=False):
                    img_bytes, _ = extract_image_by_source(
                        st.session_state.working_pdf_bytes or st.session_state.file_bytes,
                        img,
                    )
                    if img_bytes:
                        st.image(img_bytes, width=300)
                    st.info("This image will be excluded from the Word document output.")
            st.divider()

        if skipped:
            st.error(
                f"**{len(skipped)} image{'s' if len(skipped) != 1 else ''} "
                f"{'were' if len(skipped) != 1 else 'was'} skipped or not reviewed.** "
                "You can confirm anyway or go back to classify them."
            )
            for img in skipped:
                pg   = img.get("page", "?")
                with st.expander(f"Page {pg} — Not classified", expanded=False):
                    img_bytes, _ = extract_image_by_source(
                        st.session_state.working_pdf_bytes or st.session_state.file_bytes,
                        img,
                    )
                    if img_bytes:
                        st.image(img_bytes, width=300)
                    st.caption("No decision was made for this image.")
            st.divider()

        for img in others:
            img_key = img["img_key"]
            result  = results.get(img_key, {})
            pg      = result.get("page", img.get("page", "?"))
            sev     = result.get("severity", "")
            label   = SEVERITY_LABELS.get(sev, "")
            kind    = result.get("type", "")
            alt     = result.get("alt_text", "")

            header = f"Page {pg} — {label}"
            with st.expander(header, expanded=False):
                img_bytes, _ = extract_image_by_source(
                    st.session_state.working_pdf_bytes or st.session_state.file_bytes,
                    img,
                )
                if img_bytes:
                    st.image(img_bytes, width=300)

                if kind == "decorative":
                    st.success("Marked as decorative — no description needed.")
                else:
                    st.markdown(f"**Description:** {alt}")

        st.divider()

        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "Yes, that looks good.",
                type="primary",
                use_container_width=True,
                key="alt_confirm",
            ):
                # Mark the issue as resolved and store results for DOCX output.
                if "missing_alt_text" not in st.session_state.resolved_ids:
                    st.session_state.resolved_ids.append("missing_alt_text")
                    st.session_state.resolved_count += 1

                # Trigger re-analysis check per CLAUDE.md (every 3 resolved issues).
                if (
                    st.session_state.resolved_count % 3 == 0
                    and not st.session_state.reanalysis_done
                ):
                    st.session_state.step = "re_analysis"
                else:
                    st.session_state.reanalysis_done = False
                    st.session_state.step = "continue_or_stop"
                st.rerun()

        with col2:
            if st.button(
                "No, let me go back and edit.",
                use_container_width=True,
                key="alt_redo",
            ):
                # Restart from the first image without clearing results —
                # existing answers are preserved so the faculty member only
                # needs to change what's wrong.
                st.session_state.alt_text_index = 0
                st.session_state.alt_text_phase = "question_1"
                st.rerun()
        _render_abandon_button()
        return

    # -------------------------------------------------------------------------
    # Safety net — if index is out of range, jump to summary
    # -------------------------------------------------------------------------
    if index >= total:
        st.session_state.alt_text_phase = "summary"
        st.rerun()
        return

    # -------------------------------------------------------------------------
    # PER-IMAGE HEADER — progress indicator and image previews
    # -------------------------------------------------------------------------
    current  = images[index]
    img_key  = current["img_key"]
    page_num = current["page"] - 1  # convert 1-based page number to 0-based index

    st.markdown(f"**Image {index + 1} of {total}** — found on page {current['page']}")
    st.progress(index / total)
    st.divider()

    pdf_source = st.session_state.working_pdf_bytes or st.session_state.file_bytes

    # --- Panel 1: Page context view with red highlight ---
    # Shows the full page so faculty can see exactly which image is being assessed.
    # Critical when a page has multiple images or when the extracted crop is ambiguous.
    # Branch on source: pymupdf images use xref, docling images use bbox.
    if current.get("source") == "docling":
        page_view = render_page_with_image_highlight(
            pdf_source, page_num, bbox=current.get("bbox")
        )
    else:
        page_view = render_page_with_image_highlight(
            pdf_source, page_num, xref=current.get("xref")
        )
    if page_view:
        st.caption("📍 The image being assessed is highlighted in red:")
        st.image(page_view, use_container_width=True)
    st.divider()

    # --- Panel 2: Extracted image at fixed width ---
    # Clean crop of the image alone — easier to read the content in detail.
    img_bytes, _ = extract_image_by_source(pdf_source, current)
    if img_bytes:
        st.caption("🖼 Extracted image:")
        st.image(img_bytes, width=500)

        # --- Panel 3: Full-size zoom (collapsed by default) ---
        with st.expander("🔍 View full size", expanded=False):
            st.image(img_bytes, use_container_width=True)
    elif not page_view:
        # Both views unavailable — show a fallback message.
        st.caption("_(Image preview unavailable for this item)_")

    st.divider()

    # -------------------------------------------------------------------------
    # QUESTION 1 — Classify the image: decorative, artifact, or informational?
    # Per CLAUDE.md: three options. Show auto-classification badge if available.
    # -------------------------------------------------------------------------
    if phase == "question_1":
        # Auto-classification suggestion badge — shown when confidence is below
        # the HIGH_CONFIDENCE threshold (those were handled in bulk_review) but
        # still meaningfully non-"content". We show all non-content suggestions here.
        clf = auto_classifications.get(img_key, {})
        if clf.get("classification") and clf["classification"] != "content":
            label_map = {
                "background":    "a page background",
                "edge_artifact": "a scan edge artifact",
                "artifact":      "a scan artifact",
            }
            suggestion = label_map.get(clf["classification"], clf["classification"])
            conf_pct   = int(clf.get("confidence", 0) * 100)
            st.warning(
                f"⚠️ **I think this might be {suggestion}** ({conf_pct}% confidence). "
                f"Reason: {clf.get('reason', '')}  \n"
                "You can accept my suggestion by choosing the matching option below, "
                "or override it if I got it wrong."
            )

        st.markdown("### What is this image?")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("**Decorative**")
            st.caption(
                "Exists only for visual appeal — a border, texture, or stock photo "
                "with no connection to the topic. No description needed."
            )
            if st.button(
                "It's decorative",
                type="primary",
                use_container_width=True,
                key=f"q1_decorative_{index}",
            ):
                # Decorative → Green, empty alt text. No further review needed.
                results[img_key] = {
                    "type": "decorative",
                    "severity": "green",
                    "alt_text": "",
                    "page": current["page"],
                }
                st.session_state.alt_text_results = results
                _advance_alt_text_image()

        with col2:
            st.markdown("**Scan artifact**")
            st.caption(
                "A smudge, shadow, dark edge, or other visual noise from scanning. "
                "Will be excluded from the accessible output."
            )
            if st.button(
                "It's a scan artifact — remove it",
                use_container_width=True,
                key=f"q1_artifact_{index}",
            ):
                # Artifact → Green, excluded from DOCX output.
                # The PDF working copy is not modified — exclusion applies at build time.
                results[img_key] = {
                    "type": "artifact",
                    "severity": "green",
                    "alt_text": "",
                    "page": current["page"],
                }
                st.session_state.alt_text_results = results
                _advance_alt_text_image()

        with col3:
            st.markdown("**Informational**")
            st.caption(
                "Carries content students need — a chart, diagram, photo, figure, "
                "or any image that isn't purely decorative. Needs a description."
            )
            if st.button(
                "It carries information",
                use_container_width=True,
                key=f"q1_info_{index}",
            ):
                st.session_state.alt_text_phase = "question_2"
                st.rerun()

        _render_image_nav_buttons(index, total, img_key, current, results)

    # -------------------------------------------------------------------------
    # QUESTION 2 — Is it critical to understanding?
    # Per CLAUDE.md: verbatim question with example types.
    # -------------------------------------------------------------------------
    elif phase == "question_2":
        st.markdown("### Does understanding this image affect how your students understand the document?")
        st.markdown(
            "An image is critical to understanding if it presents information that "
            "isn't described in the surrounding text — for example, a chart, diagram, "
            "map, or figure that students need to interpret the content."
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "Yes, it's important for understanding.",
                type="primary",
                use_container_width=True,
                key=f"q2_yes_{index}",
            ):
                # Critical → stays Red, full description required.
                results[img_key] = {
                    "type": "critical",
                    "severity": "red",
                    "alt_text": "",
                    "page": current["page"],
                }
                st.session_state.alt_text_results = results
                st.session_state.alt_text_phase = "enter_text"
                st.rerun()

        with col2:
            if st.button(
                "No, it adds context but isn't essential.",
                use_container_width=True,
                key=f"q2_no_{index}",
            ):
                # Non-critical → downgraded to Yellow, brief description needed.
                results[img_key] = {
                    "type": "supplementary",
                    "severity": "yellow",
                    "alt_text": "",
                    "page": current["page"],
                }
                st.session_state.alt_text_results = results
                st.session_state.alt_text_phase = "enter_text"
                st.rerun()

        # Back to Question 1 to change the classification entirely.
        st.divider()
        if st.button(
            "← Back to classification",
            key=f"q2_back_{index}",
            use_container_width=False,
        ):
            st.session_state.alt_text_phase = "question_1"
            st.rerun()

        _render_image_nav_buttons(index, total, img_key, current, results)

    # -------------------------------------------------------------------------
    # ENTER TEXT — write the alt text description
    # Per CLAUDE.md: generate a description and ask the faculty member to
    # review, edit, or approve it before it is applied.
    # -------------------------------------------------------------------------
    elif phase == "enter_text":
        result   = results.get(img_key, {})
        severity = result.get("severity", "red")
        existing = result.get("alt_text", "")

        if severity == "red":
            st.markdown("### 🔴 This image needs a full text description")
            st.markdown(
                "Write a description that conveys everything a student would need "
                "to know from this image — what it shows, why it matters in context, "
                "and any data or key information it presents. A student who cannot "
                "see this image should be able to fully understand it from your words."
            )
        else:
            st.markdown("### 🟡 This image needs a brief description")
            st.markdown(
                "Write a short description that gives students useful context — "
                "what the image is and its general relationship to the surrounding "
                "content. It doesn't need to capture every detail."
            )

        new_text = st.text_area(
            "Text description (alt text):",
            value=existing,
            height=120,
            key=f"alt_input_{index}",
            placeholder="Describe the image here...",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "Save this description.",
                type="primary",
                use_container_width=True,
                key=f"alt_save_{index}",
            ):
                if new_text.strip():
                    results[img_key]["alt_text"] = new_text.strip()
                    st.session_state.alt_text_results = results
                    _advance_alt_text_image()
                else:
                    st.warning("Please enter a description before saving.")

        with col2:
            if st.button(
                "← Back to question",
                use_container_width=True,
                key=f"alt_back_{index}",
            ):
                st.session_state.alt_text_phase = "question_2"
                st.rerun()

        _render_image_nav_buttons(index, total, img_key, current, results)
        _render_abandon_button()


# -----------------------------------------------------------------------------
# STEP: source_file_question
# Intercept screen shown when the user selects the "Untagged PDF" issue.
# Asks whether they still have the original Word / Google Docs / PowerPoint /
# Google Slides source file before we attempt any PDF-level fix.
# Per CLAUDE.md: two choices only — yes (→ cuny_resources) or no (→ resolving_issue).
# -----------------------------------------------------------------------------
def render_source_file_question():
    """
    Purpose: Before tackling an untagged PDF, find out whether the faculty
    member has the original source file. Working from the source is always
    the more reliable path — we want to route them there when possible.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    st.markdown("### Before we dig in — a quick question.")
    st.markdown(
        "The most reliable fix for an untagged PDF is to update the **original "
        "source file** directly. We're building a dedicated tool to help you do "
        "that — but in the meantime, great resources are available."
    )
    st.markdown(
        "Do you still have the original **Word, Google Docs, PowerPoint, or "
        "Google Slides** file you used to create this PDF?"
    )

    st.divider()

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Yes, I have the original file.",
            type="primary",
            use_container_width=True,
        ):
            st.session_state.step = "cuny_resources"
            st.rerun()
    with col2:
        if st.button(
            "No, I only have the PDF.",
            use_container_width=True,
        ):
            # Skip the source-file path and go straight to the untagged PDF
            # fix screen, same as any other issue.
            st.session_state.step = "resolving_issue"
            st.rerun()

    _render_abandon_button()


# -----------------------------------------------------------------------------
# STEP: cuny_resources
# Shown when the faculty member confirms they have the original source file.
# Presents CUNY accessibility guides and two forward paths:
#   1. Go work on the source file (mark the untagged issue as acknowledged
#      and return to the issue list for remaining issues).
#   2. Keep going with the PDF (proceed to the untagged PDF fix screen).
# Per CLAUDE.md: exactly two choices.
# -----------------------------------------------------------------------------
def render_cuny_resources():
    """
    Purpose: Give faculty members who have their original source file the
    CUNY accessibility resources they need to fix it properly, then let them
    choose whether to work on that file now or continue remediating the PDF.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    st.markdown(
        "Working from your original source file will give you the best result — "
        "Word, Google Docs, and PowerPoint all export properly tagged, accessible "
        "PDFs when the source file is set up correctly."
    )

    st.info(
        "We're building a source-file editor directly into this tool. "
        "While that's in progress, CUNY's accessibility guides walk you through "
        "the process step by step."
    )

    # Display CUNY resources as clearly labeled, clickable links.
    # All three links open in a new tab so the user doesn't leave the tool.
    st.markdown("### CUNY Accessibility Resources")
    st.markdown(
        """
        - 📚 [**CUNY Accessibility Toolkit**](https://guides.cuny.edu/accessibility)
          — Overview of accessibility standards and tools across file formats

        - 📝 [**Making Word Documents Accessible**](https://guides.cuny.edu/accessibility/microsoft_word)
          — Step-by-step guide for heading styles, alt text, reading order, and more in Word

        - 📊 [**Making PowerPoint Presentations Accessible**](https://guides.cuny.edu/accessibility/powerpoint)
          — Step-by-step guide for slide titles, alt text, reading order, and more in PowerPoint
        """
    )

    st.markdown(
        "_If your file is a **Google Doc or Google Slides** presentation, "
        "open it, go to **File → Download**, choose **Microsoft Word (.docx)** "
        "or **Microsoft PowerPoint (.pptx)**, then follow the Word or PowerPoint "
        "guide above._"
    )

    st.divider()

    st.markdown("**What would you like to do next?**")

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Review these resources and update my original file.",
            type="primary",
            use_container_width=True,
        ):
            # Mark the untagged PDF issue as acknowledged so it no longer
            # appears in the issue list. The faculty member is handling it
            # via their source file — no PDF-level fix needed.
            if "untagged_pdf" not in st.session_state.resolved_ids:
                st.session_state.resolved_ids.append("untagged_pdf")
                st.session_state.resolved_count += 1
            st.session_state.current_issue_id = None
            st.session_state.step = "issue_list"
            st.rerun()
    with col2:
        if st.button(
            "Keep going with this PDF.",
            use_container_width=True,
        ):
            # Proceed to the standard untagged PDF fix screen.
            st.session_state.step = "resolving_issue"
            st.rerun()

    _render_abandon_button()


# -----------------------------------------------------------------------------
# STEP: resolving_issue
# Shows the issue details, applies a placeholder fix, then goes to QA step.
# -----------------------------------------------------------------------------
def render_resolving_issue():
    """
    Purpose: Show the selected issue and its proposed fix for faculty review.
    Per CLAUDE.md Per-Issue QA Process:
      Step 1 — Internal self-check (not shown to user — we simulate this).
      Step 2 — Faculty confirmation with exact verbatim message.
    When the faculty member clicks "No, let's try again.", an issue-specific
    alternative is offered before falling back to returning to the issue list.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Find the current issue by ID
    issue = next(
        (i for i in st.session_state.issues if i["id"] == st.session_state.current_issue_id),
        None,
    )

    if issue is None:
        # Safety net: if we can't find the issue, go back to the list
        st.session_state.step = "issue_list"
        st.rerun()
        return

    # Show which issue we're working on
    st.markdown(f"### {SEVERITY_LABELS[issue['severity']]} — {issue['title']}")
    st.markdown(f"*{issue['wcag']} | {issue['ada']}*")
    st.divider()

    # Explain the issue and why it matters
    st.markdown("**Why this matters for your students:**")
    st.markdown(issue["description"])

    st.divider()

    # =========================================================================
    # Shared apply-and-advance logic used by both standard and alternative mode.
    # Extracted here to avoid repeating the resolved_count / step logic below.
    # =========================================================================
    def _apply_and_advance():
        """Apply the fix for the current issue and advance to the next step."""
        fix_fn = FIX_DISPATCH.get(issue["id"])
        if fix_fn and st.session_state.working_pdf_bytes:
            new_bytes, _ = fix_fn(
                st.session_state.working_pdf_bytes,
                issue.get("fix_data", {}),
            )
            st.session_state.working_pdf_bytes = new_bytes
        st.session_state.resolved_ids.append(issue["id"])
        st.session_state.resolved_count      += 1
        st.session_state.current_issue_id    = None
        st.session_state.proposed_fix        = None
        st.session_state.showing_alternative = False
        if (
            st.session_state.resolved_count % 3 == 0
            and not st.session_state.reanalysis_done
        ):
            st.session_state.step = "re_analysis"
        else:
            st.session_state.reanalysis_done = False
            st.session_state.step = "continue_or_stop"
        st.rerun()

    def _skip_issue():
        """Mark the current issue as skipped and return to the issue list."""
        st.session_state.skipped_ids.append(issue["id"])
        st.session_state.current_issue_id    = None
        st.session_state.showing_alternative = False
        st.session_state.step                = "issue_list"
        st.rerun()

    def _docx_path_buttons(primary_label="Got it — apply it anyway"):
        """Two-button footer shared by every DOCX-path alternative-mode branch."""
        col1, col2 = st.columns(2)
        with col1:
            if st.button(primary_label, type="primary", use_container_width=True, key="alt_yes"):
                _apply_and_advance()
        with col2:
            if st.button("Skip this issue for now", use_container_width=True, key="alt_skip"):
                _skip_issue()

    # =========================================================================
    # ALTERNATIVE MODE
    # Shown after the faculty member clicks "No, let's try again."
    # Each issue type gets a custom path: more control, or a plain explanation
    # of how the fix works so they can decide whether to proceed or skip it.
    # =========================================================================
    if st.session_state.get("showing_alternative", False):

        st.markdown("### Let's try a different approach")

        # --- missing_title: let them type a completely fresh title ---
        if issue["id"] == "missing_title":
            st.markdown(
                "No problem — the field is yours. Clear it and type any title you'd like "
                "to use for this document:"
            )
            new_title = st.text_input(
                "Document title:",
                value="",
                key="title_input_alt",
                help="Type a title from scratch — anything that describes the document.",
            )
            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "Save this title",
                    type="primary",
                    use_container_width=True,
                    key="alt_yes",
                ):
                    value = (new_title or "").strip()
                    if not value:
                        st.warning("Please enter a title before saving.")
                    else:
                        issue["fix_data"]["confirmed_title"] = value
                        _apply_and_advance()
            with col2:
                if st.button("Skip this issue for now", use_container_width=True, key="alt_skip"):
                    _skip_issue()

        # --- missing_lang: show a language picker dropdown ---
        elif issue["id"] == "missing_lang":
            # Ordered dict of display names → BCP 47 codes for the selectbox
            lang_options = {
                "English (United States)": "en-US",
                "English (United Kingdom)": "en-GB",
                "Spanish":                 "es-ES",
                "French":                  "fr-FR",
                "German":                  "de-DE",
                "Portuguese":              "pt-PT",
                "Italian":                 "it-IT",
                "Chinese (Simplified)":    "zh-CN",
                "Chinese (Traditional)":   "zh-TW",
                "Japanese":                "ja",
                "Korean":                  "ko",
                "Arabic":                  "ar",
                "Russian":                 "ru",
                "Hindi":                   "hi",
                "Dutch":                   "nl-NL",
                "Polish":                  "pl",
                "Turkish":                 "tr",
            }
            current_code = issue.get("fix_data", {}).get("lang_code", "en-US")
            current_display = next(
                (k for k, v in lang_options.items() if v == current_code),
                "English (United States)",
            )
            st.markdown(
                f"The auto-detected language was **{_lang_display_name(current_code)}**. "
                "If that's not right, choose the correct language from the list below:"
            )
            selected_name = st.selectbox(
                "Document language:",
                options=list(lang_options.keys()),
                index=list(lang_options.keys()).index(current_display),
                key="lang_select_alt",
            )
            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "Apply this language",
                    type="primary",
                    use_container_width=True,
                    key="alt_yes",
                ):
                    issue["fix_data"]["lang_code"] = lang_options[selected_name]
                    _apply_and_advance()
            with col2:
                if st.button("Skip this issue for now", use_container_width=True, key="alt_skip"):
                    _skip_issue()

        # --- missing_bookmarks: show the generated bookmark list for review ---
        elif issue["id"] == "missing_bookmarks":
            st.markdown(
                "Here's the bookmark structure I would add, based on the headings I found "
                "in your document. Take a look — if it looks right, go ahead and add them."
            )
            toc = []
            if st.session_state.working_pdf_bytes:
                try:
                    doc_preview = fitz.open(
                        stream=st.session_state.working_pdf_bytes, filetype="pdf"
                    )
                    toc = _generate_toc_from_font_sizes(doc_preview)
                    doc_preview.close()
                except Exception:
                    toc = []

            if toc:
                # Indent by heading level using non-breaking spaces so the
                # hierarchy is visible in the Streamlit markdown output.
                level_indent = {1: "", 2: "    ", 3: "        "}
                lines = [
                    f"{level_indent.get(level, '')}• **{title}** — page {page}"
                    for level, title, page in toc
                ]
                st.markdown("  \n".join(lines))
            else:
                st.warning(
                    "I wasn't able to detect any headings in your document — it may not have "
                    "a clear heading structure, so no bookmarks can be generated automatically. "
                    "You can skip this issue or add headings to your source document and re-upload."
                )

            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "Add these bookmarks",
                    type="primary",
                    use_container_width=True,
                    key="alt_yes",
                    disabled=(not toc),
                ):
                    _apply_and_advance()
            with col2:
                if st.button("Skip this issue for now", use_container_width=True, key="alt_skip"):
                    _skip_issue()

        # --- untagged_pdf: explain the DOCX conversion path, offer proceed or skip ---
        elif issue["id"] == "untagged_pdf":
            st.info(
                "An untagged PDF can't be fully fixed by editing the PDF directly — adding "
                "a proper structure tree requires specialized software like Adobe Acrobat Pro. "
                "Instead, this tool converts your document to a Word file (.docx) with real "
                "heading styles, list formatting, and reading order built in. You can then "
                "re-export that Word file as a properly tagged PDF from Word or Google Docs.\n\n"
                "If you download your file as a Word document at the end of this session, "
                "the structural improvements will be included automatically."
            )
            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "Got it — apply it anyway",
                    type="primary",
                    use_container_width=True,
                    key="alt_yes",
                ):
                    _apply_and_advance()
            with col2:
                if st.button("Skip this issue for now", use_container_width=True, key="alt_skip"):
                    _skip_issue()

        # --- heading_hierarchy: explain the heuristic, offer proceed or skip ---
        elif issue["id"] == "heading_hierarchy":
            fix_data   = issue.get("fix_data", {})
            issue_type = fix_data.get("issue_type", "skip")
            if issue_type == "skip":
                st.info(
                    "This fix uses font sizes to infer which lines are headings and what level "
                    "each should be, then rewrites the heading tags so they follow a logical "
                    "sequence (Heading 1 → Heading 2 → Heading 3) with no skipped levels.\n\n"
                    "This works well for most documents, but it relies on consistent font sizing "
                    "for headings — so if your document uses unusual formatting, the result may "
                    "not be perfect. You can always review the output before finalizing."
                )
            else:
                st.info(
                    "This fix corrects the heading structure so it follows a logical sequence. "
                    "Screen reader users rely on headings to navigate — without a consistent "
                    "hierarchy, they can't tell which sections are major and which are sub-sections."
                )
            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "Got it — apply it anyway",
                    type="primary",
                    use_container_width=True,
                    key="alt_yes",
                ):
                    _apply_and_advance()
            with col2:
                if st.button("Skip this issue for now", use_container_width=True, key="alt_skip"):
                    _skip_issue()

        # --- readability barriers: explain the DOCX font-size fix path ---
        elif issue["id"] in ("readability_barrier_severe", "readability_barrier_moderate"):
            pct = issue.get("fix_data", {}).get("pct_affected", 0)
            st.info(
                f"About {pct}% of your document's text is below 9pt. This fix works "
                "by enforcing a minimum font size when you export to Word (.docx) — "
                "the original PDF is left unchanged, but the Word version will have "
                "all text at a comfortably readable size.\n\n"
                "If you re-export that Word file as a PDF, the corrected sizes carry "
                "through automatically."
            )
            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "Got it — apply it anyway",
                    type="primary",
                    use_container_width=True,
                    key="alt_yes",
                ):
                    _apply_and_advance()
            with col2:
                if st.button("Skip this issue for now", use_container_width=True, key="alt_skip"):
                    _skip_issue()

        # --- reading_order: explain that DOCX conversion corrects reading order ---
        elif issue["id"] == "reading_order":
            st.info(
                "This fix works through the Word document export. When your PDF is "
                "converted to a Word document, content is extracted and re-laid-out in "
                "a visual top-to-bottom sequence — which corrects cases where the PDF's "
                "internal block order doesn't match the expected reading flow.\n\n"
                "Re-exporting that Word file as PDF will produce a document with a "
                "correct, consistent reading order."
            )
            _docx_path_buttons()

        # --- inconsistent_list_tagging: explain DOCX conversion improves list structure ---
        elif issue["id"] == "inconsistent_list_tagging":
            st.info(
                "This fix is handled through the Word document export. When your PDF "
                "is converted to a Word document, lines formatted with bullets, dashes, "
                "or numbers are recognized and converted to properly structured list "
                "elements with correct indentation and tags.\n\n"
                "Re-exporting the Word file as PDF will carry the list structure into "
                "the accessible PDF."
            )
            _docx_path_buttons()

        # --- line_spacing_tight: explain the 1.5× line spacing fix in DOCX ---
        elif issue["id"] == "line_spacing_tight":
            tight_pct = issue.get("fix_data", {}).get("tight_pct", 0)
            st.info(
                f"About {tight_pct}% of your text has line spacing tighter than the "
                "recommended minimum. This fix applies 1.5× line spacing to body text "
                "in the Word document export — the original PDF is left unchanged.\n\n"
                "Re-exporting the Word file as PDF will carry the improved spacing "
                "through automatically."
            )
            _docx_path_buttons()

        # --- letter_spacing_tight: explain the tracking normalization fix ---
        elif issue["id"] == "letter_spacing_tight":
            st.info(
                "This fix normalizes letter spacing in the Word document export. When "
                "compressed character tracking is detected, the Word output uses each "
                "font's standard tracking — removing any artificial compression.\n\n"
                "Note: this check uses a geometric estimate and may occasionally flag "
                "documents with naturally narrow typefaces. Review the output to "
                "confirm the result looks right before finalizing."
            )
            _docx_path_buttons()

        # --- excessive_text_colors: explain DOCX color normalization ---
        elif issue["id"] == "excessive_text_colors":
            color_count = issue.get("fix_data", {}).get("color_count", 0)
            st.info(
                f"This document uses {color_count} distinct text colors. In the Word "
                "document export, text colors will be reduced to a small intentional "
                "set — body text and headings — removing visual noise without changing "
                "the document's content or structure.\n\n"
                "The original PDF is left unchanged. The simplified color palette only "
                "applies to the downloaded Word file."
            )
            _docx_path_buttons()

        # --- fallback for any issue type not explicitly handled above ---
        else:
            st.info(
                "This fix is applied automatically based on what we detected in your document. "
                "If you're not sure about it, you can skip this issue for now and come back later."
            )
            _docx_path_buttons("Apply it anyway")

        _render_abandon_button()
        return

    # =========================================================================
    # STANDARD MODE — initial view before any rejection
    # =========================================================================

    # -------------------------------------------------------------------------
    # Issue-specific fix UI
    #
    # Most issues show a static description of what will be applied.
    # The missing_title issue requires an editable field — the auto-detected
    # candidate may be wrong, and a blank document has no candidate at all.
    # -------------------------------------------------------------------------
    if issue["id"] == "missing_title":
        # Show an editable field so faculty can correct the auto-detected
        # candidate or type a title from scratch.
        fix_data      = issue.get("fix_data", {})
        candidate     = fix_data.get("title_candidate", "")
        st.markdown("**Review and confirm the document title:**")
        st.markdown(issue["fix_preview"])
        confirmed_title = st.text_input(
            "Document title:",
            value=candidate,
            key="title_input",
            help="Edit this if the detected title isn't quite right, or type one from scratch.",
        )

    elif issue["id"] == "heading_hierarchy":
        # Show the detected heading sequence with skips flagged inline so the
        # faculty member can see exactly where the hierarchy breaks down.
        confirmed_title = None
        fix_data    = issue.get("fix_data", {})
        issue_type  = fix_data.get("issue_type", "skip")
        heading_levels = fix_data.get("heading_levels", [])

        st.markdown("**Here's what we found — and how we'll fix it:**")
        st.info(issue["fix_preview"])

        if issue_type == "skip" and heading_levels:
            st.markdown("**Detected heading sequence** (problems marked in bold):")
            st.markdown(_describe_heading_sequence(heading_levels))

    elif issue["id"] == "untagged_pdf":
        # Nothing has been changed yet — this fix is applied at export time
        # when the document is converted to DOCX. Show both the plan and the
        # reason why, so the faculty member can make an informed decision.
        # Second button skips this issue and returns to the issue list.
        st.markdown("**Here's the plan — does this work for you?**")
        st.info(
            "**What we'll do:** I'll convert this document to a structured Word "
            "document (.docx) that preserves your headings, paragraphs, and reading "
            "order. You can then re-save it as a properly tagged PDF from Word or "
            "Google Docs.\n\n"
            "**Why a Word document?** This PDF has no accessibility tags, and adding "
            "them directly requires specialized tools like Adobe Acrobat Pro. Converting "
            "to Word gives your content real heading styles, list formatting, and a "
            "correct reading order built in from the start — and re-saving from Word "
            "produces a properly tagged, accessible PDF."
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "Yes, let's do it.",
                type="primary",
                use_container_width=True,
                key="confirm_yes",
            ):
                _apply_and_advance()
        with col2:
            if st.button(
                "No, not now.",
                use_container_width=True,
                key="confirm_no",
            ):
                _skip_issue()

        _render_abandon_button()
        return

    elif issue["id"] == "missing_lang":
        # The fix writes a /Lang tag to the PDF metadata — nothing has changed yet.
        # Show what will be added, why it matters, and let the faculty member
        # correct the auto-detected language if needed, all on one screen.
        lang_code = issue.get("fix_data", {}).get("lang_code", "en-US")
        lang_options = {
            "English (United States)": "en-US",
            "English (United Kingdom)": "en-GB",
            "Spanish":                 "es-ES",
            "French":                  "fr-FR",
            "German":                  "de-DE",
            "Portuguese":              "pt-PT",
            "Italian":                 "it-IT",
            "Chinese (Simplified)":    "zh-CN",
            "Chinese (Traditional)":   "zh-TW",
            "Japanese":                "ja",
            "Korean":                  "ko",
            "Arabic":                  "ar",
            "Russian":                 "ru",
            "Hindi":                   "hi",
            "Dutch":                   "nl-NL",
            "Polish":                  "pl",
            "Turkish":                 "tr",
        }
        current_display = next(
            (k for k, v in lang_options.items() if v == lang_code),
            "English (United States)",
        )

        st.markdown("**Here's the plan — does this work for you?**")
        st.info(
            f"**What we'll do:** Add the language tag **{_lang_display_name(lang_code)}** "
            f"(`{lang_code}`) to your document's metadata.\n\n"
            "**Why it matters:** Screen readers use the language tag to choose the right "
            "pronunciation rules and speech engine. Without it, text may be read aloud "
            "in the wrong accent or with incorrect pronunciation — especially noticeable "
            "for technical terms, proper names, and any non-English content."
        )

        st.markdown(
            f"I detected **{_lang_display_name(lang_code)}** based on your document's content. "
            "If that's not right, choose the correct language below before confirming:"
        )
        selected_name = st.selectbox(
            "Document language:",
            options=list(lang_options.keys()),
            index=list(lang_options.keys()).index(current_display),
            key="lang_select_standard",
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "Yes, let's do it.",
                type="primary",
                use_container_width=True,
                key="confirm_yes",
            ):
                issue["fix_data"]["lang_code"] = lang_options[selected_name]
                _apply_and_advance()
        with col2:
            if st.button(
                "No, not now.",
                use_container_width=True,
                key="confirm_no",
            ):
                _skip_issue()

        _render_abandon_button()
        return

    elif issue["id"] == "color_only_cue":
        # We detected colored text but can't verify or fix the accessibility
        # implications automatically. Explain what to look for, offer a DOCX
        # download so the faculty member can fix the colors in Word themselves,
        # then let them acknowledge or skip.
        color_count = issue.get("fix_data", {}).get("color_count", 0)
        file_stem = (st.session_state.file_name or "document").rsplit(".", 1)[0]

        st.markdown("**We noticed colored text in your document.**")
        st.markdown(
            f"This document uses **{color_count} distinct text colors**. "
            "Color can create accessibility barriers in a few ways — most of which "
            "require a quick manual check rather than an automated fix. "
            "Here's what to look for:"
        )
        st.markdown(
            "**1. Color as the only cue for meaning** (WCAG 2.2 SC 1.4.1)  \n"
            "If color is the only thing distinguishing categories, required items, "
            "or warnings — with no label, symbol, or text alongside it — students "
            "who are color-blind or printing in black and white will miss that "
            "meaning entirely.  \n"
            "*Fix:* Add a text label, asterisk, or symbol next to anything communicated "
            "by color alone. For example: \"Required \\*\" instead of red text alone.\n\n"
            "**2. Text contrast** (WCAG 2.2 SC 1.4.3)  \n"
            "Colored text may not meet the minimum contrast ratio against the page "
            "background — especially light colors on white.  \n"
            "*Fix:* Check that colored text meets at least 4.5:1 contrast for body "
            "text, or 3:1 for large text (18pt+). Tools like the WebAIM Contrast "
            "Checker can help."
        )

        st.divider()
        st.markdown(
            "**To make these fixes yourself:** Download your document as a Word file, "
            "open it in Microsoft Word or Google Docs, and make the color adjustments "
            "there. The Word file already has your headings, paragraphs, and structure "
            "preserved — it's ready to edit."
        )

        # Build or reuse the cached DOCX so the download is ready immediately.
        if not st.session_state.get("edited_docx_bytes"):
            with st.spinner("Preparing your Word document..."):
                pdf_bytes = st.session_state.working_pdf_bytes or st.session_state.file_bytes
                st.session_state.edited_docx_bytes = build_docx_from_pdf(pdf_bytes)
        docx_bytes = st.session_state.edited_docx_bytes

        if docx_bytes:
            st.download_button(
                label="⬇ Download as Word document (.docx)",
                data=docx_bytes,
                file_name=f"{file_stem}_edited.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="color_docx_download",
            )

        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "I've handled it — move on.",
                type="primary",
                use_container_width=True,
                key="confirm_yes",
            ):
                _apply_and_advance()
        with col2:
            if st.button(
                "No, not now.",
                use_container_width=True,
                key="confirm_no",
            ):
                _skip_issue()

        _render_abandon_button()
        return

    else:
        st.markdown("**Here's what I changed. Does this look right to you?**")
        st.info(issue["fix_preview"])
        confirmed_title = None  # not used for other issue types

    # Faculty confirmation — exactly two choices per CLAUDE.md
    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Yes, that looks good.",
            type="primary",
            use_container_width=True,
            key="confirm_yes",
        ):
            # For missing_title, validate and store the user's edited value
            # before calling the fix function. Guard against an empty field.
            proceed = True
            if issue["id"] == "missing_title":
                value = (confirmed_title or "").strip()
                if not value:
                    st.warning("Please enter a title before saving.")
                    proceed = False
                else:
                    # Write the confirmed title into fix_data so
                    # apply_fix_missing_title() uses it over the raw candidate.
                    issue["fix_data"]["confirmed_title"] = value

            if proceed:
                _apply_and_advance()

    with col2:
        if st.button(
            "No, let's try again.",
            use_container_width=True,
            key="confirm_no",
        ):
            # Switch to the alternative mode — each issue type has its own
            # custom second path instead of just returning to the issue list.
            st.session_state.showing_alternative = True
            st.rerun()

    _render_abandon_button()


# -----------------------------------------------------------------------------
# STEP: continue_or_stop
# After each resolved issue, ask whether to keep going or stop.
# -----------------------------------------------------------------------------
def render_continue_or_stop():
    """
    Purpose: Give the user a natural pause after each resolved issue.
    Per CLAUDE.md: ask "keep going?" after every issue.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    resolved = len(st.session_state.resolved_ids)
    remaining = len(st.session_state.issues) - resolved

    st.success(f"Issue resolved! You've fixed {resolved} issue(s) so far.")

    if remaining > 0:
        st.markdown(
            f"There {'is' if remaining == 1 else 'are'} still "
            f"**{remaining}** issue{'s' if remaining != 1 else ''} remaining. "
            "Would you like to keep going?"
        )

        col1, col2 = st.columns(2)
        with col1:
            if st.button(
                "Yes, let's keep going.",
                type="primary",
                use_container_width=True,
            ):
                st.session_state.step = "issue_list"
                st.rerun()
        with col2:
            if st.button(
                "No, I'm done for now.",
                use_container_width=True,
            ):
                st.session_state.step = "manual_review"
                st.rerun()
    else:
        # No issues remaining — go to manual review checklist before download
        st.markdown("You've worked through all the issues — great job!")
        if st.button("Continue to download", type="primary"):
            st.session_state.step = "manual_review"
            st.rerun()

    _render_abandon_button()


# -----------------------------------------------------------------------------
# STEP: re_analysis
# Offered after every 3 resolved issues per CLAUDE.md.
# -----------------------------------------------------------------------------
def render_re_analysis():
    """
    Purpose: Offer to re-analyze the document after every 3 resolved issues.
    Per CLAUDE.md: verbatim message and exactly two choices.
    Re-analysis calls analyze_pdf() on the current working copy so resolved
    issues are removed and any newly introduced issues are surfaced.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    # Verbatim message from CLAUDE.md
    st.info(
        "You've made some good progress — fixing some of these issues can sometimes "
        "resolve others automatically. Would you like me to take another look at the "
        "document to see if anything else has been taken care of?"
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Yes, check it again.",
            type="primary",
            use_container_width=True,
        ):
            # Re-run analysis on the working copy so any issues resolved by
            # earlier fixes are automatically removed from the list.
            with st.spinner("Re-analyzing your document..."):
                reanalysis_bytes = (
                    st.session_state.working_pdf_bytes or st.session_state.file_bytes
                )
                result = analyze_pdf(reanalysis_bytes)
                if result["status"] == "issues_found":
                    # Keep only issues that haven't already been resolved.
                    fresh_issues = [
                        i for i in result["issues"]
                        if i["id"] not in st.session_state.resolved_ids
                    ]
                    st.session_state.issues = (
                        fresh_issues
                        + [
                            i for i in st.session_state.issues
                            if i["id"] in st.session_state.resolved_ids
                        ]
                    )
            st.session_state.reanalysis_done = True
            st.session_state.step = "issue_list"
            st.rerun()
    with col2:
        if st.button(
            "No, let's keep going.",
            use_container_width=True,
        ):
            st.session_state.reanalysis_done = True
            st.session_state.step = "continue_or_stop"
            st.rerun()

    _render_abandon_button()


# -----------------------------------------------------------------------------
# STEP: manual_review
# A checklist of accessibility issues that can't be detected automatically.
# Shown after all auto-detected issues are resolved, before the download screen.
# Faculty can check each item, mark things that need attention, and read
# plain-language guidance — or skip the whole screen if they choose.
# -----------------------------------------------------------------------------
def render_manual_review():
    """
    Purpose: Surface accessibility issues that require human judgment.
    These are real WCAG 2.2 / Title II problems the app cannot detect
    algorithmically. Faculty review each item and mark it OK or needing attention.
    """
    # Checklist items — defined here since they are only used by this function.
    # Each item has: id, severity, title, a "what to check" prompt, and guidance
    # shown when the faculty member flags something as needing attention.
    ITEMS = [
        {
            "id":       "charts_graphs",
            "severity": "red",
            "title":    "Charts, graphs, and diagrams",
            "check":    "Does the alt text for each chart describe the data and key takeaway — not just the visual appearance? A description should say what the data shows and what conclusion to draw, not just 'bar chart.'",
            "guidance": "Revisit the alt text for each chart or graph. Describe the main finding or trend — your students who can't see the image should understand the point it makes, not just that a chart exists.",
        },
        {
            "id":       "video_audio",
            "severity": "red",
            "title":    "Embedded video or audio",
            "check":    "Does your document include any embedded video or audio clips? If so, do they have captions or a written transcript?",
            "guidance": "Add a written transcript or caption file to any audio or video content before sharing. Students who are deaf or hard of hearing cannot access that content without one.",
        },
        {
            "id":       "math_equations",
            "severity": "red",
            "title":    "Math and equations",
            "check":    "Are equations embedded as images, or typed as text? Even typed math notation (x², ∑, ∫) may not be read correctly by all screen readers.",
            "guidance": "For documents with significant math content, consider providing a written description of what each equation represents, or use MathType/LaTeX formatting. Image-based equations need alt text that explains their meaning, not just their appearance.",
        },
        {
            "id":       "color_only_cue",
            "severity": "yellow",
            "title":    "Color used as the only cue for meaning",
            "check":    "Does your document use color to signal something important — like red text for required items, or color-coded categories — without any other indicator like a label, symbol, or text?",
            "guidance": "Add a text label, symbol, or pattern alongside any color-coded elements. For example: 'Required *' instead of red text alone. Students who are color blind or printing in black and white will miss color-only signals entirely.",
        },
        {
            "id":       "link_text",
            "severity": "yellow",
            "title":    "Non-descriptive link text",
            "check":    "Do any links say 'click here,' 'this link,' or 'read more'? Screen readers read link text in isolation, without the surrounding sentence for context.",
            "guidance": "Edit link text to describe the destination. Instead of 'click here for the syllabus,' write 'Fall 2025 Course Syllabus.' This helps screen reader users understand where each link goes.",
        },
        {
            "id":       "table_structure",
            "severity": "yellow",
            "title":    "Table structure and headers",
            "check":    "Do your tables have proper header rows — actually marked as headers in the document, not just visually bold? Does each table still make sense when read row by row?",
            "guidance": "In your source document, mark header rows using the built-in table header option — not bold formatting. Re-export as PDF to carry the structure through. Screen readers read tables cell by cell and rely on headers to give each cell meaning.",
        },
        {
            "id":       "abbreviations",
            "severity": "green",
            "title":    "Abbreviations and acronyms",
            "check":    "Are discipline-specific abbreviations and acronyms spelled out the first time they appear?",
            "guidance": "Add a brief definition on first use (e.g., 'CUNY — City University of New York'), or include a glossary at the end of the document. This helps all readers, not just those using assistive technology.",
        },
        {
            "id":       "non_english",
            "severity": "green",
            "title":    "Non-English passages",
            "check":    "Does your document include any quotes or passages in a different language? Those sections may be read aloud with incorrect pronunciation by screen readers.",
            "guidance": "In Word: select the non-English text → Review → Language → Set Proofing Language. This tells screen readers to switch to the correct speech engine for that passage.",
        },
        {
            "id":       "decorative_images",
            "severity": "green",
            "title":    "Decorative image verification",
            "check":    "If this document had images, review any the app auto-classified as decorative or removed as scan artifacts. Was anything accidentally excluded that students need to see?",
            "guidance": "Compare the images in your original file with the output. If a content image was removed or incorrectly marked decorative, you can re-upload and work through the image workflow again.",
        },
    ]

    SEVERITY_LABEL = {
        "red":    "🔴 Red Light",
        "yellow": "🟡 Yellow Light",
        "green":  "🟢 Green Light",
    }

    st.title("♿ PDF Assistant")
    st.divider()

    st.markdown("### One more step — a quick manual check")
    st.markdown(
        "There are a few things I can't verify automatically. This should only take "
        "a minute — just a quick look at your document for each item below."
    )

    # Ensure results dict exists in session state
    if "manual_review_results" not in st.session_state:
        st.session_state.manual_review_results = {}

    results = st.session_state.manual_review_results

    for item in ITEMS:
        item_id  = item["id"]
        result   = results.get(item_id)

        st.divider()

        # Header: severity + title + status icon
        status = " ✓" if result == "ok" else " ⚠️" if result == "attention" else ""
        st.markdown(
            f"**{SEVERITY_LABEL[item['severity']]} — {item['title']}**{status}"
        )
        st.caption(item["check"])

        if result is None:
            # Not yet reviewed — show the two response buttons
            col1, col2 = st.columns(2)
            with col1:
                if st.button(
                    "Looks fine ✓",
                    key=f"mr_ok_{item_id}",
                    use_container_width=True,
                ):
                    st.session_state.manual_review_results[item_id] = "ok"
                    st.rerun()
            with col2:
                if st.button(
                    "Needs attention ⚠️",
                    key=f"mr_att_{item_id}",
                    use_container_width=True,
                ):
                    st.session_state.manual_review_results[item_id] = "attention"
                    st.rerun()

        elif result == "attention":
            # Show guidance and let them mark as done
            st.info(item["guidance"])
            if st.button(
                "Got it — mark as reviewed",
                key=f"mr_done_{item_id}",
            ):
                st.session_state.manual_review_results[item_id] = "ok"
                st.rerun()

        # result == "ok": checkmark already shown in header, nothing else needed

    st.divider()

    reviewed = len(results)
    total    = len(ITEMS)
    if reviewed < total:
        st.caption(f"{reviewed} of {total} items reviewed")

    col1, col2 = st.columns(2)
    with col1:
        if st.button(
            "Continue to download →",
            type="primary",
            use_container_width=True,
            key="mr_continue",
        ):
            st.session_state.step = "choose_format"
            st.rerun()
    with col2:
        if st.button(
            "Skip this step",
            use_container_width=True,
            key="mr_skip",
        ):
            st.session_state.step = "choose_format"
            st.rerun()

    _render_go_back("issue_list")


# -----------------------------------------------------------------------------
# STEP: choose_format
# Ask whether the user wants DOCX or PDF output.
# Per CLAUDE.md: deliver the file with "_edited" appended to the original name.
# -----------------------------------------------------------------------------
def render_choose_format():
    """
    Purpose: Let the user choose their output format (DOCX or PDF).
    Per CLAUDE.md: file name gets "_edited" appended before the extension.
    Builds the chosen format from the current working copy and serves it
    as a download via st.download_button.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    resolved = len(st.session_state.resolved_ids)
    st.success(f"Great work — you've resolved {resolved} accessibility issue(s).")

    st.markdown(
        "What format would you like for your updated file? "
        "I'll save it with your original file name, with **_edited** added at the end."
    )

    # Build the output file names for display
    base_name = st.session_state.file_name.replace(".pdf", "").replace(".PDF", "")
    docx_name = f"{base_name}_edited.docx"
    pdf_name = f"{base_name}_edited.pdf"

    # Use the working copy (which has all confirmed fixes applied).
    # Fall back to the original if working_pdf_bytes was never set.
    output_pdf_bytes = st.session_state.working_pdf_bytes or st.session_state.file_bytes

    # Build the DOCX once and cache it in session state.
    # build_docx_from_pdf() runs Docling over the entire document — it's expensive.
    # Streamlit reruns this whole function on every user interaction (e.g. clicking
    # the PDF download button), so without caching the DOCX would be rebuilt every
    # time even when the user never asked for it.
    if not st.session_state.edited_docx_bytes:
        with st.spinner("Building your Word document..."):
            st.session_state.edited_docx_bytes = build_docx_from_pdf(output_pdf_bytes)
    docx_bytes = st.session_state.edited_docx_bytes

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Word processor document (DOCX)**")
        st.caption(
            "Opens in Microsoft Word, Google Docs, or any compatible application. "
            "Best if you need to edit the content further before re-publishing, or "
            "if you want to re-export as a properly tagged PDF from Word."
        )
        st.download_button(
            label=f"Download {docx_name}",
            data=docx_bytes,
            file_name=docx_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )

    with col2:
        st.markdown("**Updated PDF document**")
        st.caption(
            "Your PDF with all confirmed fixes applied — metadata, language tag, "
            "and bookmarks are saved. Ready to share or upload directly."
        )
        st.download_button(
            label=f"Download {pdf_name}",
            data=output_pdf_bytes,
            file_name=pdf_name,
            mime="application/pdf",
            use_container_width=True,
        )

    _render_abandon_button()


# -----------------------------------------------------------------------------
# STEP: done
# Session end screen. Shown when the user exits without completing all issues.
# -----------------------------------------------------------------------------
def render_done():
    """
    Purpose: Gracefully end the session.
    Shown when user chooses to walk away (password, scanned doc) or finishes.
    """
    st.title("♿ PDF Assistant")
    st.divider()

    reason = st.session_state.get("done_reason", "")

    if reason == "password_exit":
        st.info(
            "No problem — whenever you have an unlocked version of the file, "
            "come back and I'll help you check it for accessibility issues."
        )
    elif reason == "scanned_exit":
        st.info(
            "Understood. If you change your mind and would like to convert this file "
            "to readable text, just come back and upload it again."
        )
    elif reason == "ocr_exit":
        st.success("Your converted file is ready whenever you need it.")
        st.markdown(
            "Whenever you're ready to work through the accessibility issues, "
            "come back and upload your converted file — I'll pick up right where we left off."
        )
    elif reason == "clean_exit":
        st.success("Your document is in great shape — no action needed.")
    else:
        resolved = len(st.session_state.resolved_ids)
        st.success(
            f"You've resolved {resolved} accessibility issue(s) — "
            "good progress for your students!"
        )
        st.markdown(
            "Whenever you're ready to continue, come back and upload your file again. "
            "I'll pick up right where we left off."
        )

    st.divider()
    if st.button("Start a new session", type="primary"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        init_state()
        st.rerun()


# =============================================================================
# MAIN ROUTER
# =============================================================================
# This is the traffic director. It reads the current "step" from session_state
# and calls the matching render function. Streamlit re-runs this entire block
# on every user interaction, so the right screen always appears.
# =============================================================================

STEP_HANDLERS = {
    "upload":               render_upload,
    "analyzing":            render_analyzing,
    "password_protected":   render_password_protected,
    "password_walkthrough": render_password_walkthrough,
    "scanned_doc":          render_scanned_doc,
    "scanned_goodbye":      render_scanned_goodbye,
    "running_ocr":          render_running_ocr,
    "running_easyocr":      render_running_easyocr,
    "ocr_format_select":    render_ocr_format_select,
    "no_issues":            render_no_issues,
    "issue_list":           render_issue_list,
    "image_alt_text":       render_image_alt_text,
    "source_file_question": render_source_file_question,
    "cuny_resources":       render_cuny_resources,
    "resolving_issue":      render_resolving_issue,
    "continue_or_stop":     render_continue_or_stop,
    "re_analysis":          render_re_analysis,
    "manual_review":        render_manual_review,
    "choose_format":        render_choose_format,
    "done":                 render_done,
}

# Call the render function for the current step.
# If somehow step is set to an unknown value, fall back to upload.
handler = STEP_HANDLERS.get(st.session_state.step, render_upload)
handler()


# -----------------------------------------------------------------------------
# FOOTER
# Always visible at the bottom of every screen.
# Provides legal context without being the lead message.
# -----------------------------------------------------------------------------
st.divider()
st.caption(
    "PDF Assistant supports compliance with **Title II of the ADA** (28 CFR Part 35) "
    "and **WCAG 2.2 Level AA**. It provides guidance and analysis but does not "
    "constitute legal advice. Always consult your institution's accessibility office."
)
