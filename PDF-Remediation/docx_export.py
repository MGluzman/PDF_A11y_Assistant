"""DOCX and PDF output building."""

# -------------------------------------------------------------------------------
# docx_export.py — builds Word (DOCX) and converts to accessible PDF
#
# Functions:
#   _embed_image_with_alt_text()  — add image to DOCX with wp:docPr descr attribute
#   build_docx_from_pdf()         — extract PDF content → structured DOCX
#   _convert_docx_to_pdf()        — convert DOCX → tagged PDF via Word COM
# -------------------------------------------------------------------------------

import io
import os
import tempfile
from collections import Counter

import fitz
import streamlit as st
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.oxml.ns import qn
from docling_core.types.doc import DocItemLabel
import pythoncom
from docx2pdf import convert as docx2pdf_convert

from images import extract_image_by_source


def _embed_image_with_alt_text(doc_out, img_bytes, alt_text):
    """
    Add an image to the DOCX document with proper alt text set on the
    underlying OOXML element.

    Word stores alt text in the 'descr' attribute of the <wp:docPr> element
    inside the drawing's inline block:
        w:r > w:drawing > wp:inline > wp:docPr[descr="..."]

    Setting this attribute is what populates Word's "Alt Text" panel and
    what gets written into the PDF structure tree when the faculty member
    uses "Export to PDF" in Word or Google Docs.

    Parameters:
        doc_out   (DocxDocument): The document being built.
        img_bytes (bytes):        Raw image bytes (PNG or JPEG).
        alt_text  (str):          Alt text description. Pass "" for decorative
                                  images — screen readers will skip them.

    Returns:
        bool: True if embedding succeeded, False if it failed (caller can fall
              back to a text placeholder).
    """
    try:
        para = doc_out.add_paragraph()
        run  = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(5))

        # Walk the OOXML tree to reach wp:docPr and set its 'descr' attribute.
        # The tree is: <w:r> → <w:drawing> → <wp:inline> → <wp:docPr descr="">
        drawing = run._element.find(qn("w:drawing"))
        if drawing is not None:
            inline = drawing.find(qn("wp:inline"))
            if inline is not None:
                doc_pr = inline.find(qn("wp:docPr"))
                if doc_pr is not None:
                    doc_pr.set("descr", alt_text)
        return True
    except Exception:
        return False


def build_docx_from_pdf(pdf_bytes):
    """
    Extract text and structure from a PDF and produce a structured DOCX.

    Uses Docling as the primary path when a cached DoclingDocument is available
    in session state. Docling's ML-based element labels (TITLE, SECTION_HEADER,
    LIST_ITEM, TABLE, PICTURE) map cleanly to DOCX styles and give much better
    heading and list detection than the font-size heuristic alone.

    Falls back to the PyMuPDF font-size heuristic if Docling is unavailable
    (e.g., the PDF was an OCR output generated mid-session, or Docling failed).

    Parameters:
        pdf_bytes (bytes): The working copy of the PDF (with metadata fixes applied).

    Returns:
        bytes: A DOCX file as raw bytes, ready for st.download_button.
    """
    doc_out = DocxDocument()

    # ------------------------------------------------------------------
    # PRIMARY PATH: Use the cached Docling document from session state.
    # The DoclingDocument was produced by _run_docling() during analyze_pdf()
    # and stored in st.session_state.docling_doc by render_analyzing().
    # ------------------------------------------------------------------
    docling_doc = st.session_state.get("docling_doc")

    # ------------------------------------------------------------------
    # Shared setup for both paths: read alt text results once so both
    # the Docling path and the PyMuPDF fallback can use the same data.
    # ------------------------------------------------------------------
    alt_results   = st.session_state.get("alt_text_results", {})
    alt_images    = st.session_state.get("alt_text_images", [])

    REMOVED_TYPES      = {"artifact", "background", "edge_artifact"}
    INFORMATIONAL_TYPES = {"informational"}

    if docling_doc is not None:
        try:
            # Group img_keys by page (1-based page numbers from fix_data).
            keys_by_page: dict[int, set] = {}
            for img_meta in alt_images:
                k = img_meta.get("img_key", img_meta.get("xref"))
                keys_by_page.setdefault(img_meta["page"], set()).add(k)

            fully_excluded_pages = {
                pg for pg, keys in keys_by_page.items()
                if all(alt_results.get(k, {}).get("type") in REMOVED_TYPES for k in keys)
            }

            images_by_page: dict[int, list] = {}
            for img_meta in alt_images:
                pg = img_meta.get("page")
                if pg is not None:
                    images_by_page.setdefault(pg, []).append(img_meta)

            for item, _ in docling_doc.iterate_items():
                if not hasattr(item, "label"):
                    continue

                label = item.label
                text = getattr(item, "text", "").strip()

                if label == DocItemLabel.TITLE:
                    doc_out.add_heading(text, level=1)

                elif label == DocItemLabel.SECTION_HEADER:
                    doc_out.add_heading(text, level=2)

                elif label == DocItemLabel.LIST_ITEM:
                    if text:
                        doc_out.add_paragraph(text, style="List Bullet")

                elif label == DocItemLabel.TABLE:
                    doc_out.add_paragraph("[Table — see original PDF for layout]")

                elif label == DocItemLabel.PICTURE:
                    item_page = item.prov[0].page_no if item.prov else None
                    if item_page in fully_excluded_pages:
                        continue

                    page_num = item_page or "?"

                    page_imgs = images_by_page.get(item_page, [])
                    img_meta  = page_imgs.pop(0) if page_imgs else None

                    if img_meta is None:
                        doc_out.add_paragraph(f"[Image — page {page_num} of original PDF]")
                        continue

                    img_key  = img_meta.get("img_key")
                    entry    = alt_results.get(img_key, {})
                    img_type = entry.get("type", "")

                    if img_type in REMOVED_TYPES:
                        continue

                    img_bytes_raw, _ = extract_image_by_source(pdf_bytes, img_meta)

                    if img_bytes_raw:
                        if img_type in INFORMATIONAL_TYPES:
                            alt_str = entry.get("alt_text", "").strip()
                        else:
                            alt_str = ""

                        if not _embed_image_with_alt_text(doc_out, img_bytes_raw, alt_str):
                            if alt_str:
                                doc_out.add_paragraph(f"[Image, page {page_num}: {alt_str}]")
                            else:
                                doc_out.add_paragraph(f"[Image — page {page_num} of original PDF]")
                    else:
                        alt_str = entry.get("alt_text", "").strip() if img_type in INFORMATIONAL_TYPES else ""
                        if alt_str:
                            doc_out.add_paragraph(f"[Image, page {page_num}: {alt_str}]")
                        else:
                            doc_out.add_paragraph(f"[Image — page {page_num} of original PDF]")

                elif label in (DocItemLabel.TEXT, DocItemLabel.PARAGRAPH):
                    if text:
                        doc_out.add_paragraph(text)

                else:
                    if text:
                        doc_out.add_paragraph(text)

            out = io.BytesIO()
            doc_out.save(out)
            return out.getvalue()

        except Exception:
            doc_out = DocxDocument()

    # ------------------------------------------------------------------
    # FALLBACK PATH: Font-size heuristic using PyMuPDF.
    # ------------------------------------------------------------------
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        xref_to_entry = {k: entry for k, entry in alt_results.items()}
        xref_to_meta  = {
            m.get("img_key", m.get("xref")): m
            for m in alt_images
        }

        all_sizes = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span.get("text", "").strip():
                            all_sizes.append(span.get("size", 12))
        body_size = Counter(round(s) for s in all_sizes).most_common(1)[0][0] if all_sizes else 12

        for page_num, page in enumerate(doc):
            if page_num > 0:
                doc_out.add_page_break()

            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue

                for line in block.get("lines", []):
                    line_text = " ".join(
                        span.get("text", "") for span in line.get("spans", [])
                    ).strip()
                    if not line_text:
                        continue

                    line_sizes = [span.get("size", body_size) for span in line.get("spans", [])]
                    avg_size = sum(line_sizes) / len(line_sizes) if line_sizes else body_size

                    if avg_size > body_size * 1.4:
                        doc_out.add_heading(line_text, level=1)
                    elif avg_size > body_size * 1.2:
                        doc_out.add_heading(line_text, level=2)
                    else:
                        doc_out.add_paragraph(line_text)

            for img in page.get_images(full=True):
                xref     = img[0]
                entry    = xref_to_entry.get(xref, {})
                img_type = entry.get("type", "")

                if img_type in REMOVED_TYPES:
                    continue

                img_meta = xref_to_meta.get(xref)
                img_bytes_raw = None
                if img_meta:
                    img_bytes_raw, _ = extract_image_by_source(pdf_bytes, img_meta)

                if img_bytes_raw:
                    if img_type in INFORMATIONAL_TYPES:
                        alt_str = entry.get("alt_text", "").strip()
                    else:
                        alt_str = ""

                    if not _embed_image_with_alt_text(doc_out, img_bytes_raw, alt_str):
                        if alt_str:
                            doc_out.add_paragraph(
                                f"[Image, page {page_num + 1}: {alt_str}]"
                            )
                else:
                    alt_str = entry.get("alt_text", "").strip() if img_type in INFORMATIONAL_TYPES else ""
                    if alt_str:
                        doc_out.add_paragraph(
                            f"[Image, page {page_num + 1}: {alt_str}]"
                        )
                    elif img_type not in ("decorative",):
                        doc_out.add_paragraph(
                            f"[Image — page {page_num + 1} of original PDF]"
                        )

        doc.close()

    except Exception:
        doc_out.add_paragraph(
            "The content of this document could not be extracted automatically. "
            "Please open the original PDF and copy the content manually."
        )

    out = io.BytesIO()
    doc_out.save(out)
    return out.getvalue()


def _convert_docx_to_pdf(docx_bytes):
    """
    Write DOCX bytes to a temp file, call docx2pdf to convert via Word,
    read the resulting PDF back as bytes, and clean up.

    Parameters:
        docx_bytes (bytes): The fully built DOCX file.

    Returns:
        bytes: The converted PDF file.

    Raises:
        Exception: Propagates any error from docx2pdf so the caller can
                   show a helpful message rather than crashing silently.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "edited.docx")
        pdf_path  = os.path.join(tmpdir, "edited.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # docx2pdf drives Word via COM on Windows. Word must be installed.
        # Streamlit callbacks run in non-main threads that have never called
        # CoInitialize, which causes COM to fail with -2147221008. We initialize
        # COM explicitly on this thread and release it when done.
        pythoncom.CoInitialize()
        try:
            docx2pdf_convert(docx_path, pdf_path)
        finally:
            pythoncom.CoUninitialize()

        with open(pdf_path, "rb") as f:
            return f.read()
