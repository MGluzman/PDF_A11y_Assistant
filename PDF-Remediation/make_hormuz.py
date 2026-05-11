import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import tempfile, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

tmp = tempfile.mkdtemp()

# ---------------------------------------------------------------------------
# CHART — trade share by industry
# ---------------------------------------------------------------------------
industries = [
    'Crude Oil',
    'Liquefied Natural Gas (LNG)',
    'Refined Petroleum Products',
    'Petrochemicals & Fertilizers',
    'Other Cargo',
]
pct_global = [30, 20, 10, 8, 2]
colors = ['#1a3a5c', '#2e6da4', '#4a9fd4', '#6dbfb8', '#a8d5ba']

fig, ax = plt.subplots(figsize=(8, 4))
bars = ax.barh(industries, pct_global, color=colors, edgecolor='white', height=0.55)
for bar, val in zip(bars, pct_global):
    ax.text(val + 0.4, bar.get_y() + bar.get_height() / 2,
            f'{val}%', va='center', ha='left', fontsize=10, fontweight='bold', color='#1a3a5c')
ax.set_xlabel('Share of Global Seaborne Trade (%)', fontsize=10, color='#333333')
ax.set_title('Annual Global Trade Through the Strait of Hormuz by Industry',
             fontsize=11, fontweight='bold', color='#1a3a5c', pad=12)
ax.set_xlim(0, 38)
ax.tick_params(axis='y', labelsize=9)
ax.tick_params(axis='x', labelsize=9)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
ax.set_facecolor('#f8f9fa')
fig.patch.set_facecolor('white')
ax.text(0.99, -0.13,
        'Sources: U.S. Energy Information Administration; International Energy Agency (2024 estimates)',
        transform=ax.transAxes, ha='right', fontsize=7, color='#888888', style='italic')
plt.tight_layout()
chart_path = os.path.join(tmp, 'chart.png')
plt.savefig(chart_path, dpi=150, bbox_inches='tight')
plt.close()
print('Chart saved')

# ---------------------------------------------------------------------------
# MAP — schematic regional diagram
# ---------------------------------------------------------------------------
fig2, ax2 = plt.subplots(figsize=(8, 5))
ax2.set_facecolor('#a8d0e6')
fig2.patch.set_facecolor('white')

iran = plt.Polygon([
    (55.0, 27.5), (56.5, 27.2), (57.5, 26.8), (58.5, 26.5), (60.0, 26.3),
    (60.0, 30.0), (55.0, 30.0)
], closed=True, facecolor='#d4c99a', edgecolor='#7a6a40', linewidth=1)

uae = plt.Polygon([
    (54.0, 23.5), (56.0, 23.5), (56.5, 24.5), (55.5, 25.2), (54.5, 25.5), (54.0, 25.0)
], closed=True, facecolor='#d4c99a', edgecolor='#7a6a40', linewidth=1)

oman = plt.Polygon([
    (56.2, 23.5), (57.5, 23.3), (58.8, 22.5), (59.5, 22.0), (60.0, 22.5),
    (60.0, 24.0), (58.5, 24.5), (57.5, 24.8), (56.5, 25.5), (56.0, 25.0), (56.2, 24.0)
], closed=True, facecolor='#d4c99a', edgecolor='#7a6a40', linewidth=1)

saudi = plt.Polygon([
    (50.0, 24.0), (54.0, 23.5), (54.0, 25.0), (52.0, 26.5), (50.0, 26.5)
], closed=True, facecolor='#d4c99a', edgecolor='#7a6a40', linewidth=1)

iraq = plt.Polygon([
    (47.5, 29.5), (50.0, 29.0), (50.0, 30.5), (47.5, 30.5)
], closed=True, facecolor='#d4c99a', edgecolor='#7a6a40', linewidth=1)

for patch in [iran, uae, oman, saudi, iraq]:
    ax2.add_patch(patch)

# Shipping lane arrows through the strait
ax2.annotate('', xy=(58.5, 25.9), xytext=(56.8, 26.5),
             arrowprops=dict(arrowstyle='->', color='#1a3a5c', lw=2.0))
ax2.annotate('', xy=(56.8, 25.3), xytext=(58.5, 24.7),
             arrowprops=dict(arrowstyle='->', color='#2e6da4', lw=2.0))

# Country labels
ax2.text(57.5, 28.5, 'IRAN', fontsize=10, fontweight='bold', color='#4a3a10', ha='center')
ax2.text(54.8, 24.1, 'UAE', fontsize=8, fontweight='bold', color='#4a3a10', ha='center')
ax2.text(58.5, 23.2, 'OMAN', fontsize=9, fontweight='bold', color='#4a3a10', ha='center')
ax2.text(51.5, 25.2, 'Saudi\nArabia', fontsize=8, fontweight='bold', color='#4a3a10', ha='center')
ax2.text(48.8, 29.8, 'Iraq/Kuwait', fontsize=7, fontweight='bold', color='#4a3a10', ha='center')

# Body of water labels
ax2.text(51.0, 27.5, 'Persian Gulf', fontsize=10, color='#1a4a6c', ha='center', style='italic')
ax2.text(59.5, 24.2, 'Gulf of\nOman', fontsize=9, color='#1a4a6c', ha='center', style='italic')

# Strait callout
ax2.annotate('Strait of\nHormuz', xy=(57.2, 26.0), xytext=(55.2, 27.8),
             fontsize=9, fontweight='bold', color='#8b0000',
             arrowprops=dict(arrowstyle='->', color='#8b0000', lw=1.5),
             bbox=dict(boxstyle='round,pad=0.3', facecolor='#fff3f3', edgecolor='#8b0000', alpha=0.9))

# North arrow
ax2.text(60.5, 29.6, 'N', fontsize=12, fontweight='bold', color='#333333', ha='center')
ax2.annotate('', xy=(60.5, 29.9), xytext=(60.5, 29.4),
             arrowprops=dict(arrowstyle='->', color='#333333', lw=2))

ax2.set_xlim(47, 61.5)
ax2.set_ylim(21.5, 31)
ax2.set_xlabel('Longitude (degrees E)', fontsize=8, color='#666')
ax2.set_ylabel('Latitude (degrees N)', fontsize=8, color='#666')
ax2.set_title('The Strait of Hormuz and Surrounding Region',
              fontsize=11, fontweight='bold', color='#1a3a5c', pad=10)
ax2.tick_params(labelsize=8)
ax2.spines['top'].set_visible(False)
ax2.spines['right'].set_visible(False)
ax2.text(0.01, -0.10, 'Schematic map -- not to precise geographic scale',
         transform=ax2.transAxes, fontsize=7, color='#999', style='italic')
plt.tight_layout()
map_path = os.path.join(tmp, 'map.png')
plt.savefig(map_path, dpi=150, bbox_inches='tight')
plt.close()
print('Map saved')

# ---------------------------------------------------------------------------
# WORD DOCUMENT
# ---------------------------------------------------------------------------
doc = Document()

# Title
title = doc.add_heading('Trade Through the Strait of Hormuz', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('History, Strategic Importance, and Global Economic Significance')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sub.runs[0].font.italic = True
doc.add_paragraph()

# --- Section 1: Geography ---
doc.add_heading('Geography and Strategic Position', level=1)
doc.add_paragraph(
    'The Strait of Hormuz is a narrow waterway connecting the Persian Gulf to the Gulf of Oman '
    'and, beyond it, the Arabian Sea. Stretching approximately 104 miles in length and narrowing '
    'to just 24 miles at its tightest point, it forms the only maritime exit from the Persian Gulf '
    'to the open ocean. Iran occupies the northern shore; the Omani Musandam Peninsula and a '
    'portion of the United Arab Emirates form the southern boundary.'
)
doc.add_paragraph(
    'Ships follow a traffic separation scheme composed of two-mile-wide inbound and outbound lanes '
    'divided by a two-mile buffer zone. Despite the strait\'s narrow profile, this corridor carries '
    'roughly a fifth of all global oil trade -- making it the single most consequential maritime '
    'chokepoint on Earth.'
)

# Map figure
doc.add_paragraph()
map_para = doc.add_paragraph()
map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
map_para.add_run().add_picture(map_path, width=Inches(5.8))
cap1 = doc.add_paragraph(
    'Figure 1. Schematic map of the Strait of Hormuz and surrounding region. '
    'Arrows indicate inbound and outbound shipping lanes.'
)
cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap1.runs[0].font.size = Pt(9)
cap1.runs[0].font.italic = True
cap1.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

# --- Section 2: History ---
doc.add_heading('A History of Trade', level=1)

doc.add_heading('Ancient and Medieval Periods', level=2)
doc.add_paragraph(
    'The strait\'s commercial importance predates recorded history. The 1st-century CE maritime '
    'handbook Periplus of the Erythraean Sea documents trade routes through the Persian Gulf, '
    'connecting the Roman Empire with India, Arabia, and East Africa. Merchants moved spices, '
    'incense, silk, pearls, and precious stones through these waters for centuries.'
)
doc.add_paragraph(
    'The Kingdom of Ormus -- a powerful trading state centered on Hormuz Island -- dominated the '
    'strait from roughly the 10th through the 17th centuries. At its height, Ormus was one of the '
    'wealthiest trading ports in the world, levying tolls on goods passing between the Indian Ocean '
    'and the Persian Gulf. Pearl diving in the surrounding Gulf waters supported an additional '
    'economy that remained significant well into the 20th century.'
)

doc.add_heading('European Arrival: 15th to 18th Centuries', level=2)
doc.add_paragraph(
    'Portuguese explorer Afonso de Albuquerque seized Hormuz Island in 1507, establishing Iberian '
    'dominance over strait traffic as part of Portugal\'s broader strategy to control the spice '
    'trade. The Portuguese built fortifications and extracted tolls from merchants for over a '
    'century, reshaping the economics of the entire region.'
)
doc.add_paragraph(
    'In 1622, a combined Persian and English force expelled the Portuguese, ending their grip on '
    'the waterway. The British East India Company leveraged this alliance to establish trading '
    'privileges throughout the Gulf, a foothold that expanded into full imperial influence over '
    'the following two centuries. By the 19th century, Britain had signed a series of treaties '
    'with Gulf rulers -- the Trucial States -- cementing its role as the dominant naval and '
    'commercial power in the region.'
)

doc.add_heading('The Oil Age: 20th Century to Present', level=2)
doc.add_paragraph(
    'The discovery of vast petroleum reserves beneath the Gulf basin transformed the strait\'s '
    'economic character entirely. Commercial oil production began in Bahrain in 1932, followed '
    'rapidly by Saudi Arabia, Kuwait, the UAE, and Iraq. Within decades the Strait of Hormuz had '
    'evolved from a spice and pearl route into the world\'s most critical energy corridor.'
)
doc.add_paragraph(
    'The 1973 Arab oil embargo, the 1979 Iranian Revolution, and the 1980 to 1988 Iran-Iraq '
    'Tanker War each demonstrated the strait\'s vulnerability to geopolitical disruption. American '
    'naval presence in the region, formalized through the 1980 Carter Doctrine, was explicitly '
    'designed to guarantee the flow of oil through the waterway. By 2018 daily throughput had '
    'reached 21 million barrels of crude oil, with a total daily cargo value exceeding one '
    'billion dollars.'
)

# --- Section 3: Trade by Industry ---
doc.add_heading('Global Trade by Industry', level=1)
doc.add_paragraph(
    'Today the strait carries a disproportionate share of several global commodity markets. '
    'Crude oil dominates by volume, but the waterway is equally critical for liquefied natural '
    'gas, refined petroleum products, and fertilizers -- industries whose concentration in the '
    'Gulf region has no parallel anywhere else on the planet.'
)

# Chart figure
doc.add_paragraph()
chart_para = doc.add_paragraph()
chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
chart_para.add_run().add_picture(chart_path, width=Inches(5.8))
cap2 = doc.add_paragraph(
    'Figure 2. Estimated share of global seaborne trade transiting the Strait of Hormuz, '
    'by industry. Sources: U.S. Energy Information Administration; '
    'International Energy Agency (2024 estimates).'
)
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap2.runs[0].font.size = Pt(9)
cap2.runs[0].font.italic = True
cap2.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

doc.add_paragraph(
    'Crude oil accounts for approximately 30 percent of all global seaborne crude shipments. '
    'Liquefied natural gas follows at around 20 percent, driven by Qatar\'s position as one of '
    'the world\'s largest LNG exporters. Refined petroleum products add another 10 percent of '
    'global seaborne supply, while petrochemicals and fertilizers -- particularly urea and ammonia '
    'from Saudi Arabia and Qatar -- represent roughly 8 percent of internationally traded volumes. '
    'Over 85 percent of the oil and gas transiting the strait is destined for Asian markets, '
    'primarily China, India, Japan, and South Korea.'
)

doc.add_heading('Strategic Vulnerability', level=2)
doc.add_paragraph(
    'The strait\'s concentration of global energy supply in a single narrow channel creates '
    'systemic economic risk. Extended closure -- whether through military conflict, blockade, or '
    'interdiction -- would immediately affect global oil prices and trigger supply crises for '
    'petroleum-dependent nations with limited strategic reserves. While Saudi Arabia and the UAE '
    'have invested in overland pipeline alternatives, no route currently matches the strait\'s '
    'capacity, ensuring its status as an irreplaceable node in the global energy system for the '
    'foreseeable future.'
)

doc.save('hormuz_trade.docx')
print('Done: hormuz_trade.docx')
